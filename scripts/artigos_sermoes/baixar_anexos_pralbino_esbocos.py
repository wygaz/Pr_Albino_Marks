import argparse
import csv
import datetime
import email
import hashlib
import json
import os
import re
import unicodedata
from email.header import decode_header
from email.utils import parsedate_to_datetime
from io import BytesIO
from pathlib import Path

from imapclient import IMAPClient

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

try:
    from docx import Document
except ImportError:
    Document = None


TEXT_EXT_DEFAULT = [".docx", ".txt"]
TERMOS_DEFAULT = ["esboço", "esboco", "esbo"]


def find_upwards(start: Path, filename: str, max_levels: int = 6) -> Path | None:
    cur = start
    for _ in range(max_levels + 1):
        cand = cur / filename
        if cand.exists():
            return cand
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


def find_repo_root(start: Path, max_levels: int = 8) -> Path | None:
    cur = start
    for _ in range(max_levels + 1):
        if (cur / "manage.py").exists() or (cur / ".git").exists():
            return cur
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


def strip_accents(s: str) -> str:
    s = s or ""
    return "".join(ch for ch in unicodedata.normalize("NFKD", s) if not unicodedata.combining(ch))


def normalize_for_match(s: str) -> str:
    s = strip_accents(s).lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s


def contains_any_term(text: str, termos: list[str]) -> bool:
    norm = normalize_for_match(text)
    return any(normalize_for_match(t) in norm for t in termos)


def parse_data(s: str) -> datetime.date:
    s = s.strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    raise ValueError("Formato invalido. Use AAAA-MM-DD ou DD/MM/AAAA.")


def perguntar_periodo(data_ini_arg: str | None = None, data_fim_arg: str | None = None) -> tuple[datetime.date, datetime.date]:
    if data_ini_arg:
        data_ini = parse_data(data_ini_arg)
    else:
        data_ini = datetime.date(2024, 9, 20)

    if data_fim_arg:
        data_fim = parse_data(data_fim_arg)
    else:
        data_fim = datetime.date.today()

    if data_ini > data_fim:
        raise ValueError("Data inicial maior que a final.")
    return data_ini, data_fim


def nome_unico(path: Path) -> Path:
    if not path.exists():
        return path
    base = path.with_suffix("")
    ext = path.suffix
    i = 2
    while True:
        cand = Path(f"{base}_{i}{ext}")
        if not cand.exists():
            return cand
        i += 1


def next_lote_name(output_base: Path, base_date: datetime.date | None = None) -> str:
    base = (base_date or datetime.date.today()).strftime("%Y_%m_%d")
    i = 1
    while True:
        candidate = f"{base}-{i:02d}"
        if not (output_base / candidate).exists():
            return candidate
        i += 1


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def strip_sm_code(s: str) -> str:
    return re.sub(r"^\s*sm\s*\d+\s*[-–—:]\s*|\s*^\s*sm\s*\d+\s*", "", s or "", flags=re.IGNORECASE).strip()


def strip_prefixos_sujos(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^\s*\d{4}-\d{2}-\d{2}\s*[-–—:]\s*", "", s)
    s = re.sub(r"^\s*\d{2}-\d{2}-\d{4}\s*[-–—:]\s*", "", s)
    s = strip_sm_code(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def safe_filename(name: str, max_len: int = 160) -> str:
    name = strip_prefixos_sujos(name)
    name = re.sub(r'[<>:"/\\|?*]', "", name).strip().rstrip(".")
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "SEM_TITULO"
    return name[:max_len] if len(name) > max_len else name


def decode_filename(raw_filename: str) -> str:
    decoded_parts = decode_header(raw_filename)
    out: list[str] = []
    for chunk, enc in decoded_parts:
        if isinstance(chunk, (bytes, bytearray)):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(chunk)
    return "".join(out)


def decode_mime_header(value: str | None) -> str:
    if not value:
        return ""
    decoded_parts = decode_header(value)
    out: list[str] = []
    for chunk, enc in decoded_parts:
        if isinstance(chunk, (bytes, bytearray)):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(chunk)
    return "".join(out)


def parse_email_datetime(value: str | None) -> tuple[str, str, str]:
    if not value:
        now = datetime.datetime.now(datetime.timezone.utc)
        return "", now.isoformat(), "UTC"

    try:
        dt = parsedate_to_datetime(value)
    except Exception:
        now = datetime.datetime.now(datetime.timezone.utc)
        return value, now.isoformat(), "UTC"

    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=datetime.timezone.utc)

    return value, dt.isoformat(), dt.tzname() or str(dt.tzinfo) or "UTC"


def iso_to_fs_stamp(iso_value: str) -> str:
    try:
        dt = datetime.datetime.fromisoformat(iso_value)
    except Exception:
        dt = datetime.datetime.now(datetime.timezone.utc)
    return dt.strftime("%Y%m%d_%H%M%S")


def build_saved_filename(base_name: str, ext: str, email_dt_iso: str, forcar_prefixo_esboco: bool = True) -> str:
    clean = safe_filename(base_name).upper()
    if forcar_prefixo_esboco and "ESBO" not in normalize_for_match(clean).upper():
        clean = f"ESBOCO - {clean}"
    return f"{iso_to_fs_stamp(email_dt_iso)}__{clean}{ext}"


def docx_best_title_from_bytes(data: bytes) -> str | None:
    if not Document:
        return None
    try:
        doc = Document(BytesIO(data))
        title = (doc.core_properties.title or "").strip()
        title = strip_prefixos_sujos(title)
        if title:
            return title
        for p in doc.paragraphs:
            txt = strip_prefixos_sujos(p.text or "")
            if txt:
                return txt
    except Exception:
        return None
    return None


def set_docx_internal_title(path: Path, title: str) -> None:
    if not Document:
        return
    try:
        doc = Document(str(path))
        doc.core_properties.title = title
        doc.save(str(path))
    except Exception:
        pass


def anexo_eh_esboco(
    filename: str,
    ext: str,
    payload: bytes,
    termos: list[str],
    incluir_assunto: bool,
    assunto_email: str,
) -> tuple[bool, str, str | None]:
    nome_sem_ext = os.path.splitext(filename)[0]

    if contains_any_term(nome_sem_ext, termos):
        return True, "nome_anexo", None

    if incluir_assunto and contains_any_term(assunto_email, termos):
        return True, "assunto_email", None

    if ext == ".docx":
        titulo_docx = docx_best_title_from_bytes(payload)
        if titulo_docx and contains_any_term(titulo_docx, termos):
            return True, "titulo_interno", titulo_docx
        return False, "", titulo_docx

    return False, "", None


def anexo_aceito(
    modo: str,
    filename: str,
    ext: str,
    payload: bytes,
    termos: list[str],
    incluir_assunto: bool,
    assunto_email: str,
) -> tuple[bool, bool, str, str | None]:
    eh_esboco, motivo, titulo_docx = anexo_eh_esboco(
        filename=filename,
        ext=ext,
        payload=payload,
        termos=termos,
        incluir_assunto=incluir_assunto,
        assunto_email=assunto_email,
    )

    if modo == "esbocos":
        return eh_esboco, eh_esboco, motivo, titulo_docx

    if ext == ".docx" and not eh_esboco:
        return True, False, "docx_nao_esboco", titulo_docx

    return False, eh_esboco, "", titulo_docx


def write_manifest(rows: list[dict], manifest_dir: Path) -> tuple[Path, Path]:
    manifest_dir.mkdir(parents=True, exist_ok=True)
    csv_path = manifest_dir / "manifest_anexos.csv"
    json_path = manifest_dir / "manifest_anexos.json"
    fieldnames = [
        "email_datetime_original",
        "email_datetime_iso",
        "timezone",
        "message_id",
        "thread_id",
        "subject",
        "from",
        "attachment_filename_original",
        "saved_filename",
        "saved_path",
        "size_bytes",
        "sha256",
        "downloaded_at",
        "eh_esboco",
        "motivo",
        "titulo_docx",
        "status_download",
        "canonical_sha256",
    ]

    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})

    json_path.write_text(
        json.dumps({"rows": rows, "count": len(rows)}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return csv_path, json_path


def main():
    ap = argparse.ArgumentParser(
        description="Baixa anexos de esboço, registra metadata do e-mail e deduplica por hash mantendo o mais recente."
    )
    ap.add_argument("--ini", default="2024-09-20", help="Data inicial AAAA-MM-DD.")
    ap.add_argument("--fim", default=datetime.date.today().isoformat(), help="Data final AAAA-MM-DD.")
    ap.add_argument("--remetente", default="pralbino@gmail.com")
    ap.add_argument("--so", nargs="*", default=TEXT_EXT_DEFAULT, help="Extensoes permitidas.")
    ap.add_argument("--modo", choices=["esbocos", "artigos"], default="esbocos", help="Seleciona esbocos ou artigos DOCX.")
    ap.add_argument("--termos", nargs="*", default=TERMOS_DEFAULT, help="Termos que caracterizam esboco.")
    ap.add_argument("--incluir-assunto", action="store_true", help="Tambem considera o assunto do email.")
    ap.add_argument("--output-base", default="", help="Default: <repo>\\Apenas_Local\\anexos_filtrados")
    ap.add_argument("--lote", default="", help="Default: data atual com sufixo incremental em AAAA_MM_DD-01")
    args = ap.parse_args()

    scripts_dir = Path(__file__).resolve().parent
    repo_root = find_repo_root(scripts_dir)
    if not repo_root:
        raise RuntimeError("Nao foi possivel localizar a raiz do projeto.")

    output_base = Path(args.output_base).resolve() if args.output_base else repo_root / "Apenas_Local" / "anexos_filtrados"

    if load_dotenv:
        env_path = find_upwards(scripts_dir, ".env.local", max_levels=6)
        if env_path:
            load_dotenv(env_path, override=True)

    email_user = os.getenv("EMAIL_USER") or "wygazeta@gmail.com"
    email_pass = os.getenv("EMAIL_PASS")
    if not email_pass:
        raise RuntimeError("EMAIL_PASS nao encontrado. Coloque no .env.local.")

    extensoes = [e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.so]
    data_ini, data_fim = perguntar_periodo(args.ini, args.fim)

    lote_nome = args.lote.strip() or next_lote_name(output_base, datetime.date.today())
    out_dir = output_base / lote_nome
    dup_dir = out_dir / "_duplicados_hash"
    manifest_dir = out_dir / "_manifest"
    out_dir.mkdir(parents=True, exist_ok=True)
    dup_dir.mkdir(parents=True, exist_ok=True)
    manifest_dir.mkdir(parents=True, exist_ok=True)

    criteria = ["FROM", args.remetente]
    criteria += ["SINCE", data_ini.strftime("%d-%b-%Y")]
    criteria += ["BEFORE", (data_fim + datetime.timedelta(days=1)).strftime("%d-%b-%Y")]

    print("\nFiltro IMAP:", criteria)
    print("Saida:", out_dir)
    print("Modo:", args.modo)
    print("Termos de busca:", args.termos)
    print("Considerar assunto do email:", "SIM" if args.incluir_assunto else "NAO")

    emails_encontrados = 0
    anexos_avaliados = 0
    anexos_selecionados = 0
    mantidos = 0
    duplicados_hash = 0
    registros: list[dict] = []

    with IMAPClient("imap.gmail.com") as server:
        server.login(email_user, email_pass)
        server.select_folder("INBOX", readonly=True)

        messages = server.search(criteria)
        emails_encontrados = len(messages)
        print(f"\nEmails encontrados: {emails_encontrados}")
        fetched = server.fetch(messages, ["RFC822", "X-GM-THRID"])

        for uid, message_data in fetched.items():
            email_message = email.message_from_bytes(message_data[b"RFC822"])
            assunto_email = decode_mime_header(email_message.get("Subject", ""))
            email_dt_original, email_dt_iso, timezone_name = parse_email_datetime(email_message.get("Date"))
            message_id = (email_message.get("Message-ID") or "").strip()
            thread_raw = message_data.get(b"X-GM-THRID")
            thread_id = str(thread_raw) if thread_raw is not None else ""
            from_header = decode_mime_header(email_message.get("From", ""))

            for part in email_message.walk():
                if part.get_content_maintype() == "multipart":
                    continue
                if part.get("Content-Disposition") is None:
                    continue

                raw = part.get_filename()
                if not raw:
                    continue

                filename = decode_filename(raw)
                ext = os.path.splitext(filename)[1].lower()
                if ext not in extensoes:
                    continue

                anexos_avaliados += 1
                payload = part.get_payload(decode=True) or b""
                aceito, eh_esboco, motivo, titulo_docx = anexo_aceito(
                    modo=args.modo,
                    filename=filename,
                    ext=ext,
                    payload=payload,
                    termos=args.termos,
                    incluir_assunto=args.incluir_assunto,
                    assunto_email=assunto_email,
                )
                if not aceito:
                    continue

                anexos_selecionados += 1
                nome_sem_ext = os.path.splitext(filename)[0]
                base_name = titulo_docx or nome_sem_ext
                saved_ext = ".txt" if ext == ".txt" else ext
                saved_filename = build_saved_filename(
                    base_name,
                    saved_ext,
                    email_dt_iso,
                    forcar_prefixo_esboco=(args.modo == "esbocos"),
                )

                registros.append(
                    {
                        "email_datetime_original": email_dt_original,
                        "email_datetime_iso": email_dt_iso,
                        "timezone": timezone_name,
                        "message_id": message_id,
                        "thread_id": thread_id,
                        "subject": assunto_email,
                        "from": from_header,
                        "attachment_filename_original": filename,
                        "saved_filename": saved_filename,
                        "saved_path": "",
                        "size_bytes": len(payload),
                        "sha256": sha256_bytes(payload),
                        "downloaded_at": datetime.datetime.now(datetime.timezone.utc).isoformat(),
                        "eh_esboco": eh_esboco,
                        "motivo": motivo,
                        "titulo_docx": titulo_docx or "",
                        "status_download": "candidato",
                        "canonical_sha256": "",
                        "_payload": payload,
                    }
                )

    grupos_por_hash: dict[str, list[dict]] = {}
    for reg in registros:
        grupos_por_hash.setdefault(reg["sha256"], []).append(reg)

    for sha, itens in grupos_por_hash.items():
        ordenados = sorted(
            itens,
            key=lambda x: (x["email_datetime_iso"], x["downloaded_at"], x["attachment_filename_original"]),
            reverse=True,
        )
        principal = ordenados[0]
        principal["status_download"] = "mantido"
        principal["canonical_sha256"] = sha
        out_path = nome_unico(out_dir / principal["saved_filename"])
        with out_path.open("wb") as f:
            f.write(principal["_payload"])
        principal["saved_path"] = str(out_path)
        if out_path.suffix.lower() == ".docx":
            set_docx_internal_title(out_path, Path(principal["saved_filename"]).stem)
        mantidos += 1
        print(f"Salvo principal: {out_path.name}   [email={principal['email_datetime_iso']}]")

        for duplicado in ordenados[1:]:
            duplicado["status_download"] = "duplicado_hash"
            duplicado["canonical_sha256"] = sha
            dup_path = nome_unico(dup_dir / duplicado["saved_filename"])
            with dup_path.open("wb") as f:
                f.write(duplicado["_payload"])
            duplicado["saved_path"] = str(dup_path)
            if dup_path.suffix.lower() == ".docx":
                set_docx_internal_title(dup_path, Path(duplicado["saved_filename"]).stem)
            duplicados_hash += 1
            print(
                f"Duplicado hash: {duplicado['attachment_filename_original']}   "
                f"[principal={principal['saved_filename']}]"
            )

    for reg in registros:
        reg.pop("_payload", None)

    manifest_csv, manifest_json = write_manifest(
        sorted(registros, key=lambda x: (x["attachment_filename_original"], x["email_datetime_iso"], x["sha256"])),
        manifest_dir,
    )

    print("\nFinalizado.")
    print(f"Emails encontrados : {emails_encontrados}")
    print(f"Anexos avaliados   : {anexos_avaliados}")
    print(f"Anexos selecionados: {anexos_selecionados}")
    print(f"Arquivos mantidos  : {mantidos}")
    print(f"Duplicados hash    : {duplicados_hash}")
    print(f"Pasta de saida     : {out_dir}")
    print(f"Manifest CSV       : {manifest_csv}")
    print(f"Manifest JSON      : {manifest_json}")


if __name__ == "__main__":
    main()
