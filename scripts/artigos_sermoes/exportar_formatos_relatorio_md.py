import argparse
import html
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SECTION_RE = re.compile(r"^(\d+)\)\s+(.*)$")
SITE_URL = "https://www.albinomarks.com.br"
BRAND_NAVY = RGBColor(0x1D, 0x33, 0x4A)
BRAND_GOLD = RGBColor(0xC9, 0xA2, 0x27)
BODY_INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x5B, 0x6B, 0x7A)


def ascii_slug(text: str) -> str:
    text = unicodedata.normalize("NFKD", text or "")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-").lower()
    return text or "relatorio-tecnico"


def clean_base_name(md_path: Path | str) -> str:
    value = Path(md_path).stem
    value = re.sub(r"__dossie$", "", value, flags=re.IGNORECASE)
    value = re.sub(r"__relatorio_tecnico__.*$", "", value, flags=re.IGNORECASE)
    value = re.sub(r"^\d+__", "", value)
    return value.replace("__", " ").strip()


def canonical_base_stem(md_path: Path) -> str:
    raw = md_path.stem
    match = re.match(r"^(?P<prefix>\d{1,3})__(?P<body>.+)$", raw)
    prefix = f"{int(match.group('prefix')):02d}__" if match else ""
    body = match.group("body") if match else raw
    return f"{prefix}{ascii_slug(clean_base_name(body))}"


def parse_relatorio_md(text: str, md_path: Path) -> Dict:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    sections: List[Tuple[str, List[str]]] = []
    current_title = None
    current_lines: List[str] = []

    for raw in lines:
        line = raw.rstrip()
        match = SECTION_RE.match(line.strip())
        if match:
            if current_title is not None:
                sections.append((current_title, current_lines))
            current_title = match.group(2).strip()
            current_lines = []
        else:
            if current_title is None:
                continue
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, current_lines))

    section_map = {title.lower(): [item.strip() for item in body if item.strip()] for title, body in sections}
    identificacao = section_map.get("identificação do artigo", []) or section_map.get("identificacao do artigo", [])

    def pick(prefix: str) -> str:
        prefix_lower = prefix.lower()
        for item in identificacao:
            normalized = item.lstrip("-").strip()
            if normalized.lower().startswith(prefix_lower):
                return normalized.split(":", 1)[1].strip() if ":" in normalized else normalized
        return ""

    article_title = pick("título") or pick("titulo") or clean_base_name(md_path)
    tema = pick("tema central") or ""
    article_type = pick("tipo de artigo") or ""
    author = pick("autor") or "Pr. Albino Marks"
    return {
        "article_title": article_title,
        "title": f"Relatório técnico - {article_title}",
        "tema": tema,
        "article_type": article_type,
        "author": author,
        "sections": sections,
    }


def md_lines_to_html(lines: List[str]) -> str:
    chunks: List[str] = []
    in_list = False
    for line in lines:
        s = line.strip()
        if not s:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            continue
        if s.startswith("- "):
            if not in_list:
                chunks.append('<ul class="bullet-list">')
                in_list = True
            chunks.append(f"<li>{html.escape(s[2:].strip())}</li>")
        else:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            chunks.append(f"<p>{html.escape(s)}</p>")
    if in_list:
        chunks.append("</ul>")
    return "\n".join(chunks)


def build_html(data: Dict) -> str:
    sections_html = []
    for title_sec, body in data["sections"]:
        sections_html.append(
            f'<section class="section-block"><h2>{html.escape(title_sec)}</h2>{md_lines_to_html(body)}</section>'
        )
    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{html.escape(data["title"])}</title>
<style>
@page {{ size: A4; margin: 18mm 16mm 18mm 16mm; }}
:root {{
  --navy:#1d334a;
  --gold:#c9a227;
  --paper:#fffdf7;
  --edge:#d7dde4;
  --muted:#607080;
  --ink:#1f2937;
}}
body {{
  margin:0;
  background:#ffffff;
  color:var(--ink);
  font-family: Georgia, "Times New Roman", serif;
}}
.page-shell {{
  max-width:100%;
  margin:0 auto;
}}
.masthead {{
  background:linear-gradient(135deg, #1d334a 0%, #2f4f4f 100%);
  color:#f7efcf;
  border-bottom:4px solid var(--gold);
  border-radius:0 0 18px 18px;
  padding:16px 20px 18px;
  margin-bottom:18px;
}}
.brand {{
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.76rem;
  text-transform:uppercase;
  letter-spacing:.18em;
  opacity:.9;
}}
.masthead h1 {{
  margin:8px 0 10px;
  line-height:1.12;
  color:#fff8da;
  font-size:22pt;
}}
.masthead-sub {{
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.92rem;
  color:#d9e4ea;
  max-width:60ch;
}}
.meta-grid {{
  display:grid;
  grid-template-columns:repeat(3, minmax(0,1fr));
  gap:12px;
  margin-top:16px;
}}
.meta-card {{
  background:rgba(255,255,255,.1);
  border:1px solid rgba(255,255,255,.16);
  border-radius:14px;
  padding:12px 14px;
}}
.meta-label {{
  display:block;
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.74rem;
  text-transform:uppercase;
  letter-spacing:.12em;
  color:#dfe8ec;
  margin-bottom:5px;
}}
.meta-value {{
  font-size:1rem;
  line-height:1.45;
  color:#fff;
}}
.reading-surface {{
  background:var(--paper);
  border:1px solid var(--edge);
  border-radius:22px;
  box-shadow:0 10px 30px rgba(20,30,40,.08);
  padding:22px 26px 26px;
}}
.section-block + .section-block {{
  margin-top:20px;
}}
.section-block h2 {{
  font-size:1.06rem;
  color:var(--navy);
  margin:0 0 10px;
  padding-top:10px;
  border-top:1px solid #e7ebef;
  font-family:"Segoe UI", Arial, sans-serif;
}}
.section-block p {{
  margin:.5rem 0 .92rem;
  line-height:1.58;
  font-size:11.2pt;
}}
.bullet-list {{
  margin:.45rem 0 .95rem 1.2rem;
  padding:0;
}}
.bullet-list li {{
  margin:.32rem 0;
  line-height:1.52;
  font-size:11pt;
}}
.footer {{
  margin-top:16px;
  padding:10px 4px 0;
  border-top:1px solid #dbe3ea;
  color:var(--muted);
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.84rem;
  display:flex;
  justify-content:space-between;
  gap:12px;
}}
.footer a {{
  color:var(--navy);
  text-decoration:none;
  font-weight:600;
}}
</style>
</head>
<body>
  <div class="page-shell">
    <header class="masthead">
      <div class="brand">Pr. Albino Marks | Apoio ao estudo</div>
      <h1>{html.escape(data["title"])}</h1>
      <div class="masthead-sub">Documento técnico para apoio hermenêutico, doutrinário e argumentativo da série editorial.</div>
      <div class="meta-grid">
        <div class="meta-card"><span class="meta-label">Artigo-base</span><div class="meta-value">{html.escape(data["article_title"])}</div></div>
        <div class="meta-card"><span class="meta-label">Autor</span><div class="meta-value">{html.escape(data["author"])}</div></div>
        <div class="meta-card"><span class="meta-label">Tipo</span><div class="meta-value">{html.escape(data["article_type"] or "Relatório teológico técnico")}</div></div>
      </div>
    </header>
    <main class="reading-surface">
      {''.join(sections_html)}
      <div class="footer">
        <span><a href="{SITE_URL}">{SITE_URL}</a></span>
        <span>Relatório técnico | Uso editorial interno</span>
      </div>
    </main>
  </div>
</body>
</html>
"""


def set_cell_shading(cell, fill: str = "F8F2E4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def ensure_style(doc: Document, name: str, style_type, *, font_name="Calibri", font_size=11, bold=False, color=None):
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, style_type)
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    return style


def write_docx(data: Dict, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    ensure_style(doc, "RelBody", WD_STYLE_TYPE.PARAGRAPH, font_name="Georgia", font_size=11, color=BODY_INK)
    ensure_style(doc, "RelHeading", WD_STYLE_TYPE.PARAGRAPH, font_name="Segoe UI", font_size=12, bold=True, color=BRAND_NAVY)

    header = section.header.paragraphs[0]
    header.text = f"Pr. Albino Marks | Relatório técnico | {SITE_URL}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = f"{SITE_URL} | Apoio ao estudo"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["title"])
    run.bold = True
    run.font.name = "Segoe UI Semibold"
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Documento técnico para apoio hermenêutico, doutrinário e argumentativo.")
    subrun.font.name = "Segoe UI"
    subrun.font.size = Pt(10)
    subrun.font.color.rgb = MUTED

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = f"Artigo-base\n{data['article_title']}"
    hdr[1].text = f"Autor\n{data['author']}"
    hdr[2].text = f"Tipo\n{data['article_type'] or 'Relatório teológico técnico'}"
    for cell in hdr:
        set_cell_shading(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = BRAND_NAVY

    doc.add_paragraph("")
    for title_sec, body in data["sections"]:
        heading = doc.add_paragraph(style="RelHeading")
        heading.add_run(title_sec)
        for line in body:
            s = line.strip()
            if not s:
                continue
            if s.startswith("- "):
                p = doc.add_paragraph(style="RelBody")
                p.style = doc.styles["List Bullet"]
                p.add_run(s[2:].strip())
            else:
                p = doc.add_paragraph(style="RelBody")
                p.add_run(s)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser(description="Exporta formatos editoriais a partir de um relatório técnico em Markdown.")
    parser.add_argument("--md", required=True)
    parser.add_argument("--outdir", required=True)
    args = parser.parse_args()

    md_path = Path(args.md)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    text = md_path.read_text(encoding="utf-8", errors="ignore")
    data = parse_relatorio_md(text, md_path)
    stem = canonical_base_stem(md_path)

    html_path = outdir / f"{stem}__dossie__a4.html"
    docx_path = outdir / f"{stem}__dossie__a4.docx"

    html_path.write_text(build_html(data), encoding="utf-8")
    write_docx(data, docx_path)

    print(f"[OK] {html_path}")
    print(f"[OK] {docx_path}")


if __name__ == "__main__":
    main()
