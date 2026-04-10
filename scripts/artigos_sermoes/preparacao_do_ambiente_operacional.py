from __future__ import annotations

import argparse
import csv
import datetime
import json
import re
import subprocess
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document


def strip_accents(text: str) -> str:
    return "".join(
        ch
        for ch in unicodedata.normalize("NFKD", text or "")
        if not unicodedata.combining(ch)
    )


def strip_editorial_prefixes(text: str) -> str:
    value = str(text or "").strip()
    value = re.sub(r"^\d{8}_\d{6}__", "", value)
    value = re.sub(r"^\d{2}__", "", value)
    value = re.sub(r"^\d+\s*(?:[-_. ]+\s*)?", "", value)
    value = re.sub(r"^[^A-Za-z0-9]+", "", value)
    while True:
        updated = re.sub(r"^(?:[A-Za-z]{1,5}\d{2,6})(?:\s*[-_.:)\]]*\s*)", "", value)
        if updated == value:
            break
        value = updated.lstrip()
        value = re.sub(r"^[^A-Za-z0-9]+", "", value)
    return value.strip()


def normalize(text: str) -> str:
    text = strip_accents(strip_editorial_prefixes(text or "")).upper()
    text = text.replace("_", " ")
    text = re.sub(r"\.DOCX?$", "", text, flags=re.I)
    text = re.sub(r"[^A-Z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def sanitize_piece(text: str) -> str:
    text = strip_editorial_prefixes(text)
    text = text.replace("_", " ")
    text = re.sub(r"[â€œâ€\"']", "", text)
    text = text.replace(",", "")
    text = re.sub(r"\s+-\s+\d+$", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r'[<>:/\\|?*]', "", text)
    return text


def canonical_display_piece(text: str) -> str:
    text = strip_editorial_prefixes(text)
    text = text.replace("_", " ")
    text = re.sub(r"\s+-\s+\d+$", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def ascii_slug(text: str) -> str:
    value = strip_accents(text or "").lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value or "documento"


def technical_name(ordem: int, titulo: str) -> str:
    return f"{ordem:02d}__{ascii_slug(sanitize_piece(titulo))}"


def canonical_article_filename(ordem: int, titulo: str) -> str:
    return f"{technical_name(ordem, titulo)}.docx"


def docx_internal_title(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception:
        return ""

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            return text
    try:
        title = (doc.core_properties.title or "").strip()
        if title:
            return title
    except Exception:
        pass
    return ""


def filename_title(path: Path) -> str:
    return sanitize_piece(path.stem)


def source_title_candidates(path: Path) -> list[str]:
    internal = sanitize_piece(docx_internal_title(path))
    by_name = filename_title(path)
    out: list[str] = []
    for value in [internal, by_name]:
        if value and value not in out:
            out.append(value)
    return out


@dataclass
class CanonicalItem:
    serie_id: int
    serie_dir_name: str
    ordem: int
    titulo_esboco: str
    norm: str


def preferred_canonical_title(path: Path, item: CanonicalItem) -> str:
    internal = canonical_display_piece(docx_internal_title(path))
    if internal:
        return internal
    return canonical_display_piece(item.titulo_esboco)


def title_divergence_alert(path: Path, item: CanonicalItem) -> str:
    internal = sanitize_piece(docx_internal_title(path))
    if not internal:
        return ""
    aliases = alias_map()
    internal_norm = aliases.get(normalize(internal), normalize(internal))
    esboco_norm = aliases.get(normalize(item.titulo_esboco), normalize(item.titulo_esboco))
    if internal_norm != esboco_norm:
        return f"TITULO_INTERNO_DIVERGENTE: DOCX='{internal}' | ESBOCO='{sanitize_piece(item.titulo_esboco)}'"
    return ""


def extract_series_from_outline(esboco_path: Path) -> dict[int, dict]:
    doc = Document(str(esboco_path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    def find_idx(*terms: str, start: int = 0) -> int:
        targets = [normalize(t) for t in terms]
        for idx in range(start, len(paras)):
            current = normalize(paras[idx])
            if all(term in current for term in targets):
                return idx
        raise ValueError(f"Cabecalho nao encontrado: {terms}")

    idx1 = find_idx("grande conflito", "cristo", "satanas")
    idx2 = find_idx("apocalipse 17", "principio do simbolismo")
    idx3 = find_idx("torah", "nomos", "mandamentos", "ordenancas", "graca")
    idx4_all = [
        i for i, text in enumerate(paras)
        if "BIBLIA" in normalize(text) and "HISTORIA DA HUMANIDADE" in normalize(text)
    ]
    if len(idx4_all) < 2:
        raise ValueError("Nao foi possivel localizar a Serie 4 no esboco.")
    idx4_first = idx4_all[0]
    idx4_second = idx4_all[1]

    blocks = {
        1: paras[idx1 : idx2],
        2: paras[idx2 : idx3],
        3: paras[idx3:idx4_first],
        4: paras[idx4_second:],
    }

    series: dict[int, dict] = {}
    for serie_id, lines in blocks.items():
        lines = [line for line in lines if normalize(line) != "ESBOCO"]
        if not lines:
            raise ValueError(f"Serie {serie_id} vazia no esboco.")
        serie_title = sanitize_piece(lines[0])
        titles = list(lines)
        if len(titles) >= 2 and normalize(titles[0]) == normalize(titles[1]):
            titles = [titles[0]] + titles[2:]
        series[serie_id] = {
            "serie_id": serie_id,
            "dir_name": f"Serie_{serie_id}__{ascii_slug(serie_title)}",
            "first_title": serie_title,
            "titles": titles,
        }
    return series


def alias_map() -> dict[str, str]:
    return {
        "O DILUVIO": "O DILUVIO E A ARCA DE NOE",
        "ORIGEM DO GRANDE CONFLITO COSMICO DE CONCEITOS ESPIRITUAIS": "A ORIGEM DO GRANDE CONFLITO DE CONCEITOS ESPIRITUAIS",
        "O PRINCIPIO BIBLICO DO SIMBOLISMO": "O PRINCIPIO DO SIMBOLISMO",
        "O CONFLITO COSMICO E OS DOIS PODERES EM CONFRONTO": "O CONFLITO E OS DOIS PODERES EM CONFRONTO",
        "A RECONQUISTA DO DOMINIO PERDIDO": "A RECONQUISTA DO DOMINIO",
        "A PRIMEIRA TROMBETA E A PRIMEIRA PRAGA": "A PRIMEIRA TROMBETA E A PRIMEIRO PRAGA",
        "A QUINTA TROMBETA E O SELO DE DEUS": "A QUINTA PRAGA E O SELO DE DEUS",
        "SETIMA TROMBETA E A SETIMA PRAGA": "SETIMA TROMBETA E SETIMA PRAGA",
        "A QUEM SIMBOLIZA A BESTA NAO QUALIFICADA": "QUEM E A BESTA NAO QUALIFICADA",
        "A BESTA QUE ERA E AGORA NAO E E O OITAVO REI": "A BESTA QUE ERA NAO E E O OITAVO REI",
        "A BATALHA FINAL DO ARMAGEDOM": "A BATALHA FINAL D ARMAGEDOM",
        "A TORAH NOMOS LEI MANDAMENTOS ORDENANCAS E GRACA": "A LEI TORAH NOMOS MANDAMENTOS ORDENANCAS E A GRACA",
        "A TORAH NOMOS LEI MANDAMENTOS ORDENANCAS E A GRACA": "A LEI TORAH NOMOS MANDAMENTOS ORDENANCAS E A GRACA",
        "O NOVO TESTAMENTO JESUS A LEI E OS PROFETAS": "O NOVO TESTAMENTO JESUS E A LEI",
        "JESUS E A LEI NOMOS": "JESUS E A LEI NOMOS",
        "JESUS E A LEI (NOMOS)": "JESUS E A LEI NOMOS",
        "NINGUEM E JUSTIFICADO POR OBRAS DA LEI": "NINGUEM E JUSTIFICADO PELAS OBRAS DA LEI",
        "PAULO E A GRACA DE DEUS": "O APOSTOLO PAULO E A GRACA DE DEUS",
        "MALDICAO DA LEI": "A MALDICAO DA LEI",
        "A LEI MORAL E AS ALIANCAS DE DEUS": "A LEI MORAL E AS ALIANCAS",
        "A LEI NAO REVOGA A ALIANCA": "A LEI NAO ANULA AS ALIANCAS",
        "A LEI ACRESCENTADA POR CAUSA DAS TRANSGRESSOES": "POR CAUSA DAS TRANSGRESSOES",
        "A DIVISAO DA HISTORIA DA HUMANIDADE PREDITA": "A DIVISAO DA HISTORIA DA HUMANIDADE PREDITA",
        "OS MIL DUZENTOS E SESSENTA ANOS E SATANAS": "OS MIL DUZENTOS SESSENTA ANOS E SATANAS",
    }


def source_stamp(path: Path) -> datetime.datetime:
    match = re.match(r"^(?P<date>\d{8}_\d{6})__", path.name)
    if not match:
        return datetime.datetime.min
    try:
        return datetime.datetime.strptime(match.group("date"), "%Y%m%d_%H%M%S")
    except ValueError:
        return datetime.datetime.min


def canonical_index(series: dict[int, dict]) -> dict[str, CanonicalItem]:
    idx: dict[str, CanonicalItem] = {}
    for serie_id, data in series.items():
        for ordem, titulo in enumerate(data["titles"], start=1):
            norm = normalize(titulo)
            idx[norm] = CanonicalItem(
                serie_id=serie_id,
                serie_dir_name=data["dir_name"],
                ordem=ordem,
                titulo_esboco=sanitize_piece(titulo),
                norm=norm,
            )
    return idx


def match_candidate(candidates: list[str], idx: dict[str, CanonicalItem], cutoff: float) -> tuple[CanonicalItem | None, str, float]:
    aliases = alias_map()
    best_item: CanonicalItem | None = None
    best_method = ""
    best_score = 0.0

    for raw in candidates:
        norm = normalize(raw)
        if not norm:
            continue
        norm = aliases.get(norm, norm)
        if norm in idx:
            return idx[norm], "exato", 1.0
        for key, item in idx.items():
            score = SequenceMatcher(None, norm, key).ratio()
            if score > best_score:
                best_item = item
                best_score = score
                best_method = "aproximado"

    if best_item and best_score >= cutoff:
        return best_item, best_method, best_score
    return None, "", 0.0


def iter_source_files(source_dir: Path, extra_dirs: list[Path]) -> list[Path]:
    files = sorted(source_dir.glob("*.docx"))
    for extra in extra_dirs:
        files.extend(sorted(extra.glob("*.docx")))
    unique: list[Path] = []
    seen: set[Path] = set()
    for path in files:
        if path not in seen:
            seen.add(path)
            unique.append(path)
    return unique


def ensure_dirs(output_dir: Path, series: dict[int, dict]) -> dict[int, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    serie_dirs: dict[int, Path] = {}
    for serie_id, data in series.items():
        serie_dir = output_dir / data["dir_name"]
        serie_dir.mkdir(parents=True, exist_ok=True)
        serie_dirs[serie_id] = serie_dir
    return serie_dirs


def ensure_support_dirs(output_dir: Path) -> tuple[Path, Path, Path, Path]:
    support_root = output_dir.parent / "_preparacao_operacional"
    nao_classificados = support_root / "Nao_Classificados"
    conflitos = support_root / "Conflitos_Titulo"
    resolvidos = support_root / "Duplicados_Resolvidos"
    report_dir = support_root / "_relatorios_preparacao"
    nao_classificados.mkdir(parents=True, exist_ok=True)
    conflitos.mkdir(parents=True, exist_ok=True)
    resolvidos.mkdir(parents=True, exist_ok=True)
    report_dir.mkdir(parents=True, exist_ok=True)
    return nao_classificados, conflitos, resolvidos, report_dir


def archive_dir_if_nonempty(path: Path) -> Path | None:
    if not path.exists() or not path.is_dir():
        return None
    try:
        next(path.iterdir())
    except StopIteration:
        return None
    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    archived = path.with_name(f"{path.name}_old_{stamp}")
    counter = 1
    while archived.exists():
        counter += 1
        archived = path.with_name(f"{path.name}_old_{stamp}_{counter}")
    shutil.move(str(path), str(archived))
    return archived


def copy_or_move(src: Path, dst: Path, move: bool) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    if move:
        shutil.move(str(src), str(dst))
    else:
        shutil.copy2(str(src), str(dst))


def resolve_single_article_target(
    arquivo: Path,
    outline: Path,
    output_dir: Path,
    cutoff: float,
) -> tuple[CanonicalItem, str, float, str, Path]:
    if not arquivo.exists():
        raise FileNotFoundError(f"Arquivo de reposicao nao encontrado: {arquivo}")
    if not outline.exists():
        raise FileNotFoundError(f"Esboco nao encontrado: {outline}")

    series = extract_series_from_outline(outline)
    idx = canonical_index(series)
    serie_dirs = ensure_dirs(output_dir, series)
    candidates = source_title_candidates(arquivo)
    item, method, score = match_candidate(candidates, idx, cutoff=cutoff)
    if not item:
        raise RuntimeError(f"Nao foi possivel encaixar o artigo: {arquivo.name}")

    titulo_final = preferred_canonical_title(arquivo, item)
    new_name = canonical_article_filename(item.ordem, titulo_final)
    dst = serie_dirs[item.serie_id] / new_name
    return item, method, score, titulo_final, dst


def run_python_step(script_path: Path, extra_args: list[str]) -> None:
    cmd = [sys.executable, str(script_path), *extra_args]
    proc = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if proc.stdout:
        print(proc.stdout.rstrip())
    if proc.returncode != 0:
        detail = (proc.stderr or proc.stdout or "").strip()
        raise RuntimeError(f"Falha em {script_path.name}: {detail}")


def scripts_dir_from_here() -> Path:
    return Path(__file__).resolve().parent


def run_organizar_lote(args: argparse.Namespace) -> int:
    source_dir = Path(args.artigos_dir).resolve()
    outline = Path(args.esboco).resolve()
    output_dir = Path(args.saida).resolve()
    extra_dirs = [Path(item).resolve() for item in args.complementar_dir]

    if not source_dir.exists():
        raise FileNotFoundError(f"Pasta de artigos nao encontrada: {source_dir}")
    if not outline.exists():
        raise FileNotFoundError(f"Esboco nao encontrado: {outline}")

    archived_output = archive_dir_if_nonempty(output_dir)

    series = extract_series_from_outline(outline)
    idx = canonical_index(series)
    serie_dirs = ensure_dirs(output_dir, series)
    nao_classificados, conflitos, resolvidos, report_dir = ensure_support_dirs(output_dir)

    report_rows: list[dict] = []
    slot_owners: dict[tuple[int, int], dict] = {}

    for src in iter_source_files(source_dir, extra_dirs):
        candidates = source_title_candidates(src)
        item, method, score = match_candidate(candidates, idx, cutoff=args.cutoff)
        if not item:
            dst = nao_classificados / sanitize_piece(src.name)
            copy_or_move(src, dst, args.mover)
            report_rows.append(
                {
                    "arquivo_origem": str(src),
                    "status": "nao_classificado",
                    "serie_id": "",
                    "ordem": "",
                    "titulo_canonico": "",
                    "titulo_detectado": " | ".join(candidates),
                    "destino": str(dst),
                    "metodo": "",
                    "score": "",
                }
            )
            continue

        slot_key = (item.serie_id, item.ordem)
        titulo_final = preferred_canonical_title(src, item)
        new_name = canonical_article_filename(item.ordem, titulo_final)
        dest_dir = serie_dirs[item.serie_id]
        dst = dest_dir / new_name

        if slot_key in slot_owners:
            current = slot_owners[slot_key]
            current_stamp = current["stamp"]
            incoming_stamp = source_stamp(src)
            if incoming_stamp > current_stamp:
                old_dst = Path(current["destino"])
                resolved_old = resolvidos / f"{item.serie_id:02d}_{item.ordem:02d}__{sanitize_piece(old_dst.name)}"
                shutil.move(str(old_dst), str(resolved_old))
                report_rows.append(
                    {
                        "arquivo_origem": current["arquivo_origem"],
                        "status": "duplicado_resolvido_antigo",
                        "serie_id": item.serie_id,
                        "ordem": item.ordem,
                        "titulo_canonico": current["titulo_canonico"],
                        "nome_tecnico": current.get("nome_tecnico", ""),
                        "titulo_detectado": current["titulo_detectado"],
                        "destino": str(resolved_old),
                        "metodo": current["metodo"],
                        "score": current["score"],
                        "alerta_titulo": "",
                    }
                )
                copy_or_move(src, dst, args.mover)
                alerta_titulo = title_divergence_alert(src, item)
                slot_owners[slot_key] = {
                    "arquivo_origem": str(src),
                    "destino": str(dst),
                    "stamp": incoming_stamp,
                    "titulo_canonico": titulo_final,
                    "nome_tecnico": Path(new_name).stem,
                    "titulo_detectado": " | ".join(candidates),
                    "metodo": method,
                    "score": f"{score:.3f}",
                    "alerta_titulo": alerta_titulo,
                }
                report_rows.append(
                    {
                        "arquivo_origem": str(src),
                        "status": "classificado_mais_recente",
                        "serie_id": item.serie_id,
                        "ordem": item.ordem,
                        "titulo_canonico": titulo_final,
                        "nome_tecnico": Path(new_name).stem,
                        "titulo_detectado": " | ".join(candidates),
                        "destino": str(dst),
                        "metodo": method,
                        "score": f"{score:.3f}",
                        "alerta_titulo": alerta_titulo,
                    }
                )
                continue

            resolved_new = resolvidos / f"{item.serie_id:02d}_{item.ordem:02d}__{sanitize_piece(src.name)}"
            copy_or_move(src, resolved_new, args.mover)
            report_rows.append(
                {
                    "arquivo_origem": str(src),
                    "status": "duplicado_resolvido_antigo",
                    "serie_id": item.serie_id,
                    "ordem": item.ordem,
                    "titulo_canonico": titulo_final,
                    "nome_tecnico": Path(new_name).stem,
                    "titulo_detectado": " | ".join(candidates),
                    "destino": str(resolved_new),
                    "metodo": method,
                    "score": f"{score:.3f}",
                    "alerta_titulo": "",
                }
            )
            continue

        copy_or_move(src, dst, args.mover)
        alerta_titulo = title_divergence_alert(src, item)
        slot_owners[slot_key] = {
            "arquivo_origem": str(src),
            "destino": str(dst),
            "stamp": source_stamp(src),
            "titulo_canonico": titulo_final,
            "nome_tecnico": Path(new_name).stem,
            "titulo_detectado": " | ".join(candidates),
            "metodo": method,
            "score": f"{score:.3f}",
            "alerta_titulo": alerta_titulo,
        }
        report_rows.append(
            {
                "arquivo_origem": str(src),
                "status": "classificado",
                "serie_id": item.serie_id,
                "ordem": item.ordem,
                "titulo_canonico": titulo_final,
                "nome_tecnico": Path(new_name).stem,
                "titulo_detectado": " | ".join(candidates),
                "destino": str(dst),
                "metodo": method,
                "score": f"{score:.3f}",
                "alerta_titulo": alerta_titulo,
            }
        )

    missing_rows: list[dict] = []
    for serie_id, data in series.items():
        for ordem, titulo in enumerate(data["titles"], start=1):
            if (serie_id, ordem) not in slot_owners:
                missing_rows.append(
                    {
                        "serie_id": serie_id,
                        "ordem": ordem,
                        "titulo_canonico": canonical_display_piece(titulo),
                        "nome_tecnico": technical_name(ordem, titulo),
                    }
                )

    csv_path = report_dir / "preparacao_lote.csv"
    json_path = report_dir / "preparacao_lote.json"
    faltantes_path = report_dir / "faltantes.csv"
    alertas_path = report_dir / "alertas_titulo.csv"

    with csv_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "arquivo_origem",
                "status",
                "serie_id",
                "ordem",
                "titulo_canonico",
                "nome_tecnico",
                "titulo_detectado",
                "destino",
                "metodo",
                "score",
                "alerta_titulo",
            ],
        )
        writer.writeheader()
        writer.writerows(report_rows)

    with faltantes_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=["serie_id", "ordem", "titulo_canonico", "nome_tecnico"])
        writer.writeheader()
        writer.writerows(missing_rows)

    alertas_rows = []
    for (serie_id, ordem), owner in sorted(slot_owners.items()):
        alerta = owner.get("alerta_titulo", "")
        if not alerta:
            continue
        m = re.match(r"TITULO_INTERNO_DIVERGENTE: DOCX='(?P<docx>.*?)' \| ESBOCO='(?P<esboco>.*?)'$", alerta)
        alertas_rows.append(
            {
                "arquivo": Path(str(owner.get("arquivo_origem", ""))).name,
                "caminho_arquivo": str(owner.get("arquivo_origem", "")),
                "serie_id": serie_id,
                "ordem": ordem,
                "titulo_docx": m.group("docx") if m else "",
                "titulo_canonico": owner.get("titulo_canonico", ""),
                "nome_tecnico": owner.get("nome_tecnico", ""),
                "titulo_esboco": m.group("esboco") if m else "",
            }
        )

    with alertas_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=["arquivo", "caminho_arquivo", "serie_id", "ordem", "titulo_docx", "titulo_canonico", "nome_tecnico", "titulo_esboco"],
        )
        writer.writeheader()
        writer.writerows(alertas_rows)

    json_path.write_text(
        json.dumps(
            {
                "source_dir": str(source_dir),
                "extra_dirs": [str(p) for p in extra_dirs],
                "output_dir": str(output_dir),
                "support_dir": str(report_dir.parent),
                "count_total": len(report_rows),
                "count_classificados": sum(1 for row in report_rows if row["status"] == "classificado"),
                "count_conflitos": sum(1 for row in report_rows if row["status"] == "conflito_titulo"),
                "count_duplicados_resolvidos": sum(1 for row in report_rows if row["status"] == "duplicado_resolvido_antigo"),
                "count_nao_classificados": sum(1 for row in report_rows if row["status"] == "nao_classificado"),
                "count_faltantes": len(missing_rows),
                "count_alertas_titulo": sum(1 for row in report_rows if row.get("alerta_titulo")),
                "rows": report_rows,
                "faltantes": missing_rows,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    count_classificados = sum(1 for row in report_rows if row['status'] == 'classificado')
    count_mais_recentes = sum(1 for row in report_rows if row['status'] == 'classificado_mais_recente')
    count_resolvidos = sum(1 for row in report_rows if row['status'] == 'duplicado_resolvido_antigo')
    count_conflitos = sum(1 for row in report_rows if row['status'] == 'conflito_titulo')
    count_nao_classificados = sum(1 for row in report_rows if row['status'] == 'nao_classificado')
    count_alertas_titulo = sum(1 for row in report_rows if row.get('alerta_titulo'))
    print(f"Classificados     : {count_classificados}")
    print(f"Mais recentes     : {count_mais_recentes}")
    print(f"Resolvidos        : {count_resolvidos}")
    print(f"Conflitos         : {count_conflitos}")
    print(f"Nao classificados : {count_nao_classificados}")
    print(f"Alertas titulo    : {count_alertas_titulo}")
    print(f"Faltantes         : {len(missing_rows)}")
    print(f"Relatorio CSV     : {csv_path}")
    print(f"Relatorio JSON    : {json_path}")
    print(f"Faltantes CSV     : {faltantes_path}")
    print(f"Alertas CSV       : {alertas_path}")
    if archived_output:
        print(f"Saida anterior    : {archived_output}")
    if count_alertas_titulo:
        print()
        print("\a" * 3, end="")
        print("[ERRO] ALERTA DE TITULO CANONICO")
        print("[ERRO] A preparacao foi INTERROMPIDA para evitar contaminacao de titulo, slug, serie e publicacao.")
        print(f"[ERRO] Foram encontrados {count_alertas_titulo} alerta(s) residual(is) de titulo.")
        print(f"[ERRO] Revise o arquivo: {alertas_path}")
        print("[ERRO] Corrija o(s) caso(s) e rode novamente o step 2.")
        return 2
    return 0


def run_reposicao(args: argparse.Namespace) -> int:
    arquivo = Path(args.arquivo).resolve()
    outline = Path(args.esboco).resolve()
    output_dir = Path(args.saida).resolve()
    substituidos_dir = output_dir / "_substituidos_reposicao"
    item, method, score, titulo_final, dst = resolve_single_article_target(
        arquivo=arquivo,
        outline=outline,
        output_dir=output_dir,
        cutoff=args.cutoff,
    )

    if not args.dry_run:
        if dst.exists():
            substituidos_dir.mkdir(parents=True, exist_ok=True)
            backup = substituidos_dir / dst.name
            shutil.move(str(dst), str(backup))
        copy_or_move(arquivo, dst, args.mover)

    print(f"Serie             : {item.serie_id}")
    print(f"Ordem             : {item.ordem:02d}")
    print(f"Titulo canonico   : {titulo_final}")
    print(f"Metodo            : {method}")
    print(f"Score             : {score:.3f}")
    print(f"Destino           : {dst}")
    print(f"Modo              : {'dry-run' if args.dry_run else 'aplicado'}")
    return 0


def run_organizar_artigo(args: argparse.Namespace) -> int:
    arquivo = Path(args.arquivo).resolve()
    outline = Path(args.esboco).resolve()
    output_dir = Path(args.saida).resolve()
    item, method, score, titulo_final, dst = resolve_single_article_target(
        arquivo=arquivo,
        outline=outline,
        output_dir=output_dir,
        cutoff=args.cutoff,
    )

    if not args.dry_run:
        copy_or_move(arquivo, dst, args.mover)

    print(f"Serie             : {item.serie_id}")
    print(f"Ordem             : {item.ordem:02d}")
    print(f"Titulo canonico   : {titulo_final}")
    print(f"Metodo            : {method}")
    print(f"Score             : {score:.3f}")
    print(f"Destino           : {dst}")
    print(f"Modo              : {'dry-run' if args.dry_run else 'aplicado'}")
    return 0


def run_reprocessar_artigo(args: argparse.Namespace) -> int:
    arquivo = Path(args.arquivo).resolve()
    outline = Path(args.esboco).resolve()
    output_dir = Path(args.saida).resolve()
    item, method, score, titulo_final, dst = resolve_single_article_target(
        arquivo=arquivo,
        outline=outline,
        output_dir=output_dir,
        cutoff=args.cutoff,
    )

    if args.dry_run:
        print(f"Serie             : {item.serie_id}")
        print(f"Ordem             : {item.ordem:02d}")
        print(f"Titulo canonico   : {titulo_final}")
        print(f"Metodo            : {method}")
        print(f"Score             : {score:.3f}")
        print(f"Destino           : {dst}")
        print("Pipeline          : dry-run")
        return 0

    if args.mover or arquivo.resolve() != dst.resolve():
        dst.parent.mkdir(parents=True, exist_ok=True)
        if dst.exists() and dst.resolve() != arquivo.resolve():
            dst.unlink()
        if args.mover:
            shutil.move(str(arquivo), str(dst))
        elif arquivo.resolve() != dst.resolve():
            shutil.copy2(str(arquivo), str(dst))

    scripts_dir = scripts_dir_from_here()
    publish_kinds = args.publish_kinds

    run_python_step(
        scripts_dir / "gerar_prompts_imagens_operacional.py",
        ["--series-root", str(output_dir), "--docx-path", str(dst)],
    )
    if args.run_images:
        run_python_step(
            scripts_dir / "gerar_imagens_lote_operacional.py",
            ["--docx-path", str(dst), "--run", "--overwrite"],
        )
    run_python_step(
        scripts_dir / "gerar_pdfs_artigos_operacional.py",
        ["--series-root", str(output_dir), "--docx-path", str(dst), "--overwrite"],
    )
    run_python_step(
        scripts_dir / "publicar_artigos_operacional.py",
        [
            "--series-root",
            str(output_dir),
            "--docx-path",
            str(dst),
            "--publish-kinds",
            publish_kinds,
            "--overwrite-media",
        ],
    )

    print(f"Serie             : {item.serie_id}")
    print(f"Ordem             : {item.ordem:02d}")
    print(f"Titulo canonico   : {titulo_final}")
    print(f"Metodo            : {method}")
    print(f"Score             : {score:.3f}")
    print(f"Destino           : {dst}")
    print(f"Pipeline          : prompts + {'imagens + ' if args.run_images else ''}pdf + publicacao")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Preparacao do ambiente operacional para lotes e reposicoes de artigos."
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    lote = sub.add_parser("organizar-lote", help="Distribui um lote de artigos em series a partir do esboco.")
    lote.add_argument("--artigos-dir", required=True, help="Diretorio com DOCX do lote principal.")
    lote.add_argument("--complementar-dir", action="append", default=[], help="Diretorio adicional com artigos complementares.")
    lote.add_argument("--esboco", required=True, help="DOCX do esboco geral.")
    lote.add_argument("--saida", required=True, help="Diretorio final do ambiente operacional.")
    lote.add_argument("--cutoff", type=float, default=0.84, help="Similaridade minima para encaixe aproximado.")
    lote.add_argument("--mover", action="store_true", help="Move os arquivos em vez de copiar.")
    lote.set_defaults(func=run_organizar_lote)

    artigo = sub.add_parser("organizar-artigo", help="Classifica e normaliza um unico artigo isolado.")
    artigo.add_argument("--arquivo", required=True, help="DOCX do artigo corrigido.")
    artigo.add_argument("--esboco", required=True, help="DOCX do esboco geral.")
    artigo.add_argument("--saida", required=True, help="Diretorio final do ambiente operacional.")
    artigo.add_argument("--cutoff", type=float, default=0.84, help="Similaridade minima para encaixe aproximado.")
    artigo.add_argument("--mover", action="store_true", help="Move o artigo em vez de copiar.")
    artigo.add_argument("--dry-run", action="store_true", help="So calcula o destino final sem gravar arquivo.")
    artigo.set_defaults(func=run_organizar_artigo)

    reprocessar = sub.add_parser("reprocessar-artigo", help="Reclassifica 1 artigo e reprojeta os artefatos/publicacao.")
    reprocessar.add_argument("--arquivo", required=True, help="DOCX do artigo corrigido.")
    reprocessar.add_argument("--esboco", required=True, help="DOCX do esboco geral.")
    reprocessar.add_argument("--saida", required=True, help="Diretorio final do ambiente operacional.")
    reprocessar.add_argument("--cutoff", type=float, default=0.84, help="Similaridade minima para encaixe aproximado.")
    reprocessar.add_argument("--mover", action="store_true", help="Move o artigo para o destino canonico.")
    reprocessar.add_argument("--run-images", action="store_true", help="Gera tambem a imagem real via API.")
    reprocessar.add_argument("--publish-kinds", default="all", help="all | docx | pdf | img | docx,pdf | docx,img | pdf,img")
    reprocessar.add_argument("--dry-run", action="store_true", help="So calcula o destino final sem executar pipeline.")
    reprocessar.set_defaults(func=run_reprocessar_artigo)

    reposicao = sub.add_parser("reposicao", help="Encaixa um artigo revisado isolado no ambiente operacional.")
    reposicao.add_argument("--arquivo", required=True, help="DOCX revisado.")
    reposicao.add_argument("--esboco", required=True, help="DOCX do esboco geral.")
    reposicao.add_argument("--saida", required=True, help="Diretorio final do ambiente operacional.")
    reposicao.add_argument("--cutoff", type=float, default=0.84, help="Similaridade minima para encaixe aproximado.")
    reposicao.add_argument("--mover", action="store_true", help="Move o arquivo de reposicao em vez de copiar.")
    reposicao.add_argument("--dry-run", action="store_true", help="So calcula o destino final sem gravar arquivo.")
    reposicao.set_defaults(func=run_reposicao)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())


