from __future__ import annotations

import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document


SERIE_DISPLAY_BY_ID = {
    1: "O GRANDE CONFLITO ENTRE CRISTO E SATANÁS",
    2: "APOCALIPSE 17 E O PRINCÍPIO DO SIMBOLISMO",
    3: "A LEI TORAH NÓMOS MANDAMENTOS ORDENANÇAS E A GRAÇA",
    4: "A BÍBLIA E A HISTÓRIA DA HUMANIDADE",
}


@dataclass
class ArticleEntry:
    serie_dir: Path
    serie_nome: str
    ordem: int
    titulo: str
    docx_path: Path


def repo_root_from_here() -> Path:
    cur = Path(__file__).resolve().parent
    for _ in range(8):
        if (cur / "manage.py").exists() or (cur / ".git").exists():
            return cur
        if cur.parent == cur:
            break
        cur = cur.parent
    raise RuntimeError("Raiz do projeto nao encontrada.")


def strip_accents(text: str) -> str:
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", text or "") if not unicodedata.combining(ch)
    )


def ascii_slug(text: str) -> str:
    text = strip_accents(text or "").lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text or "documento"


def normalize_text(text: str) -> str:
    text = strip_accents(text or "")
    text = text.upper()
    text = re.sub(r"\s+", " ", text).strip()
    return text


def strip_editorial_prefixes(text: str) -> str:
    value = str(text or "").strip()
    value = re.sub(r"^\d{8}_\d{6}__", "", value)
    value = re.sub(r"^\d{2}__", "", value)
    value = re.sub(r"^\d+\s*(?:[-_. ]+\s*)?", "", value)
    value = re.sub(r"^[^A-Za-z0-9]+", "", value)
    while True:
        updated = re.sub(r"^(?:[A-Za-z]{1,5}\d{2,6})(?:\s*[-_.:)\]]*\s*)", "", value)
        if updated == value:
            break
        value = updated.lstrip()
        value = re.sub(r"^[^A-Za-z0-9]+", "", value)
    return value.strip()


def clean_article_title(name_or_title: str) -> str:
    text = Path(name_or_title).stem
    text = strip_editorial_prefixes(text)
    text = text.replace("_", " ").strip()
    text = re.sub(r"\s+", " ", text).strip()
    return text


def docx_internal_title(docx_path: Path) -> str:
    try:
        doc = Document(str(docx_path))
    except Exception:
        return ""
    for para in doc.paragraphs:
        text = " ".join((para.text or "").split()).strip()
        if text:
            return clean_article_title(text)
    try:
        title = (doc.core_properties.title or "").strip()
        if title:
            return clean_article_title(title)
    except Exception:
        pass
    return ""


def serie_nome_from_dirname(dirname: str) -> str:
    m = re.match(r"^Serie_(\d+)__(.+)$", dirname)
    if m:
        serie_id = int(m.group(1))
        return SERIE_DISPLAY_BY_ID.get(serie_id, m.group(2).replace("-", " ").strip()).strip()
    return dirname.replace("-", " ").strip()


def article_entries(series_root: Path) -> list[ArticleEntry]:
    entries: list[ArticleEntry] = []
    for serie_dir in sorted([p for p in series_root.glob("Serie_*") if p.is_dir()], key=lambda p: p.name):
        serie_nome = serie_nome_from_dirname(serie_dir.name)
        for docx_path in sorted(serie_dir.glob("*.docx"), key=lambda p: p.name):
            m = re.match(r"^(?P<ordem>\d{2})__", docx_path.name)
            ordem = int(m.group("ordem")) if m else 0
            titulo = docx_internal_title(docx_path).strip() or clean_article_title(docx_path.name)
            entries.append(
                ArticleEntry(
                    serie_dir=serie_dir,
                    serie_nome=serie_nome,
                    ordem=ordem,
                    titulo=titulo,
                    docx_path=docx_path,
                )
            )
    return entries


def summary_from_docx(docx_path: Path, n_par: int = 5, max_chars: int = 420) -> str:
    doc = Document(str(docx_path))
    paras: list[str] = []
    for p in doc.paragraphs:
        text = " ".join((p.text or "").split()).strip()
        if not text:
            continue
        paras.append(text)
        if len(paras) >= n_par:
            break
    text = " ".join(paras).strip()
    if len(text) > max_chars:
        text = text[:max_chars].rstrip()
    return text


def safe_filename(name: str) -> str:
    text = name.strip()
    text = re.sub(r'[<>:"/\\|?*]', "", text)
    text = re.sub(r"\s+", " ", text).strip().rstrip(". ")
    return text


def safe_slug_filename(name: str, suffix: str) -> str:
    return f"{ascii_slug(name)}{suffix}"


def find_soffice() -> str:
    env = os.getenv("SOFFICE_PATH", "").strip()
    if env and Path(env).exists():
        return env
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for cand in candidates:
        if Path(cand).exists():
            return cand
    raise RuntimeError("LibreOffice nao encontrado. Defina SOFFICE_PATH.")
