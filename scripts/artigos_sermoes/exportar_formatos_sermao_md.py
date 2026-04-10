import argparse
import html
import re
import unicodedata
from pathlib import Path
from typing import Dict, List, Tuple
from artigos_operacional_utils import strip_editorial_prefixes

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SECTION_RE = re.compile(r"^(\d+)\)\s+(.*)$")

BRAND_NAVY = RGBColor(0x1D, 0x33, 0x4A)
BRAND_TEAL = RGBColor(0x2F, 0x4F, 0x4F)
BRAND_GOLD = RGBColor(0xC9, 0xA2, 0x27)
BODY_INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x5B, 0x6B, 0x7A)
SITE_URL = "www.albinomarks.com.br"


def set_cell_shading(cell, fill: str = "F8F2E4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def parse_sermao_md(text: str) -> Dict:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    sections: List[Tuple[str, List[str]]] = []
    current_title = None
    current_lines: List[str] = []

    for raw in lines:
        line = raw.rstrip()
        match = SECTION_RE.match(line.strip())
        if match:
            if current_title is not None:
                sections.append((current_title, current_lines))
            current_title = match.group(2).strip()
            current_lines = []
        else:
            if current_title is None:
                continue
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, current_lines))

    data = {"sections": sections}

    def cleaned(src: List[str]) -> List[str]:
        out = []
        for line in src:
            value = line.strip()
            if value:
                out.append(value)
        return out

    def norm(label: str) -> str:
        return ascii_slug(label).replace("-", " ")

    sec_map = {norm(title): cleaned(lines) for title, lines in sections}
    data["title"] = sec_map.get("titulo do sermao", [""])[0] if sec_map.get("titulo do sermao") else "Sermão"
    data["verso_chave"] = sec_map.get("verso chave", [""])[0] if sec_map.get("verso chave") else ""
    data["oracao"] = sec_map.get("sugestao breve de oracao", [""])[0] if sec_map.get("sugestao breve de oracao") else ""
    data["leitura"] = sec_map.get("leitura do texto biblico central", [""])[0] if sec_map.get("leitura do texto biblico central") else ""
    return data


def ascii_slug(text: str) -> str:
    text = unicodedata.normalize("NFKD", text or "")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-").lower()
    return text or "sermao"


def clean_workspace_stem(value: str) -> str:
    text = Path(value or "").stem
    text = re.sub(r"__dossie$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"__sermao$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"__relatorio_tecnico__.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"__sermao__.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\d{1,3}__+", "", text)
    text = re.sub(r"^\d{1,3}_+", "", text)
    text = strip_editorial_prefixes(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" _-") or "sermao"


def derive_output_stem(md_path: Path, title: str) -> str:
    raw = md_path.stem
    match = re.match(r"^(?P<prefix>\d{1,3})__(?P<body>.+)$", raw)
    prefix = f"{int(match.group('prefix')):02d}__" if match else ""
    body = match.group("body") if match else raw
    base = ascii_slug(clean_workspace_stem(body))
    return f"{prefix}{base}__sermao"


def md_lines_to_html(lines: List[str]) -> str:
    chunks = []
    in_list = False
    for line in lines:
        value = line.strip()
        if not value:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            continue
        if value.startswith("- "):
            if not in_list:
                chunks.append('<ul class="bullet-list">')
                in_list = True
            chunks.append(f"<li>{html.escape(value[2:].strip())}</li>")
        else:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            chunks.append(f"<p>{html.escape(value)}</p>")
    if in_list:
        chunks.append("</ul>")
    return "\n".join(chunks)


def build_html(data: Dict, page_mode: str, base_name: str) -> str:
    title = html.escape(data["title"])
    verse = html.escape(data["verso_chave"])
    prayer = html.escape(data["oracao"])
    leitura = html.escape(data["leitura"])
    current_mode_label = "Leitura online: Tablet" if page_mode == "tablet" else "Leitura online: A4"
    pdf_a4_target = f"{base_name}__a4.pdf"

    if page_mode == "tablet":
        mode_class = "tablet"
        page_css = """
        @page { size: 160mm 240mm; margin: 0; }
        body { background:#eef2f3; }
        .page-shell { max-width: 880px; margin: 0 auto; padding: 20px 16px 36px; }
        .reading-surface { padding: 34px 34px 42px; }
        h1 { font-size: 2.2rem; }
        .reading-body p, .reading-body li { font-size: 1.16rem; line-height: 1.8; }
        """
    elif page_mode == "a5":
        mode_class = "a5"
        page_css = """
        @page { size: A5; margin: 12mm 12mm 14mm 12mm; }
        body { background:#ffffff; }
        .page-shell { max-width: 100%; margin: 0 auto; padding: 0; }
        .reading-surface { padding: 16px 18px 22px; }
        h1 { font-size: 18pt; }
        .reading-body p, .reading-body li { font-size: 10.7pt; line-height: 1.52; }
        """
    else:
        mode_class = "a4"
        page_css = """
        @page { size: A4; margin: 18mm 16mm 18mm 16mm; }
        body { background:#ffffff; }
        .page-shell { max-width: 100%; margin: 0 auto; padding: 0; }
        .reading-surface { padding: 22px 26px 26px; }
        h1 { font-size: 23pt; }
        .reading-body p, .reading-body li { font-size: 11.2pt; line-height: 1.58; }
        """

    sections_html = []
    for title_sec, lines in data["sections"]:
        lower = ascii_slug(title_sec).replace("-", " ")
        if lower in {
            "titulo do sermao",
            "verso chave",
            "sugestao breve de oracao",
            "leitura do texto biblico central",
        }:
            continue
        sections_html.append(
            f'<section class="section-block"><h2>{html.escape(title_sec)}</h2>{md_lines_to_html(lines)}</section>'
        )

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{title}</title>
<style>
:root {{
  --navy:#1d334a;
  --teal:#2f4f4f;
  --gold:#c9a227;
  --paper:#fffdf7;
  --edge:#d7dde4;
  --muted:#607080;
  --ink:#1f2937;
}}
{page_css}
body {{
  margin:0;
  color:var(--ink);
  font-family: Georgia, "Times New Roman", serif;
}}
.page-shell {{
  min-height:100vh;
}}
.masthead {{
  background:linear-gradient(180deg, #143f74 0%, #123a69 100%);
  color:#f7fbff;
  border:1px solid rgba(201,162,39,.55);
  border-radius:20px;
  padding:18px 22px 16px;
  margin-bottom:16px;
  box-shadow:0 14px 28px rgba(18,58,105,.16);
}}
.masthead-top {{
  display:flex;
  justify-content:space-between;
  gap:12px;
  align-items:flex-start;
}}
.brand {{
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.76rem;
  text-transform:uppercase;
  letter-spacing:.18em;
  color:#d7e7f7;
}}
.masthead h1 {{
  margin:10px 0 8px;
  line-height:1.12;
  color:#ffffff;
}}
.masthead-sub {{
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.92rem;
  color:#e6eff7;
  max-width:58ch;
}}
.action-row {{
  display:flex;
  gap:10px;
  flex-wrap:wrap;
  margin-top:14px;
}}
.action-pill {{
  display:inline-flex;
  align-items:center;
  justify-content:center;
  border-radius:999px;
  text-decoration:none;
  padding:9px 14px;
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.84rem;
  line-height:1;
  white-space:nowrap;
  font-weight:700;
}}
.action-pill.primary {{
  background:#d8b45a;
  color:#24364a;
  border:1px solid #c9a227;
}}
.action-pill.secondary {{
  background:rgba(255,255,255,.12);
  color:#f8fbff;
  border:1px solid rgba(255,255,255,.24);
}}
.action-pill.ghost {{
  background:rgba(255,255,255,.08);
  color:#dbe8f5;
  border:1px solid rgba(255,255,255,.18);
}}
.reading-surface {{
  background:var(--paper);
  border:1px solid var(--edge);
  border-radius:22px;
  box-shadow:0 10px 30px rgba(20,30,40,.08);
}}
.key-panel {{
  margin-top:14px;
  background:#d9e9fb;
  border:1px solid #c9a227;
  border-radius:18px;
  padding:14px 16px;
  color:#1d334a;
}}
.key-panel p {{
  margin:.25rem 0;
  font-family:"Segoe UI", Arial, sans-serif;
  line-height:1.55;
}}
.reading-body h2 {{
  font-size:1.08rem;
  color:var(--navy);
  margin:28px 0 10px;
  padding-top:10px;
  border-top:1px solid #e7ebef;
  font-family:"Segoe UI", Arial, sans-serif;
}}
.reading-body h3,
.reading-body h4,
.reading-body h5,
.reading-body h6 {{
  color:var(--navy);
  font-size:1.02rem;
  line-height:1.35;
  margin:1rem 0 .55rem;
  font-family:"Segoe UI", Arial, sans-serif;
}}
.reading-body p {{
  margin:.55rem 0 .95rem;
}}
.reading-body p strong,
.reading-body li strong {{
  font-size:1em;
  color:var(--navy);
}}
.bullet-list {{
  padding-left:1.2rem;
  margin:.35rem 0 .95rem;
}}
.bullet-list li {{
  margin:.25rem 0;
}}
.footer-note {{
  margin-top:16px;
  text-align:center;
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.8rem;
  color:var(--muted);
}}
.site-url {{
  display:block;
  margin-top:6px;
  font-weight:600;
  color:var(--navy);
}}
.mode-chip {{
  display:inline-flex;
  align-items:center;
  justify-content:center;
  border-radius:999px;
  border:1px solid rgba(255,255,255,.25);
  background:#f4e8c8;
  color:#5b4b2d;
  padding:6px 10px;
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.78rem;
  white-space:nowrap;
  font-weight:700;
}}
.artifact-chip {{
  display:inline-flex;
  align-items:center;
  justify-content:center;
  border-radius:999px;
  background:#f8fbff;
  color:var(--navy);
  border:1px solid rgba(255,255,255,.28);
  padding:6px 10px;
  font-family:"Segoe UI", Arial, sans-serif;
  font-size:.78rem;
  white-space:nowrap;
  font-weight:700;
}}
.tablet .masthead h1 {{ font-size:2.18rem; }}
.a4 .masthead h1 {{ font-size:22pt; }}
.a5 .masthead h1 {{ font-size:18pt; }}
@media (max-width: 860px) {{
  .masthead-top {{ flex-direction:column; align-items:flex-start; }}
}}
</style>
</head>
<body class="{mode_class}">
  <main class="page-shell">
    <header class="masthead">
      <div class="masthead-top">
        <div>
          <div style="display:flex; gap:8px; flex-wrap:wrap; align-items:center;">
            <div class="artifact-chip">Sermão</div>
            <div class="mode-chip">{page_mode.upper()}</div>
          </div>
          <h1>{title}</h1>
          <div class="masthead-sub">Sermão diagramado para leitura, estudo e distribuição editorial, em continuidade visual com o acervo do projeto.</div>
        </div>
        <div class="brand">Sermão: Pr. Albino Marks</div>
      </div>
      <div class="action-row">
        <span class="action-pill ghost">{current_mode_label}</span>
        <a class="action-pill primary" href="{pdf_a4_target}">PDF A4</a>
      </div>
    </header>
    <section class="key-panel">
      <p><strong>Texto-chave:</strong> {verse}</p>
      <p><strong>Texto bíblico central:</strong> {leitura}</p>
    </section>
    <article class="reading-surface">
      <div class="reading-body">
        {''.join(sections_html)}
      </div>
    </article>
    <div class="footer-note">
      Projeto Pr. Albino Marks · Material editorial para uso espiritual, educacional e teológico.
      <span class="site-url">{SITE_URL}</span>
    </div>
  </main>
</body>
</html>
"""


def ensure_style(styles, name: str, base: str = "Normal", size: float = 11):
    if name not in [s.name for s in styles]:
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style = styles[name]
    style.base_style = styles[base]
    style.font.name = "Georgia"
    style.font.size = Pt(size)
    return style


def add_header_footer(section, title: str):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("Projeto Pr. Albino Marks")
    hr.bold = True
    hr.font.name = "Segoe UI"
    hr.font.size = Pt(9)
    hr.font.color.rgb = BRAND_TEAL

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"{title} · Material editorial de uso espiritual, educacional e teológico.")
    fr.font.name = "Segoe UI"
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = MUTED
    fp.add_run("\n")
    sr = fp.add_run(SITE_URL)
    sr.font.name = "Segoe UI"
    sr.font.size = Pt(8.5)
    sr.bold = True
    sr.font.color.rgb = BRAND_TEAL


def add_docx_paragraphs_from_lines(doc: Document, lines: List[str], body_style: str = "Body Text"):
    for line in lines:
        value = line.strip()
        if not value:
            continue
        if value.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(value[2:].strip())
            set_paragraph_spacing(paragraph, after=2, line=1.18)
        else:
            paragraph = doc.add_paragraph(style=body_style)
            paragraph.add_run(value)
            set_paragraph_spacing(paragraph, after=4, line=1.22)


def build_docx_a4(data: Dict, outpath: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    add_header_footer(sec, data["title"])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    body = ensure_style(styles, "Body Text", size=11)
    meta_style = ensure_style(styles, "MetaBox", size=10.5)
    quote_style = ensure_style(styles, "QuoteBlock", size=10.6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = BRAND_NAVY
    set_paragraph_spacing(title, after=6, line=1.08)

    deck = doc.add_paragraph()
    deck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = deck.add_run("Sermão diagramado para leitura, impressão e distribuição editorial.")
    dr.italic = True
    dr.font.size = Pt(10.5)
    dr.font.color.rgb = MUTED
    set_paragraph_spacing(deck, after=10, line=1.0)

    for label, value in [
        ("Verso-chave", data["verso_chave"]),
        ("Sugestão breve de oração", data["oracao"]),
        ("Texto bíblico central", data["leitura"]),
    ]:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F8F2E4")
        paragraph = cell.paragraphs[0]
        paragraph.style = meta_style
        run_label = paragraph.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.color.rgb = BRAND_TEAL
        run_value = paragraph.add_run(value)
        run_value.font.color.rgb = BODY_INK
        set_paragraph_spacing(paragraph, after=0, line=1.15)
        doc.add_paragraph()

    for title_sec, lines in data["sections"]:
        lower = ascii_slug(title_sec).replace("-", " ")
        if lower in {
            "titulo do sermao",
            "verso chave",
            "sugestao breve de oracao",
            "leitura do texto biblico central",
        }:
            continue
        heading = doc.add_paragraph()
        run_heading = heading.add_run(title_sec)
        run_heading.bold = True
        run_heading.font.name = "Segoe UI"
        run_heading.font.size = Pt(13)
        run_heading.font.color.rgb = BRAND_NAVY
        set_paragraph_spacing(heading, before=8, after=4, line=1.1)
        add_docx_paragraphs_from_lines(doc, lines, body_style=body.name)

    closing = doc.add_paragraph(style=quote_style.name)
    closing_run = closing.add_run("Projeto Pr. Albino Marks · Material editorial para uso espiritual, educacional e teológico.")
    closing_run.italic = True
    closing_run.font.color.rgb = MUTED
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(closing, before=14, after=0, line=1.0)

    doc.save(outpath)


def export_formats(md_path: Path, outdir: Path):
    text = md_path.read_text(encoding="utf-8")
    data = parse_sermao_md(text)
    base = derive_output_stem(md_path, data["title"])
    outdir.mkdir(parents=True, exist_ok=True)

    html_tablet = outdir / f"{base}__tablet.html"
    html_a4 = outdir / f"{base}__a4.html"
    html_a5 = outdir / f"{base}__a5.html"
    docx_a4 = outdir / f"{base}__a4.docx"

    html_tablet.write_text(build_html(data, "tablet", base), encoding="utf-8")
    html_a4.write_text(build_html(data, "a4", base), encoding="utf-8")
    html_a5.write_text(build_html(data, "a5", base), encoding="utf-8")
    build_docx_a4(data, docx_a4)

    return html_tablet, html_a4, html_a5, docx_a4


def main():
    ap = argparse.ArgumentParser(description="Exporta um sermão em .md para HTML tablet, HTML A4/A5 e DOCX A4.")
    ap.add_argument("--md", required=True, help="Arquivo .md do sermão")
    ap.add_argument("--outdir", required=True, help="Pasta de saída")
    args = ap.parse_args()

    md_path = Path(args.md)
    outdir = Path(args.outdir)

    if not md_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {md_path}")

    outputs = export_formats(md_path, outdir)
    for path in outputs:
        print(f"[OK] {path}")


if __name__ == "__main__":
    main()

