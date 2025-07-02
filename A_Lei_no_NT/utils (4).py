
import os
import re
import uuid
from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(titulo__startswith=titulo_base).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)

def docx_para_html(arquivo_docx):
    doc = Document(arquivo_docx)
    html = ""
    titulo = ""
    titulo_definido = False
    lista_aberta = False
    pilha_listas = []

    def detectar_estilo_manual(paragraph):
        bold = any(run.bold for run in paragraph.runs)
        underline = any(run.underline for run in paragraph.runs)
        font_size = any(run.font.size and run.font.size.pt >= 14 for run in paragraph.runs)
        return bold or underline or font_size

    for par in doc.paragraphs:
        texto = par.text.strip()
        if not texto:
            html += "<br>
"
            continue

        # Detectar título
        if not titulo_definido:
            if par.style.name.lower().startswith("heading") or detectar_estilo_manual(par):
                titulo = texto
                titulo_definido = True
                html += f"<h2>{titulo}</h2>
"
                continue

        # Detectar listas
        marcador = re.match(r"^(\s*)([-*•–]|\d+[.)]|[a-zA-Z][.)])\s+", texto)
        if marcador:
            indent = len(marcador.group(1)) // 2
            tipo = marcador.group(2)
            tag = "ul" if re.match(r"^[-*•–]$", tipo) else "ol"

            while len(pilha_listas) > indent:
                html += f"</li></{pilha_listas.pop()}>
"

            if len(pilha_listas) < indent:
                html += f"<{tag}><li>"
                pilha_listas.append(tag)
            elif pilha_listas and pilha_listas[-1] == tag:
                html += "</li><li>"
            else:
                html += f"</li></{pilha_listas.pop()}><{tag}><li>"
                pilha_listas.append(tag)

            html += texto[marcador.end():].strip()
            continue

        while pilha_listas:
            html += f"</li></{pilha_listas.pop()}>
"

        # Formatação inline
        html_formatado = ""
        for run in par.runs:
            run_text = run.text.replace('<', '&lt;').replace('>', '&gt;')
            if run.bold:
                run_text = f"<b>{run_text}</b>"
            if run.italic:
                run_text = f"<i>{run_text}</i>"
            if run.underline:
                run_text = f"<u>{run_text}</u>"
            html_formatado += run_text
        html += f"<p>{html_formatado}</p>
"

    while pilha_listas:
        html += f"</li></{pilha_listas.pop()}>
"

    # Tabelas
    for table in doc.tables:
        html += "<table border='1'>
"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text.strip()}</td>"
            html += "</tr>
"
        html += "</table>
"

    return html, titulo or "Sem Título"
