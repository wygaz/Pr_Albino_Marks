import os, re, shutil
import unicodedata
from io import BytesIO
from unidecode import unidecode
from bs4 import BeautifulSoup
from django.utils.text import slugify
from uuid import uuid4
from docx import Document
from io import BytesIO
from datetime import datetime
from django.utils.html import strip_tags
from re import split
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.db import transaction
from django.apps import apps
from pathlib import Path
from django.conf import settings

# utils.py (trecho necessário para o auditor)
# Padrões de numeração que queremos remover do fim do título
_PADROES_NUM = [
    r"\s*-\s*\d+\s+de\s+\d+\s*$",          # " - 1 de 3"
    r"\s*\(\s*\d+\s*/\s*\d+\s*\)\s*$",     # "(1/3)"
    r"\s*\(\s*\d+\s+de\s+\d+\s*\)\s*$",    # "(1 de 3)"
    r"\s*n[ºo]\.?\s*\d+\s*$",              # "nº 1"
    r"\s*-\s*parte\s*\d+\s*$",             # "- parte 1"
]

def path_docx_por_slug(slug: str) -> Path:
    return Path(settings.MEDIA_ROOT) / "uploads" / f"{slug}.docx"

def localizar_docx(slug: str) -> Path | None:
    """
    Procura DOCX no padrão novo e nos legados:
      uploads/<slug>.docx
      uploads/artigo_<slug>*.docx
      (se slug == '<base>-k-de-N') tenta também: uploads/artigo_<base>-k-de-*.docx
      uploads/artigo_temp_*.docx (fallback)
    """
    uploads = Path(settings.MEDIA_ROOT) / "uploads"
    uploads.mkdir(parents=True, exist_ok=True)

    alvo = uploads / f"{slug}.docx"
    if alvo.exists():
        return alvo

    legados = sorted(uploads.glob(f"artigo_{slug}*.docx"))
    if legados:
        return max(legados, key=lambda p: p.stat().st_mtime)

    m = re.match(r"^(?P<base>.+)-(?P<k>\d+)-de-(?P<n>\d+)$", slug)
    if m:
        base = m.group("base")
        k = m.group("k")
        flex = sorted(uploads.glob(f"artigo_{base}-{k}-de-*.docx"))
        if flex:
            return max(flex, key=lambda p: p.stat().st_mtime)

    temps = sorted(uploads.glob("artigo_temp_*.docx"))
    if temps:
        return max(temps, key=lambda p: p.stat().st_mtime)

    return None

def normalizar_docx_para_padrao(slug: str, origem: Path) -> Path:
    """Move/renomeia para uploads/<slug>.docx (overwrite seguro)."""
    destino = path_docx_por_slug(slug)
    destino.parent.mkdir(parents=True, exist_ok=True)
    if origem != destino:
        destino.unlink(missing_ok=True)           # overwrite seguro
        shutil.move(origem.as_posix(), destino.as_posix())
    return destino

def encontrar_capa_existente(slug: str) -> Path | None:
    """Retorna a capa existente para um slug (qualquer extensão conhecida)."""
    base = Path(settings.MEDIA_ROOT) / "imagens" / "artigos"
    for ext in (".png", ".jpg", ".jpeg", ".webp"):
        p = base / f"{slug}{ext}"
        if p.exists():
            return p
    # tenta variações 'temp_' geradas em uploads
    for ext in (".png", ".jpg", ".jpeg", ".webp"):
        temp = base / f"temp_{slug}{ext}"
        if temp.exists():
            return temp
    return None

def limpar_numeracao(titulo: str) -> str:
    """
    Remove qualquer marcação de numeração no final do título e devolve o 'título-base'.
    """
    base = (titulo or "").strip()
    for pat in _PADROES_NUM:
        base = re.sub(pat, "", base, flags=re.IGNORECASE).strip()
    return base

def path_capa_por_slug(slug: str, ext=".jpg") -> Path:
    return Path(settings.MEDIA_ROOT) / "imagens" / "artigos" / f"{slug}{ext}"

def gerar_slug(titulo: str) -> str:
    """
    Gera um slug único a partir de um título já 'limpo'.
    Não mexe mais no texto do título, só acrescenta sufixo no slug se colidir.
    """
    from .models import Artigo

    titulo = (titulo or "").strip()
    if not titulo:
        titulo = "Artigo"

    slug_base = slugify(unidecode(titulo))
    if not slug_base:
        # fallback bem neutro se o título não gerar slug
        import uuid
        slug_base = f"artigo-{uuid.uuid4().hex[:6]}"

    slug = slug_base
    contador = 2
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{slug_base}-{contador}"
        contador += 1
    return slug


def remover_autor_do_conteudo(html: str, autor: str | None) -> str:
    """
    Remove o autor do topo do HTML.
    Mesmo que o autor tenha virado <li> (por causa do conversor de listas),
    removemos apenas no cabeçalho (primeiros blocos) para não apagar citações no corpo.
    """
    if not html or not autor:
        return html

    autor_n = _norm_cmp(autor)
    soup = BeautifulSoup(html, "html.parser")

    # olha só o topo (primeiros 6 blocos)
    top = []
    for el in list(soup.contents):
        if getattr(el, "name", None) in ("p", "ol", "ul", "h1", "h2", "h3"):
            top.append(el)
        if len(top) >= 6:
            break

    def parece_autor(txt: str) -> bool:
        t = (txt or "").strip()
        if not t:
            return False
        # limpa marcadores comuns
        t = re.sub(r"^(autor\s*:\s*|por\s+)", "", t, flags=re.I).strip()
        t = re.sub(r"^(\d+|[ivxlc]+|[a-z])[\.\)\-–:]\s+", "", t, flags=re.I).strip()
        return _norm_cmp(t) == autor_n

    for el in top:
        if el.name == "p":
            if parece_autor(el.get_text(" ", strip=True)):
                el.decompose()
                break

        if el.name in ("ol", "ul"):
            lis = el.find_all("li", recursive=True)
            if len(lis) == 1 and parece_autor(lis[0].get_text(" ", strip=True)):
                el.decompose()
                break

    return str(soup)


        
@transaction.atomic
def gerar_titulo_numerado(titulo_base: str, ordem_por: str = "id") -> str:
    """
    [LEGADO] Mantida só por compatibilidade.
    Não gera mais 'Título - 1 de 3' ou 'Título (1 de 3)'.

    Agora apenas remove qualquer numeração no final e devolve o título-base.
    """
    base = limpar_numeracao(titulo_base or "")
    if not base:
        base = "Artigo"
    return base


def detectar_titulo_possivel(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.style.name.lower().startswith('heading'):
            return texto
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.alignment in [1, 2]:
            return texto
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12:
            return texto
    return None

def detectar_autor(paragrafos):
    for i, p in enumerate(paragrafos[1:4], start=1):
        texto = ''.join(run.text for run in p.runs).strip()
        if not texto or len(texto) > 100:
            continue
        if i in (1, 2) and len(texto.split()) <= 5 and len(p.runs) <= 3:
            return texto
        if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return texto
    return None

def formatar_paragrafo(paragrafo):
    partes = []

    for run in paragrafo.runs:
        texto = run.text.replace('\n', '')  # removendo quebras de linha
        if not texto.strip():
            continue

        estilo = ''
        if run.bold:
            estilo += 'font-weight:bold;'
        if run.italic:
            estilo += 'font-style:italic;'
        if run.underline:
            estilo += 'text-decoration:underline;'

        if estilo:
            partes.append(f'<span style="{estilo}">{texto}</span>')
        else:
            partes.append(texto)

    return ''.join(partes)

def renomear_com_slug(caminho_arquivo, slug):
    ext = os.path.splitext(caminho_arquivo)[1]
    return f"{slug}{ext}"


def renomear_arquivo_word(arquivo):
    nome_base = os.path.basename(arquivo.name)
    novo_nome = gerar_slug(nome_base.replace('.docx', '')) + '.docx'
    return novo_nome

def renomear_imagem_capa(nome_original):
    nome_base, ext = os.path.splitext(nome_original)
    nome_slug = gerar_slug(nome_base)
    return nome_slug + ext

def converter_subtitulos_manualmente_numerados(html):
    padrao = r'<p>\s*(?P<numero>\d+)[\.\-–)]\s+(?P<texto>.*?)</p>'

    def substituir(match):
        numero = match.group("numero")
        texto = match.group("texto").strip()
        resultado = f'<h2>{numero}. <strong>{texto}</strong></h2>'
        print(f'🔎 número: {numero}')
        print(f'🔎 texto: {texto}')
        print(f'✅ resultado: {resultado}\n')
        return resultado

    return re.sub(padrao, substituir, html)

def converter_para_html(paragrafos):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    html = ""
    lista_aberta = False
    tipo_lista = None
    contador_ord = 0
    contador_nao_ord = 0

    def fecha_lista():
        nonlocal html, lista_aberta, tipo_lista
        if lista_aberta:
            html += f'</{tipo_lista}>\n'
            lista_aberta = False
            tipo_lista = None

    def detectar_item_manual(texto):
        import re
        if re.match(r"^\s*[\-\u2022\*]\s+", texto):  # - • *
            return "ul"
        if re.match(r"^\s*(\d+|[a-zA-Z]+)[\.\)\-–]\s+", texto):  # 1. a) I) A. etc
            return "ol"
        return None

    def eh_titulo_exemplo(texto):
        import re
        return bool(re.match(r"^\d+\.\s+Lista", texto.strip(), re.IGNORECASE))

    for p in paragrafos:
        texto_corrigido = ''.join(run.text for run in p.runs).strip()
        alinhado_central = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        estilo = p.style.name.lower()

        if not texto_corrigido:
            fecha_lista()
            html += "<br>\n"
            continue

        # 🛑 Ignora título de exemplo de lista (ex: "1. Lista Não Ordenada")
        if eh_titulo_exemplo(texto_corrigido):
            fecha_lista()
            html += f"<p><strong>{texto_corrigido}</strong></p>\n"
            continue

        tipo_item = detectar_item_manual(texto_corrigido)

        # ▶️ Item de lista
        if tipo_item:
            if not lista_aberta or tipo_lista != tipo_item:
                fecha_lista()
                html += f"<{tipo_item}>\n"
                lista_aberta = True
                tipo_lista = tipo_item
            item = texto_corrigido.split(' ', 1)[1] if ' ' in texto_corrigido else texto_corrigido
            html += f"<li>{item}</li>\n"
        else:
            fecha_lista()
            # 🔤 Tratamento de formatação inline
            partes = []
            for run in p.runs:
                t = run.text.replace('\n', '<br>').strip()
                if not t:
                    continue
                if run.bold:
                    t = f"<strong>{t}</strong>"
                if run.italic:
                    t = f"<em>{t}</em>"
                if run.underline:
                    t = f"<u>{t}</u>"
                partes.append(t)
            linha_formatada = ' '.join(partes)
            if alinhado_central:
                html += f"<p style='text-align: center'>{linha_formatada}</p>\n"
            else:
                html += f"<p>{linha_formatada}</p>\n"

    fecha_lista()
    return html

def aplicar_estrutura_listas(html):
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, 'html.parser')

    novas_tags = []
    lista_atual = []
    tipo_lista = None

    print("\n🔍 INICIANDO AGRUPAMENTO DE LISTAS...")

    for i, el in enumerate(soup.contents):
        print(f"\n📄 Elemento {i}: {el.name} — {str(el)[:60]}...")

        if el.name in ['ul', 'ol']:
            itens = el.find_all('li')
            print(f"  ➕ Detectado: lista <{el.name}> com {len(itens)} itens")

            if not itens:
                print("  ⚠️ Lista ignorada (sem itens <li>)")
                continue

            if tipo_lista == el.name:
                lista_atual.extend(itens)
                print(f"  🔁 Continuando lista do tipo <{el.name}> com +{len(itens)} itens")
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                    print(f"  ✅ Lista <{tipo_lista}> finalizada e adicionada com {len(lista_atual)} itens")

                tipo_lista = el.name
                lista_atual = itens
                print(f"  🔄 Nova lista <{el.name}> iniciada")
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                print(f"  ✅ Lista <{tipo_lista}> finalizada e adicionada com {len(lista_atual)} itens")

                lista_atual = []
                tipo_lista = None

            novas_tags.append(el)
            print(f"  📌 Elemento não-lista adicionado normalmente: <{el.name}>")

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)
        print(f"  ✅ Lista <{tipo_lista}> finalizada no final com {len(lista_atual)} itens")

    print("🏁 AGRUPAMENTO DE LISTAS FINALIZADO.\n")
    return str(BeautifulSoup(''.join(str(tag) for tag in novas_tags), 'html.parser'))


SM_PREFIX_RE = re.compile(r"^\s*sm\s*\d+[a-z]?\s+", re.IGNORECASE)

def remover_sm_prefix(texto: str) -> str:
    return SM_PREFIX_RE.sub("", (texto or "").strip()).strip()

def _norm_cmp(texto: str) -> str:
    t = remover_sm_prefix(texto or "")
    t = re.sub(r"\s+", " ", t).strip().lower()
    return t


def docx_para_html(arquivo):
    doc = Document(arquivo)
    paragrafos = doc.paragraphs

    # -------- 1) detectar título (sem depender de negrito do 1º run)
    titulo = None
    idx_titulo = None

    for i, p in enumerate(paragrafos[:8]):
        txt = ''.join(run.text for run in p.runs).strip()
        if not txt:
            continue

        estilo = (p.style.name or "").lower() if p.style else ""
        alinhado_centro = (p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER)
        tem_negrito = any((run.bold and run.text.strip()) for run in p.runs)

        # Regra forte: começa com Sm#### -> é título (quase sempre)
        if SM_PREFIX_RE.match(txt):
            titulo = remover_sm_prefix(txt)
            idx_titulo = i
            break

        # Headings
        if estilo.startswith("heading"):
            titulo = remover_sm_prefix(txt)
            idx_titulo = i
            break

        # Centro + curto
        if alinhado_centro and len(txt.split()) <= 18:
            titulo = remover_sm_prefix(txt)
            idx_titulo = i
            break

        # Negrito em algum run + curto (nos primeiros)
        if i <= 2 and tem_negrito and len(txt.split()) <= 18:
            titulo = remover_sm_prefix(txt)
            idx_titulo = i
            break

    # remove o parágrafo do título (se achou)
    work = [p for j, p in enumerate(paragrafos) if j != idx_titulo]

    # -------- 2) detectar autor (e REMOVER ANTES de converter para evitar virar <li>)
    autor = None
    idx_autor = None

    for j, p in enumerate(work[:8]):
        txt = ''.join(run.text for run in p.runs).strip()
        if not txt:
            continue

        # não confundir com o próprio título repetido
        if titulo and _norm_cmp(txt) == _norm_cmp(titulo):
            continue

        # marcadores explícitos
        if re.match(r"^\s*(autor\s*:|por\s+)", txt, flags=re.I):
            autor = re.sub(r"^\s*(autor\s*:\s*|por\s+)", "", txt, flags=re.I).strip()
            idx_autor = j
            break

        # alinhado à direita e curto
        if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT and len(txt.split()) <= 6:
            autor = txt.strip()
            idx_autor = j
            break

        # heurística simples no topo
        if j <= 2 and len(txt.split()) <= 5 and not re.search(r"[.!?]\s*$", txt):
            autor = txt.strip()
            idx_autor = j
            break

    if idx_autor is not None:
        work = [p for k, p in enumerate(work) if k != idx_autor]

    # -------- 3) converter parágrafos para HTML
    html = converter_para_html(work)

    # -------- 4) adicionar tabelas (simples, sem “tabelas aninhadas”)
    for tabela in doc.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(pp.text for pp in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    # -------- 5) limpeza final: remove autor se ainda aparecer no topo
    html = remover_autor_do_conteudo(html, autor)

    # -------- 6) remove título duplicado no topo (caso apareça por repetição no DOCX)
    if titulo:
        soup = BeautifulSoup(html, "html.parser")
        for _ in range(6):
            first = None
            for el in list(soup.contents):
                if getattr(el, "name", None) in ("p", "h1", "h2", "h3"):
                    first = el
                    break
            if not first:
                break
            t = first.get_text(" ", strip=True)
            if not t:
                first.decompose()
                continue
            if _norm_cmp(t) == _norm_cmp(titulo):
                first.decompose()
                continue
            # se ainda vier “Sm#### Título”
            if SM_PREFIX_RE.match(t) and _norm_cmp(remover_sm_prefix(t)) == _norm_cmp(titulo):
                first.decompose()
                continue
            break
        html = str(soup)

    # -------- 7) sem prints
    return html, titulo, autor



# =====================================================
# Rotinas adicionais de normalização de mídia (DOCX + Capas)
# =====================================================
import datetime

PREFER_CAPA_EXTS = (".png", ".jpg", ".jpeg", ".webp")

def _quarentena_dir(kind: str) -> Path:
    stamp = datetime.datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    base = Path(settings.MEDIA_ROOT) / "_legacy" / stamp / kind
    base.mkdir(parents=True, exist_ok=True)
    return base

def localizar_docx_legados(slug: str) -> list[Path]:
    up = Path(settings.MEDIA_ROOT) / "uploads"
    up.mkdir(parents=True, exist_ok=True)
    out = list(up.glob(f"artigo_{slug}*.docx"))
    m = re.match(r"^(?P<base>.+)-(?P<k>\d+)-de-(?P<n>\d+)$", slug)
    if m:
        base, k = m.group("base"), m.group("k")
        out += list(up.glob(f"artigo_{base}-{k}-de-*.docx"))
    out += list(up.glob("artigo_temp_*.docx"))
    return [p for p in {p.resolve() for p in out} if p.exists()]

def normalizar_docx_com_limpeza(slug: str, *, dry_run: bool=False) -> dict:
    alvo = path_docx_por_slug(slug)
    uploads = alvo.parent
    uploads.mkdir(parents=True, exist_ok=True)
    legados = localizar_docx_legados(slug)
    movidos, mantido = [], alvo if alvo.exists() else None
    quar = _quarentena_dir("docx")
    candidatos = [p for p in legados if p != alvo]
    if not mantido and legados:
        mantido = max(legados, key=lambda p: p.stat().st_mtime)
    if mantido and mantido != alvo:
        if not dry_run:
            alvo.unlink(missing_ok=True)
            shutil.move(mantido.as_posix(), alvo.as_posix())
        movidos.append(mantido)
        mantido = alvo
    extras = [p for p in legados if p != mantido]
    for p in extras:
        if not dry_run and p.exists():
            shutil.move(p.as_posix(), (quar / p.name).as_posix())
        movidos.append(p)
    return {"mantido": mantido if mantido and mantido.exists() else None,
            "movidos": movidos, "quarentena": quar}

def localizar_capas_relacionadas(slug: str) -> list[Path]:
    base = Path(settings.MEDIA_ROOT) / "imagens" / "artigos"
    base.mkdir(parents=True, exist_ok=True)
    out = []
    for ext in PREFER_CAPA_EXTS:
        out += list(base.glob(f"{slug}{ext}"))
        out += list(base.glob(f"temp_{slug}{ext}"))
        out += list(base.glob(f"{slug}_*.{ext.lstrip('.')}"))
    return [p for p in {p.resolve() for p in out} if p.exists()]

def _escolher_melhor_capa(candidatos: list[Path]) -> Path | None:
    if not candidatos: return None
    for ext in PREFER_CAPA_EXTS:
        exts = [p for p in candidatos if p.suffix.lower() == ext]
        if exts:
            return max(exts, key=lambda p: p.stat().st_mtime)
    return max(candidatos, key=lambda p: p.stat().st_mtime)

def normalizar_capas_com_limpeza(slug: str, *, dry_run: bool=False) -> dict:
    candidatos = localizar_capas_relacionadas(slug)
    movidos, mantida = [], None
    quar = _quarentena_dir("capas")
    if not candidatos:
        return {"mantida": None, "movidos": [], "quarentena": quar}
    melhor = _escolher_melhor_capa(candidatos)
    dest = path_capa_por_slug(slug, PREFER_CAPA_EXTS[0])
    if melhor and dest:
        if not dry_run:
            dest.parent.mkdir(parents=True, exist_ok=True)
            if melhor != dest:
                dest.unlink(missing_ok=True)
                shutil.move(melhor.as_posix(), dest.as_posix())
        movidos.append(melhor)
        mantida = dest
    for p in candidatos:
        if mantida and p.resolve() == mantida.resolve():
            continue
        if not dry_run and p.exists():
            shutil.move(p.as_posix(), (quar / p.name).as_posix())
        movidos.append(p)
    return {"mantida": mantida if mantida and mantida.exists() else None,
            "movidos": movidos, "quarentena": quar}
