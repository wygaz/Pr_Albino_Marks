from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

SPACE = '\u00A0'  # Espaço incondicional
MAX_RUN_LENGTH = 100

def aplicar_formatacao(run):
    texto = run.text.replace('\xa0', SPACE).replace(' ', SPACE)
    if not texto.strip():
        return texto
    if run.bold:
        texto = f"<strong>{texto}</strong>"
    if run.italic:
        texto = f"<em>{texto}</em>"
    if run.underline:
        texto = f"<u>{texto}</u>"
    return texto

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if p.style.name in ['Título', 'Heading 1', 'Título 1'] or (len(texto) < 80 and p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER):
            return texto
    return ""

def remover_autor_por_heuristica(paragrafo, nomes_autores_conhecidos):
    texto = ''.join(run.text.strip() for run in paragrafo.runs)
    if texto.startswith("Autor:"):
        return True
    if paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        if len(texto) <= MAX_RUN_LENGTH:
            for nome in nomes_autores_conhecidos:
                if nome.lower() in texto.lower():
                    return True
    return False

def docx_para_html(caminho_arquivo, nomes_autores_conhecidos=None):
    if nomes_autores_conhecidos is None:
        nomes_autores_conhecidos = []

    doc = Document(caminho_arquivo)
    html = ""
    titulo = detectar_titulo_possivel(doc.paragraphs)

    for p in doc.paragraphs:
        texto = ''.join(aplicar_formatacao(run) for run in p.runs).strip()

        # Ignora título duplicado
        if texto.replace(SPACE, ' ') == titulo.replace(SPACE, ' '):
            continue

        # Heurística para pular menção ao autor
        if remover_autor_por_heuristica(p, nomes_autores_conhecidos):
            continue

        # Detecta lista não ordenada manual
        if texto.startswith('- '):
            texto = texto[2:]
            html += f"<ul><li>{texto}</li></ul>"
        else:
            html += f"<p>{texto}</p>"

    soup = BeautifulSoup(html, 'html.parser')
    return str(soup), titulo
