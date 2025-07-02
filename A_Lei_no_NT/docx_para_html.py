
def docx_para_html(docx_file):
    from docx import Document
    from io import BytesIO
    from .models import Artigo
    from django.utils.text import slugify

    doc = Document(BytesIO(docx_file.read()))
    html = ''
    titulo = None

    in_ol = False
    in_ul = False
    list_buffer = ''

    def flush_list():
        nonlocal html, list_buffer, in_ol, in_ul
        if in_ol:
            html += f"<ol>{list_buffer}</ol>"
        elif in_ul:
            html += f"<ul>{list_buffer}</ul>"
        list_buffer = ''
        in_ol = False
        in_ul = False

    for para in doc.paragraphs:
        estilo = para.style.name.lower()
        texto = ''
        for run in para.runs:
            trecho = run.text.replace('\n', '<br>')
            if run.bold:
                trecho = f"<strong>{trecho}</strong>"
            if run.italic:
                trecho = f"<em>{trecho}</em>"
            if run.underline:
                trecho = f"<u>{trecho}</u>"
            texto += trecho

        # Título do artigo (captura apenas o primeiro título)
        if not titulo and texto.strip():
            titulo = texto.strip()

        if 'heading' in estilo:
            flush_list()
            nivel = ''.join(filter(str.isdigit, estilo)) or '1'
            html += f"<h{nivel}>{texto}</h{nivel}>"

        elif 'list number' in estilo:
            if not in_ol:
                flush_list()
                in_ol = True
            list_buffer += f"<li>{texto}</li>"

        elif 'list bullet' in estilo:
            if not in_ul:
                flush_list()
                in_ul = True
            list_buffer += f"<li>{texto}</li>"

        else:
            flush_list()
            if texto.strip():
                html += f"<p>{texto}</p>"

    flush_list()  # Fechar qualquer lista pendente

    # Tabelas
    for table in doc.tables:
        html += '<table border="1">'
        for row in table.rows:
            html += '<tr>'
            for cell in row.cells:
                cell_text = '<br>'.join(para.text for para in cell.paragraphs)
                html += f"<td>{cell_text}</td>"
            html += '</tr>'
        html += '</table>'

    return html, titulo
