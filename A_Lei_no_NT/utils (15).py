
import os
import uuid
import re
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document
from html import escape

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos:
        texto = p.text.strip()
        if texto and len(texto.split()) >= 3 and texto == texto.upper():
            return texto.title()
        if texto.startswith(("Título:", "Titulo:")):
            return texto.split(":", 1)[1].strip()
        if p.style.name.lower().startswith('heading') and len(texto) > 5:
            return texto.strip()
    return None

def docx_para_html(arquivo):
    doc = Document(arquivo)
    html = []
    pilha_listas = []
    titulo_extraido = detectar_titulo_possivel(doc.paragraphs)

    for p in doc.paragraphs:
        texto = escape(p.text.strip())
        estilo = p.style.name.lower()

        # Quebra de página
        if "page break" in p._element.xml.lower():
            html.append('<hr class="quebra-pagina"/>')
            continue

        # Lista detectada por estilo
        if p._element.xpath('.//w:numPr'):
            nivel = int(p._element.xpath('.//w:ilvl')[0].text)
            tipo_lista = "ol" if "numbered" in estilo else "ul"

            # Ajustar níveis
            while len(pilha_listas) < nivel + 1:
                pilha_listas.append(tipo_lista)
                html.append(f"<{tipo_lista}>")

            while len(pilha_listas) > nivel + 1:
                tipo_fechamento = pilha_listas.pop()
                html.append(f"</{tipo_fechamento}>")

            html.append(f"<li>{texto}</li>")
            continue

        # Lista manual com hífen, asterisco ou número
        if re.match(r"^[-*•]\s+", p.text.strip()) or re.match(r"^\d+[.)]\s+", p.text.strip()):
            tipo_manual = "ul" if re.match(r"^[-*•]\s+", p.text.strip()) else "ol"
            if not pilha_listas:
                pilha_listas.append(tipo_manual)
                html.append(f"<{tipo_manual}>")
            html.append(f"<li>{texto[2:].strip()}</li>")
            continue
        else:
            # Fecha listas abertas
            while pilha_listas:
                html.append(f"</{pilha_listas.pop()}>")

        # Cabeçalhos
        if "heading" in estilo:
            nivel = re.sub(r'\D', '', estilo) or "2"
            html.append(f"<h{nivel}>{texto}</h{nivel}>")
        elif texto:
            # Estilo direto (negrito, itálico etc.)
            negrito = any(r.bold for r in p.runs if r.bold)
            italico = any(r.italic for r in p.runs if r.italic)
            sublinhado = any(r.underline for r in p.runs if r.underline)

            estilo = ""
            if negrito: estilo += "<strong>"
            if italico: estilo += "<em>"
            if sublinhado: estilo += "<u>"

            fechamento = ""
            if sublinhado: fechamento = "</u>" + fechamento
            if italico: fechamento = "</em>" + fechamento
            if negrito: fechamento = "</strong>" + fechamento

            html.append(f"<p>{estilo}{texto}{fechamento}</p>")

    # Fecha listas restantes
    while pilha_listas:
        html.append(f"</{pilha_listas.pop()}>")

    # Tabelas
    for table in doc.tables:
        html.append("<table class='tabela-docx'>")
        for row in table.rows:
            html.append("<tr>")
            for cell in row.cells:
                html.append(f"<td>{escape(cell.text.strip())}</td>")
            html.append("</tr>")
        html.append("</table>")

    # Se não encontrou título, gera a partir do nome do arquivo
    if not titulo_extraido:
        nome_arquivo = os.path.splitext(os.path.basename(arquivo.name))[0]
        titulo_extraido = nome_arquivo.replace("_", " ").replace("-", " ").title()

    return "\n".join(html), titulo_extraido

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r" - \d+ de \d+$")
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub("", artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join("capas", novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join("uploads", novo_nome)
