
import os
import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document

# === Funções auxiliares ===

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(titulo__startswith=titulo_base).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)

# === Função principal ===

def docx_para_html(caminho_docx):
    doc = Document(caminho_docx)
    html = ''
    titulo = None

    def detectar_titulo_manual(paragraph):
        if not paragraph.text.strip():
            return False
        if len(paragraph.text.split()) <= 3:
            return True
        if paragraph.style.name.lower().startswith("heading"):
            return True
        if paragraph.alignment and paragraph.alignment == 1:  # centralizado
            return True
        try:
            for run in paragraph.runs:
                if run.bold or run.font.size and run.font.size.pt > 14:
                    return True
        except:
            pass
        return False

    def escapar(texto):
        return texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    lista_aberta = False
    tipo_lista = None

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            html += '<br>'
            continue
        if not titulo and detectar_titulo_manual(p):
            titulo = txt
            html += f'<h2>{escapar(txt)}</h2>'
            continue

        if p.style.name.startswith("List Bullet"):
            if not lista_aberta:
                html += '<ul>'
                lista_aberta = True
                tipo_lista = 'ul'
            html += f'<li>{escapar(txt)}</li>'
            continue
        elif p.style.name.startswith("List Number"):
            if not lista_aberta:
                html += '<ol>'
                lista_aberta = True
                tipo_lista = 'ol'
            html += f'<li>{escapar(txt)}</li>'
            continue
        else:
            if lista_aberta:
                html += f'</{tipo_lista}>'
                lista_aberta = False
                tipo_lista = None
            negrito = any(run.bold for run in p.runs)
            italico = any(run.italic for run in p.runs)
            sublinhado = any(run.underline for run in p.runs)
            texto_formatado = escapar(txt)
            if negrito:
                texto_formatado = f"<strong>{texto_formatado}</strong>"
            if italico:
                texto_formatado = f"<em>{texto_formatado}</em>"
            if sublinhado:
                texto_formatado = f"<u>{texto_formatado}</u>"
            html += f'<p>{texto_formatado}</p>'

    if lista_aberta:
        html += f'</{tipo_lista}>'

    # Tabela (básico)
    for tabela in doc.tables:
        html += '<table border="1">'
        for row in tabela.rows:
            html += '<tr>'
            for cell in row.cells:
                html += f'<td>{escapar(cell.text.strip())}</td>'
            html += '</tr>'
        html += '</table>'

    if not titulo:
        titulo = "Sem Título"

    soup = BeautifulSoup(html, 'html.parser')
    return str(soup), titulo
