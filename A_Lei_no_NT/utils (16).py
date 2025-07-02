import docx
from django.utils.text import slugify
from datetime import datetime
import re

def gerar_slug(titulo):
    data = datetime.now().strftime("%Y%m%d%H%M%S")
    return slugify(f"{titulo}-{data}")

def detectar_titulo_possivel(doc):
    for p in doc.paragraphs[:10]:
        if p.text.strip() and len(p.text.split()) <= 12:
            if p.style.name.startswith('Heading') or p.alignment == 1 or any(run.bold for run in p.runs):
                return p.text.strip()
    return None

def formatar_run(run):
    texto = run.text.replace("\n", "<br>")  # Tratamento de quebra de linha manual
    if not texto.strip():
        return ""
    if run.bold:
        texto = f"<strong>{texto}</strong>"
    if run.italic:
        texto = f"<em>{texto}</em>"
    if run.underline:
        texto = f"<u>{texto}</u>"
    return texto

def docx_para_html(arquivo):
    doc = docx.Document(arquivo)
    html = ""
    titulo_detectado = detectar_titulo_possivel(doc)

    if titulo_detectado:
        html += f"<h2>{titulo_detectado}</h2>\n"

    in_list = False
    list_type = None
    last_ilvl = 0

    for p in doc.paragraphs:
        estilo = p.style.name.lower()
        is_list = p._element.xpath('.//w:numPr')

        if is_list:
            ilvl_elements = p._element.xpath('.//w:ilvl')
            nivel = int(ilvl_elements[0].text) if ilvl_elements else 0
            tipo_lista = "ol" if "numbered" in estilo else "ul"

            if not in_list or tipo_lista != list_type or nivel != last_ilvl:
                if in_list:
                    html += f"</{list_type}>\n"
                html += f"<{tipo_lista}>\n"
                in_list = True
                list_type = tipo_lista
                last_ilvl = nivel

            conteudo = "".join([formatar_run(run) for run in p.runs])
            html += f"<li>{conteudo}</li>\n"

        else:
            if in_list:
                html += f"</{list_type}>\n"
                in_list = False
            conteudo = "".join([formatar_run(run) for run in p.runs])
            if not conteudo.strip():
                continue
            elif estilo.startswith('heading'):
                html += f"<h3>{conteudo}</h3>\n"
            else:
                html += f"<p>{conteudo}</p>\n"

    if in_list:
        html += f"</{list_type}>\n"

    return html, titulo_detectado or ""
