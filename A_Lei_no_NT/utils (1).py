
from docx import Document
from docx.table import _Cell
from django.utils.text import slugify
import os
import re

def limpar_texto(texto):
    return texto.replace('\n', ' ').replace('\t', ' ').strip()

def formatar_run(run):
    texto = run.text
    if not texto:
        return ""
    if run.bold:
        texto = f"<strong>{texto}</strong>"
    if run.italic:
        texto = f"<em>{texto}</em>"
    if run.underline:
        texto = f"<u>{texto}</u>"
    return texto

def extrair_texto_com_formatacao(paragrafo):
    return ''.join(formatar_run(run) for run in paragrafo.runs)

def detectar_lista_manual(texto):
    return re.match(r"^([\-*•·]|\d+[.)]|[a-zA-Z][.)])\s+", texto.strip())

def limpar_marcador_manual(texto):
    return re.sub(r"^([\-*•·]|\d+[.)]|[a-zA-Z][.)])\s+", "", texto.strip())

def docx_para_html(arquivo):
    doc = Document(arquivo)
    html = ""
    titulo = None
    lista_stack = []
    ultimo_nivel = 0

    for elem in doc.element.body.iter():
        if elem.tag.endswith('}p'):
            p = docx_paragraph_from_element(doc, elem)
            if not p:
                continue
            estilo = p.style.name.lower()
            texto_formatado = extrair_texto_com_formatacao(p).strip()
            if not texto_formatado:
                continue

            if not titulo and estilo.startswith('heading'):
                titulo = limpar_texto(p.text)
                continue

            if estilo.startswith('heading'):
                html += f"<h2>{texto_formatado}</h2>"
            elif 'pageBreakBefore' in p._element.xml:
                html += '<div class="page-break"></div>'
            elif detectar_lista_manual(p.text):
                nivel = p.paragraph_format.left_indent or 0
                nivel = int(nivel.pt // 18) if nivel else 0
                texto_item = limpar_marcador_manual(texto_formatado)

                while len(lista_stack) > nivel:
                    html += f"</li></ul>"
                    lista_stack.pop()

                if len(lista_stack) < nivel + 1:
                    html += "<ul>"
                    lista_stack.append("ul")

                html += f"<li>{texto_item}"
            else:
                while lista_stack:
                    html += f"</li></ul>"
                    lista_stack.pop()
                html += f"<p>{texto_formatado}</p>"

        elif elem.tag.endswith('}tbl'):
            tabela = docx_table_from_element(doc, elem)
            if not tabela:
                continue
            html += "<table border='1'>"
            for linha in tabela.rows:
                html += "<tr>"
                for celula in linha.cells:
                    html += f"<td>{extrair_texto_com_formatacao(celula.paragraphs[0])}</td>"
                html += "</tr>"
            html += "</table>"

    while lista_stack:
        html += f"</li></ul>"
        lista_stack.pop()

    return html.strip(), titulo or "Sem Título"

def gerar_slug(titulo):
    return slugify(titulo)

def gerar_nome_arquivo(titulo):
    return slugify(titulo) + ".jpg"

# Métodos auxiliares para evitar quebra no iter
def docx_paragraph_from_element(doc, elem):
    for p in doc.paragraphs:
        if p._element == elem:
            return p
    return None

def docx_table_from_element(doc, elem):
    for t in doc.tables:
        if t._element == elem:
            return t
    return None
