
import os
import re
import uuid
from django.utils.text import slugify
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos:
        if not p.text.strip():
            continue
        estilo = p.style.name.lower()
        if 'título' in estilo or 'heading' in estilo:
            return p.text.strip()
        run = p.runs[0] if p.runs else None
        if run:
            font = run.font
            if font.size and font.size.pt > 14:
                return p.text.strip()
            if font.bold or font.underline:
                return p.text.strip()
        if len(p.text.strip()) <= 30:
            return p.text.strip()
    return None

def docx_para_html(docx_file):
    document = Document(docx_file)
    html = ''
    lista_pilha = []
    titulo = detectar_titulo_possivel(document.paragraphs) or "Sem Título"

    for par in document.paragraphs:
        texto = par.text.strip()
        if not texto:
            continue

        estilo = par.style.name.lower()
        if 'heading' in estilo or 'título' in estilo:
            html += f"<h2>{texto}</h2>"
            continue

        if par.runs and par.runs[0].font.bold:
            texto = f"<strong>{texto}</strong>"
        if par.runs and par.runs[0].font.italic:
            texto = f"<em>{texto}</em>"
        if par.runs and par.runs[0].font.underline:
            texto = f"<u>{texto}</u>"

        match_ol = re.match(r"^\(?P<numero>\(?:(?:\d+)|(?:[a-zA-Z]))[\.)])\s+(?P<resto>.+)$", texto)
        match_ul = re.match(r"^[\-*•]\s+(?P<resto>.+)$", texto)

        nivel = par._element.xpath('count(ancestor::w:tbl) + count(ancestor::w:tr) + count(ancestor::w:pStyle)')
        nivel = min(int(nivel), 5)

        if match_ol:
            while len(lista_pilha) > nivel:
                html += f"</{lista_pilha.pop()}>"
            while len(lista_pilha) < nivel:
                html += "<ol>"
                lista_pilha.append('ol')
            html += f"<li>{match_ol.group('resto')}</li>"
        elif match_ul:
            while len(lista_pilha) > nivel:
                html += f"</{lista_pilha.pop()}>"
            while len(lista_pilha) < nivel:
                html += "<ul>"
                lista_pilha.append('ul')
            html += f"<li>{match_ul.group('resto')}</li>"
        else:
            while lista_pilha:
                html += f"</{lista_pilha.pop()}>"
            html += f"<p>{texto}</p>"

    while lista_pilha:
        html += f"</{lista_pilha.pop()}>"
    return html, titulo
