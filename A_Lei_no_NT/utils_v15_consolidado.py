import os
import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

def gerar_slug(titulo):
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    from .models import Artigo
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos = Artigo.objects.filter(titulo__startswith=titulo_base).order_by(ordem_por)
    total = artigos.count() + 1
    artigos_limpos = []
    for artigo in artigos:
        titulo_limpo = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_limpo
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    return os.path.join('uploads', f"{nome_base}{ext.lower()}")

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    return os.path.join('capas', f"{nome_base}-capa{ext.lower()}")

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.style.name.startswith('Heading'):
            return texto
    return ''

def detectar_autor(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if not texto:
            continue
        texto_limpo = re.sub(r'^Autor:\s*', '', texto, flags=re.IGNORECASE).strip()
        if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT or texto.startswith("Autor:"):
            return texto_limpo
    return None

def docx_para_html(arquivo):
    doc = Document(arquivo)
    paragrafos = doc.paragraphs
    html = ''
    titulo = detectar_titulo_possivel(paragrafos)
    autor = detectar_autor(paragrafos)

    if titulo:
        html += f'<h1 style="text-align: center;">{titulo}</h1>\n'
    if autor:
        html += f'<p style="text-align: right;"><strong>{autor}</strong></p>\n'

    for p in paragrafos:
        texto = ''.join(run.text for run in p.runs).strip()
        if not texto or texto == f"Autor: {autor}" or texto == "Autor:":
            continue

        estilo = p.style.name.lower()
        texto_formatado = ''
        for run in p.runs:
            t = run.text.replace('\n', '<br>').replace('\xa0', '&nbsp;')
            if '<w:br w:type="page"/>' in run._element.xml:
                texto_formatado += '<hr class="page-break">'
            if run.bold:
                t = f'<strong>{t}</strong>'
            if run.italic:
                t = f'<em>{t}</em>'
            if run.underline:
                t = f'<u>{t}</u>'
            texto_formatado += t

        if p._p.pPr is not None and p._p.pPr.numPr is not None:
            numId = p._p.pPr.numPr.numId.val
            if numId == 1:
                html += f'<ul><li>{texto_formatado}</li></ul>\n'
            elif numId == 2:
                html += f'<ol><li>{texto_formatado}</li></ol>\n'
            else:
                html += f'<p>{texto_formatado}</p>\n'
        else:
            html += f'<p>{texto_formatado}</p>\n'

    for tabela in doc.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(p.text for p in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    soup = BeautifulSoup(html, 'html.parser')
    novas_tags = []
    lista_atual = []
    tipo_lista = None

    for el in soup.contents:
        if el.name in ['ul', 'ol']:
            if tipo_lista == el.name:
                lista_atual.extend(el.find_all('li'))
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                tipo_lista = el.name
                lista_atual = el.find_all('li')
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                lista_atual = []
                tipo_lista = None
            novas_tags.append(el)

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)

    soup.clear()
    for tag in novas_tags:
        soup.append(tag)

    return str(soup), titulo