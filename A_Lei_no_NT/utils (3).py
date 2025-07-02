
import os
import re
import uuid
from docx import Document
from typing import Tuple
from django.utils.text import slugify


def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')

    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)

    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)

    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()

    return f"{titulo_base} - {total} de {total}"


def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug


def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)


def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)


def detectar_titulo_possivel(documento: Document) -> str:
    for paragrafo in documento.paragraphs:
        texto = paragrafo.text.strip()
        if not texto:
            continue
        alinhado_centro = paragrafo.alignment == 1
        propriedades_formatacao = [
            run.bold or run.underline or (run.font.size and run.font.size.pt > 14)
            for run in paragrafo.runs
        ]
        curto_ou_frase = len(texto.split()) <= 8
        if (alinhado_centro or any(propriedades_formatacao)) and curto_ou_frase:
            return texto
    return ""


def docx_para_html(docx_file) -> Tuple[str, str]:
    documento = Document(docx_file)
    titulo = ""
    html = ""

    for paragrafo in documento.paragraphs:
        texto = paragrafo.text.strip()
        if not texto:
            continue
        estilo = paragrafo.style.name.lower()
        if 'heading' in estilo and not titulo:
            titulo = texto
            html += f"<h2>{texto}</h2>\n"
        else:
            html += f"<p>{texto}</p>\n"

    if not titulo:
        titulo = detectar_titulo_possivel(documento)

    return html, titulo
