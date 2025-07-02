import os
import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document

def docx_para_html(docx_file):
    doc = Document(docx_file)
    html = ""
    titulo = None

    # Critérios de heurística para título
    for p in doc.paragraphs:
        estilo = p.style.name.lower()
        texto = p.text.strip()

        if not texto:
            continue

        if "title" in estilo or "título" in estilo or estilo.startswith("heading"):
            titulo = texto
            continue

        if not titulo and (
            ("bold" in estilo or "negrito" in estilo or "sublin" in estilo or "central" in estilo)
            or p.style.font.size and p.style.font.size.pt > 14
        ):
            titulo = texto
            continue

    if not titulo:
        raise ValueError("Título não identificado automaticamente. Por favor, informe manualmente.")

    pilha_listas = []
    html += "<div class='artigo'>\n"

    for p in doc.paragraphs:
        texto = p.text.strip()
        if not texto:
            continue

        estilo = p.style.name.lower()

        # Listas com marcadores automáticos do Word
        if estilo.startswith("list bullet"):
            nivel = estilo.count(" ")  # tentativa simples de detectar nível
            while len(pilha_listas) < nivel + 1:
                html += "<ul>\n"
                pilha_listas.append("ul")
            while len(pilha_listas) > nivel + 1:
                html += "</ul>\n"
                pilha_listas.pop()
            html += f"<li>{texto}</li>\n"
            continue

        # Listas ordenadas automáticas do Word
        if estilo.startswith("list number"):
            nivel = estilo.count(" ")
            while len(pilha_listas) < nivel + 1:
                html += "<ol>\n"
                pilha_listas.append("ol")
            while len(pilha_listas) > nivel + 1:
                html += "</ol>\n"
                pilha_listas.pop()
            html += f"<li>{texto}</li>\n"
            continue

        # Fim de listas se estilo não for de lista
        while pilha_listas:
            tipo = pilha_listas.pop()
            html += f"</{tipo}>\n"

        # Negrito, itálico e sublinhado básicos
        if "**" in texto:
            texto = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", texto)
        if "_" in texto:
            texto = re.sub(r"_(.*?)_", r"<em>\1</em>", texto)
        if "__" in texto:
            texto = re.sub(r"__(.*?)__", r"<u>\1</u>", texto)

        html += f"<p>{texto}</p>\n"

    while pilha_listas:
        html += f"</{pilha_listas.pop()}>\n"

    html += "</div>"
    return html, titulo

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(titulo__startswith=titulo_base).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)