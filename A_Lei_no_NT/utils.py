import os, re, shutil
import unicodedata
from io import BytesIO
from unidecode import unidecode
from bs4 import BeautifulSoup
from django.utils.text import slugify
from uuid import uuid4
from docx import Document
from io import BytesIO
from datetime import datetime
from django.utils.html import strip_tags
from re import split
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.db import transaction
from django.apps import apps
from pathlib import Path
from django.conf import settings

# utils.py (trecho necessário para o auditor)

# Padrões de numeração que queremos remover do fim do título
_PADROES_NUM = [
    r"\s*-\s*\d+\s+de\s+\d+\s*$",
    r"\s*\(\s*\d+\s*/\s*\d+\s*\)\s*$",
    r"\s*n[ºo]\.?\s*\d+\s*$",
    r"\s*-\s*parte\s*\d+\s*$",
]


def path_docx_por_slug(slug: str) -> Path:
    return Path(settings.MEDIA_ROOT) / "uploads" / f"{slug}.docx"

def localizar_docx(slug: str) -> Path | None:
    """
    Procura DOCX no padrão novo e nos legados:
      uploads/<slug>.docx
      uploads/artigo_<slug>*.docx
      (se slug == '<base>-k-de-N') tenta também: uploads/artigo_<base>-k-de-*.docx
      uploads/artigo_temp_*.docx (fallback)
    """
    uploads = Path(settings.MEDIA_ROOT) / "uploads"
    uploads.mkdir(parents=True, exist_ok=True)

    alvo = uploads / f"{slug}.docx"
    if alvo.exists():
        return alvo

    legados = sorted(uploads.glob(f"artigo_{slug}*.docx"))
    if legados:
        return max(legados, key=lambda p: p.stat().st_mtime)

    m = re.match(r"^(?P<base>.+)-(?P<k>\d+)-de-(?P<n>\d+)$", slug)
    if m:
        base = m.group("base")
        k = m.group("k")
        flex = sorted(uploads.glob(f"artigo_{base}-{k}-de-*.docx"))
        if flex:
            return max(flex, key=lambda p: p.stat().st_mtime)

    temps = sorted(uploads.glob("artigo_temp_*.docx"))
    if temps:
        return max(temps, key=lambda p: p.stat().st_mtime)

    return None

def normalizar_docx_para_padrao(slug: str, origem: Path) -> Path:
    """Move/renomeia para uploads/<slug>.docx (overwrite seguro)."""
    destino = path_docx_por_slug(slug)
    destino.parent.mkdir(parents=True, exist_ok=True)
    if origem != destino:
        destino.unlink(missing_ok=True)           # overwrite seguro
        shutil.move(origem.as_posix(), destino.as_posix())
    return destino

def encontrar_capa_existente(slug: str) -> Path | None:
    """Retorna a capa existente para um slug (qualquer extensão conhecida)."""
    base = Path(settings.MEDIA_ROOT) / "imagens" / "artigos"
    for ext in (".png", ".jpg", ".jpeg", ".webp"):
        p = base / f"{slug}{ext}"
        if p.exists():
            return p
    # tenta variações 'temp_' geradas em uploads
    for ext in (".png", ".jpg", ".jpeg", ".webp"):
        temp = base / f"temp_{slug}{ext}"
        if temp.exists():
            return temp
    return None

def limpar_numeracao(titulo: str) -> str:
    """
    Remove qualquer marcação de numeração no final do título e devolve o 'título-base'.
    """
    base = (titulo or "").strip()
    for pat in _PADROES_NUM:
        base = re.sub(pat, "", base, flags=re.IGNORECASE).strip()
    return base

def path_capa_por_slug(slug: str, ext=".jpg") -> Path:
    return Path(settings.MEDIA_ROOT) / "imagens" / "artigos" / f"{slug}{ext}"

def gerar_slug(titulo):
    from .models import Artigo
    
    if not titulo or not titulo.strip():
        titulo = gerar_titulo_numerado()  # fallback se título for vazio ou só espaços

    slug_base = slugify(unidecode(titulo))
    if not slug_base:
        slug_base = f'artigo-{uuid4().hex[:6]}'

    slug = slug_base
    contador = 2
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{slug_base}-{contador}"
        contador += 1
    return slug

def remover_autor_do_conteudo(html, autor):
    from bs4 import BeautifulSoup
    import re

    soup = BeautifulSoup(html, 'html.parser')
    autor_lower = autor.lower().strip()
    candidatos = soup.find_all(['p', 'li'])[:3]  # verifica só os 3 primeiros parágrafos

    for tag in candidatos:
        texto = tag.get_text(strip=True).lower()

        # Remove se contiver nome do autor + poucas palavras
        if autor_lower in texto and len(texto.split()) <= 6:
            tag.decompose()

        # Remove se começar com algo como "1. Albino Marks", "I. Albino Marks", "a) Albino Marks"
        if re.match(r'^(\d+\.|[ivxlc]+\.|[a-z]\))\s*' + re.escape(autor_lower), texto):
            tag.decompose()

    return str(soup)

        
@transaction.atomic
def gerar_titulo_numerado(titulo_base: str, ordem_por: str = "id") -> str:
    """
    Recebe um título-base (SEM numeração) e calcula o título definitivo do novo artigo,
    renumerando retroativamente os existentes do mesmo tema se necessário.

    Regras:
    - Se total do tema == 1 => título fica SEM numeração.
    - Se total >= 2         => todos ficam " - k de N" e o novo recebe " - N de N".

    Observações:
    - A 'família' do tema é identificada pelo mesmo título-base no início (case-insensitive).
    - Usa transação + select_for_update para evitar corrida entre dois saves simultâneos.
    """
    Artigo = apps.get_model("A_Lei_no_NT", "Artigo")

    base = limpar_numeracao(titulo_base)
    if not base:
        base = "Artigo"

    # Bloqueia os artigos do mesmo tema durante a renumeração
    qs = (Artigo.objects
          .select_for_update()
          .filter(titulo__istartswith=base)
          .order_by(ordem_por))

    existentes = list(qs)
    for a in existentes:
        # Normaliza em memória: remove numeração antiga para garantir consistência
        a.titulo = limpar_numeracao(a.titulo)

    total = len(existentes) + 1  # incluindo o NOVO que ainda vai ser salvo

    if total == 1:
        # Será o único do tema -> sem numeração
        return base

    # Renumera retroativamente os existentes: 1..(total-1)
    for idx, a in enumerate(existentes, start=1):
        novo_titulo = f"{base} - {idx} de {total}"
        if a.titulo != novo_titulo:
            a.titulo = novo_titulo
            a.save(update_fields=["titulo"])

    # Título do NOVO artigo: "base - total de total"
    return f"{base} - {total} de {total}"


def detectar_titulo_possivel(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.style.name.lower().startswith('heading'):
            return texto
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.alignment in [1, 2]:
            return texto
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12:
            return texto
    return None

def detectar_autor(paragrafos):
    for i, p in enumerate(paragrafos[1:4], start=1):
        texto = ''.join(run.text for run in p.runs).strip()
        if not texto or len(texto) > 100:
            continue
        if i in (1, 2) and len(texto.split()) <= 5 and len(p.runs) <= 3:
            return texto
        if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return texto
    return None

def remover_autor_do_conteudo(html, autor):
    if not autor:
        return html

    soup = BeautifulSoup(html, 'html.parser')
    for p in soup.find_all('p'):
        if autor in p.text.strip():
            p.decompose()
            break  # Remove apenas a primeira ocorrência do nome do autor
    html_final = str(soup)

    html_final = converter_subtitulos_manualmente_numerados(html_final)

    return html_final

def formatar_paragrafo(paragrafo):
    partes = []

    for run in paragrafo.runs:
        texto = run.text.replace('\n', '')  # removendo quebras de linha
        if not texto.strip():
            continue

        estilo = ''
        if run.bold:
            estilo += 'font-weight:bold;'
        if run.italic:
            estilo += 'font-style:italic;'
        if run.underline:
            estilo += 'text-decoration:underline;'

        if estilo:
            partes.append(f'<span style="{estilo}">{texto}</span>')
        else:
            partes.append(texto)

    return ''.join(partes)

def converter_para_html(paragrafos):
    html = ''
    padrao_lista = re.compile(r'^(\d+\.)|([ivxlc]+\.)|([IVXLC]+\.)|([a-zA-Z]\.)|([a-zA-Z]\))\s+')
    for p in paragrafos:
        texto_bruto = ''.join(run.text for run in p.runs).strip()
        if not texto_bruto:
            continue
        texto_formatado = formatar_paragrafo(p)
        if p.style.name.lower().startswith('heading'):
            html += f'<h2>{texto_formatado}</h2>\n'
        elif padrao_lista.match(texto_bruto):
            # Remove o prefixo manual antes de inserir como item de lista
            texto_formatado = padrao_lista.sub('', texto_formatado).strip()
            html += f'<ol><li>{texto_formatado}</li></ol>\n'
        elif texto_bruto.startswith(('-', '*', '•')):
            html += f'<ul><li>{texto_formatado[1:].strip()}</li></ul>\n'
        else:
            html += f'<p>{texto_formatado}</p>\n'
    return html

def renomear_com_slug(caminho_arquivo, slug):
    ext = os.path.splitext(caminho_arquivo)[1]
    return f"{slug}{ext}"


def renomear_arquivo_word(arquivo):
    nome_base = os.path.basename(arquivo.name)
    novo_nome = gerar_slug(nome_base.replace('.docx', '')) + '.docx'
    return novo_nome

def renomear_imagem_capa(nome_original):
    nome_base, ext = os.path.splitext(nome_original)
    nome_slug = gerar_slug(nome_base)
    return nome_slug + ext

def converter_subtitulos_manualmente_numerados(html):
    padrao = r'<p>\s*(?P<numero>\d+)[\.\-–)]\s+(?P<texto>.*?)</p>'

    def substituir(match):
        numero = match.group("numero")
        texto = match.group("texto").strip()
        resultado = f'<h2>{numero}. <strong>{texto}</strong></h2>'
        print(f'🔎 número: {numero}')
        print(f'🔎 texto: {texto}')
        print(f'✅ resultado: {resultado}\n')
        return resultado

    return re.sub(padrao, substituir, html)

def converter_para_html(paragrafos):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    html = ""
    lista_aberta = False
    tipo_lista = None
    contador_ord = 0
    contador_nao_ord = 0

    def fecha_lista():
        nonlocal html, lista_aberta, tipo_lista
        if lista_aberta:
            html += f'</{tipo_lista}>\n'
            lista_aberta = False
            tipo_lista = None

    def detectar_item_manual(texto):
        import re
        if re.match(r"^\s*[\-\u2022\*]\s+", texto):  # - • *
            return "ul"
        if re.match(r"^\s*(\d+|[a-zA-Z]+)[\.\)\-–]\s+", texto):  # 1. a) I) A. etc
            return "ol"
        return None

    def eh_titulo_exemplo(texto):
        import re
        return bool(re.match(r"^\d+\.\s+Lista", texto.strip(), re.IGNORECASE))

    for p in paragrafos:
        texto_corrigido = ''.join(run.text for run in p.runs).strip()
        alinhado_central = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        estilo = p.style.name.lower()

        if not texto_corrigido:
            fecha_lista()
            html += "<br>\n"
            continue

        # 🛑 Ignora título de exemplo de lista (ex: "1. Lista Não Ordenada")
        if eh_titulo_exemplo(texto_corrigido):
            fecha_lista()
            html += f"<p><strong>{texto_corrigido}</strong></p>\n"
            continue

        tipo_item = detectar_item_manual(texto_corrigido)

        # ▶️ Item de lista
        if tipo_item:
            if not lista_aberta or tipo_lista != tipo_item:
                fecha_lista()
                html += f"<{tipo_item}>\n"
                lista_aberta = True
                tipo_lista = tipo_item
            item = texto_corrigido.split(' ', 1)[1] if ' ' in texto_corrigido else texto_corrigido
            html += f"<li>{item}</li>\n"
        else:
            fecha_lista()
            # 🔤 Tratamento de formatação inline
            partes = []
            for run in p.runs:
                t = run.text.replace('\n', '<br>').strip()
                if not t:
                    continue
                if run.bold:
                    t = f"<strong>{t}</strong>"
                if run.italic:
                    t = f"<em>{t}</em>"
                if run.underline:
                    t = f"<u>{t}</u>"
                partes.append(t)
            linha_formatada = ' '.join(partes)
            if alinhado_central:
                html += f"<p style='text-align: center'>{linha_formatada}</p>\n"
            else:
                html += f"<p>{linha_formatada}</p>\n"

    fecha_lista()
    return html

def aplicar_estrutura_listas(html):
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, 'html.parser')

    novas_tags = []
    lista_atual = []
    tipo_lista = None

    print("\n🔍 INICIANDO AGRUPAMENTO DE LISTAS...")

    for i, el in enumerate(soup.contents):
        print(f"\n📄 Elemento {i}: {el.name} — {str(el)[:60]}...")

        if el.name in ['ul', 'ol']:
            itens = el.find_all('li')
            print(f"  ➕ Detectado: lista <{el.name}> com {len(itens)} itens")

            if not itens:
                print("  ⚠️ Lista ignorada (sem itens <li>)")
                continue

            if tipo_lista == el.name:
                lista_atual.extend(itens)
                print(f"  🔁 Continuando lista do tipo <{el.name}> com +{len(itens)} itens")
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                    print(f"  ✅ Lista <{tipo_lista}> finalizada e adicionada com {len(lista_atual)} itens")

                tipo_lista = el.name
                lista_atual = itens
                print(f"  🔄 Nova lista <{el.name}> iniciada")
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                print(f"  ✅ Lista <{tipo_lista}> finalizada e adicionada com {len(lista_atual)} itens")

                lista_atual = []
                tipo_lista = None

            novas_tags.append(el)
            print(f"  📌 Elemento não-lista adicionado normalmente: <{el.name}>")

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)
        print(f"  ✅ Lista <{tipo_lista}> finalizada no final com {len(lista_atual)} itens")

    print("🏁 AGRUPAMENTO DE LISTAS FINALIZADO.\n")
    return str(BeautifulSoup(''.join(str(tag) for tag in novas_tags), 'html.parser'))


def docx_para_html(arquivo):
    from docx import Document
    from bs4 import BeautifulSoup
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(arquivo)
    paragrafos = doc.paragraphs

    # 🔍 Detectar título (máx. 12 palavras, estilo Heading, ou centralizado em negrito)
    titulo = None
    indice_titulo = None
    for i, p in enumerate(paragrafos[:5]):
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and (p.style.name.startswith('Heading') or p.runs[0].bold or p.alignment == WD_ALIGN_PARAGRAPH.CENTER):
            titulo = texto
            indice_titulo = i
            break

    # 🔐 Cópia segura dos parágrafos, sem o título
    paragrafos_sem_titulo = [p for i, p in enumerate(paragrafos) if i != indice_titulo]

    # 🎯 Detectar autor com base no conteúdo restante
    autor = detectar_autor(paragrafos_sem_titulo)

    # 🧾 Converter parágrafos em HTML (com quebra de linha, negrito, itálico etc.)
    html = converter_para_html(paragrafos_sem_titulo)

    # 📋 Adicionar tabelas do documento
    for tabela in doc.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(p.text for p in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    # 🧼 Remover nome do autor do conteúdo renderizado
    html = remover_autor_do_conteudo(html, autor)

    # 🧠 Agrupamento de listas sequenciais (ul, ol)
    soup = BeautifulSoup(html, 'html.parser')
    novas_tags = []
    lista_atual = []
    tipo_lista = None

    for el in soup.contents:
        if el.name in ['ul', 'ol']:
            if tipo_lista == el.name:
                lista_atual.extend(el.find_all('li'))
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                tipo_lista = el.name
                lista_atual = el.find_all('li')
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                lista_atual = []
                tipo_lista = None
            novas_tags.append(el)

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)

    soup.clear()
    for tag in novas_tags:
        soup.append(tag)

    # 🔧 Converte para string e aplica estrutura inteligente de listas e subtítulos
    try:
        html_final = str(soup)
        html_final = aplicar_estrutura_listas(html_final)
    except Exception as e:
        print("⚠️ Erro ao processar HTML final:", e)
        html_final = ""

    # ✅ Retorna conteúdo final + título extraído + autor identificado
    print("Retornando:", html_final, titulo)
    return html_final, titulo, autor


# =====================================================
# Rotinas adicionais de normalização de mídia (DOCX + Capas)
# =====================================================
import datetime

PREFER_CAPA_EXTS = (".png", ".jpg", ".jpeg", ".webp")

def _quarentena_dir(kind: str) -> Path:
    stamp = datetime.datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    base = Path(settings.MEDIA_ROOT) / "_legacy" / stamp / kind
    base.mkdir(parents=True, exist_ok=True)
    return base

def localizar_docx_legados(slug: str) -> list[Path]:
    up = Path(settings.MEDIA_ROOT) / "uploads"
    up.mkdir(parents=True, exist_ok=True)
    out = list(up.glob(f"artigo_{slug}*.docx"))
    m = re.match(r"^(?P<base>.+)-(?P<k>\d+)-de-(?P<n>\d+)$", slug)
    if m:
        base, k = m.group("base"), m.group("k")
        out += list(up.glob(f"artigo_{base}-{k}-de-*.docx"))
    out += list(up.glob("artigo_temp_*.docx"))
    return [p for p in {p.resolve() for p in out} if p.exists()]

def normalizar_docx_com_limpeza(slug: str, *, dry_run: bool=False) -> dict:
    alvo = path_docx_por_slug(slug)
    uploads = alvo.parent
    uploads.mkdir(parents=True, exist_ok=True)
    legados = localizar_docx_legados(slug)
    movidos, mantido = [], alvo if alvo.exists() else None
    quar = _quarentena_dir("docx")
    candidatos = [p for p in legados if p != alvo]
    if not mantido and legados:
        mantido = max(legados, key=lambda p: p.stat().st_mtime)
    if mantido and mantido != alvo:
        if not dry_run:
            alvo.unlink(missing_ok=True)
            shutil.move(mantido.as_posix(), alvo.as_posix())
        movidos.append(mantido)
        mantido = alvo
    extras = [p for p in legados if p != mantido]
    for p in extras:
        if not dry_run and p.exists():
            shutil.move(p.as_posix(), (quar / p.name).as_posix())
        movidos.append(p)
    return {"mantido": mantido if mantido and mantido.exists() else None,
            "movidos": movidos, "quarentena": quar}

def localizar_capas_relacionadas(slug: str) -> list[Path]:
    base = Path(settings.MEDIA_ROOT) / "imagens" / "artigos"
    base.mkdir(parents=True, exist_ok=True)
    out = []
    for ext in PREFER_CAPA_EXTS:
        out += list(base.glob(f"{slug}{ext}"))
        out += list(base.glob(f"temp_{slug}{ext}"))
        out += list(base.glob(f"{slug}_*.{ext.lstrip('.')}"))
    return [p for p in {p.resolve() for p in out} if p.exists()]

def _escolher_melhor_capa(candidatos: list[Path]) -> Path | None:
    if not candidatos: return None
    for ext in PREFER_CAPA_EXTS:
        exts = [p for p in candidatos if p.suffix.lower() == ext]
        if exts:
            return max(exts, key=lambda p: p.stat().st_mtime)
    return max(candidatos, key=lambda p: p.stat().st_mtime)

def normalizar_capas_com_limpeza(slug: str, *, dry_run: bool=False) -> dict:
    candidatos = localizar_capas_relacionadas(slug)
    movidos, mantida = [], None
    quar = _quarentena_dir("capas")
    if not candidatos:
        return {"mantida": None, "movidos": [], "quarentena": quar}
    melhor = _escolher_melhor_capa(candidatos)
    dest = path_capa_por_slug(slug, PREFER_CAPA_EXTS[0])
    if melhor and dest:
        if not dry_run:
            dest.parent.mkdir(parents=True, exist_ok=True)
            if melhor != dest:
                dest.unlink(missing_ok=True)
                shutil.move(melhor.as_posix(), dest.as_posix())
        movidos.append(melhor)
        mantida = dest
    for p in candidatos:
        if mantida and p.resolve() == mantida.resolve():
            continue
        if not dry_run and p.exists():
            shutil.move(p.as_posix(), (quar / p.name).as_posix())
        movidos.append(p)
    return {"mantida": mantida if mantida and mantida.exists() else None,
            "movidos": movidos, "quarentena": quar}
