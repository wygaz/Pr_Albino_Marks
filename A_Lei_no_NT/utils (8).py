
import os
import re
import uuid
from django.utils.text import slugify
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')

    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)

    total = artigos_com_titulo_base.count() + 1

    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)

    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()

    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos:
        texto = p.text.strip()
        if not texto or len(texto) < 3:
            continue
        try:
            if p.style.name.lower().startswith("título") or p.style.name.lower() in ["heading", "heading 1"]:
                return texto
        except:
            pass
        if p.alignment == 1 or any(run.bold or run.font.size and run.font.size.pt > 14 for run in p.runs):
            return texto
    return "Título não detectado"

def aplicar_formatacao(run):
    texto = run.text
    if not texto:
        return ""

    if run.bold:
        texto = f"<strong>{texto}</strong>"
    if run.italic:
        texto = f"<em>{texto}</em>"
    if run.underline:
        texto = f"<u>{texto}</u>"

    return texto

def docx_para_html(docx_file):
    if isinstance(docx_file, bytes):
        doc = Document(BytesIO(docx_file))
    else:
        doc = Document(docx_file)

    paragrafos = doc.paragraphs
    titulo = detectar_titulo_possivel(paragrafos)

    html = ""
    for p in paragrafos:
        if not p.text.strip():
            continue

        texto_formatado = ''.join([aplicar_formatacao(run) for run in p.runs])
        texto = p.text.strip()

        estilo = p.style.name.lower() if p.style and p.style.name else ""

        if estilo.startswith("título") or estilo.startswith("heading"):
            html += f"<h2>{texto_formatado}</h2>
"
        elif estilo in ["normal", "corpo de texto"] or estilo == "":
            html += f"<p>{texto_formatado}</p>
"
        elif estilo.startswith("lista") or re.match(r"^\d+[\.\)]", texto) or re.match(r"[-*+•]", texto):
            html += f"<li>{texto_formatado}</li>
"
        else:
            html += f"<p>{texto_formatado}</p>
"

    soup = BeautifulSoup(html, 'html.parser')

    # Agrupamento em listas ordenadas e não ordenadas
    novos_elementos = []
    dentro_de_lista = False
    tipo_lista = None
    itens_lista = []

    for el in soup.contents:
        if el.name == 'li':
            simbolo = el.get_text().strip()[:1]
            if simbolo in '0123456789':
                tipo = 'ol'
            else:
                tipo = 'ul'

            if not dentro_de_lista:
                tipo_lista = tipo
                dentro_de_lista = True
                itens_lista = [el]
            elif tipo == tipo_lista:
                itens_lista.append(el)
            else:
                nova_lista = soup.new_tag(tipo_lista)
                for item in itens_lista:
                    nova_lista.append(item)
                novos_elementos.append(nova_lista)
                dentro_de_lista = True
                tipo_lista = tipo
                itens_lista = [el]
        else:
            if dentro_de_lista:
                nova_lista = soup.new_tag(tipo_lista)
                for item in itens_lista:
                    nova_lista.append(item)
                novos_elementos.append(nova_lista)
                dentro_de_lista = False
                tipo_lista = None
                itens_lista = []
            novos_elementos.append(el)

    if dentro_de_lista:
        nova_lista = soup.new_tag(tipo_lista)
        for item in itens_lista:
            nova_lista.append(item)
        novos_elementos.append(nova_lista)

    soup.clear()
    for novo in novos_elementos:
        soup.append(novo)

    return str(soup), titulo
