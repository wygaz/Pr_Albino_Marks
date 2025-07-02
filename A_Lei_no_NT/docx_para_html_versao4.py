def docx_para_html(arquivo):
    from docx import Document
    from bs4 import BeautifulSoup
    from re import split
    document = Document(arquivo)
    html = ''
    titulo_extraido = ''
    buffer_paragrafo = None
    encontrou_titulo = False
    autor_extraido = ''
    corpo_html = ''

    for paragrafo in document.paragraphs:
        texto_raw = paragrafo.text.strip()
        if not texto_raw:
            continue

        texto_para_analise = "".join(run.text for run in paragrafo.runs).strip()

        # Detectar título pela heurística (negrito + fonte > 12)
        if not encontrou_titulo:
            maior_fonte = 0
            contem_negrito = False
            for run in paragrafo.runs:
                if run.bold:
                    contem_negrito = True
                if run.font.size:
                    try:
                        pontos = run.font.size.pt
                        maior_fonte = max(maior_fonte, pontos)
                    except:
                        continue
            if maior_fonte > 12 and contem_negrito:
                titulo_extraido = texto_para_analise
                encontrou_titulo = True
                continue  # pula do HTML, pois o Django já exibe {{ artigo.titulo }}

        # Detectar autor
        if texto_para_analise.lower().startswith("autor:"):
            autor_extraido = texto_para_analise
            html += f"<p><strong>{autor_extraido}</strong></p>"
            continue

        estilo = paragrafo.style.name.lower()
        texto_formatado = ''
        for run in paragrafo.runs:
            texto = run.text.replace('\n', '<br>')
            if not texto.strip() and '<w:br' not in run._element.xml:
                continue

            if '<w:br w:type="page"/>' in run._element.xml:
                texto_formatado += '<hr class="page-break">'

            if run.bold:
                texto = f'<strong>{texto}</strong>'
            if run.italic:
                texto = f'<em>{texto}</em>'
            if run.underline:
                texto = f'<u>{texto}</u>'

            if texto_formatado and texto and not texto_formatado[-1].isspace() and not texto[0].isspace():
                texto_formatado += ' '

            texto_formatado += texto

        # Títulos internos
        if estilo.startswith('heading'):
            nivel = ''.join(filter(str.isdigit, estilo)) or '1'
            html += f'<h{nivel}>{texto_formatado}</h{nivel}>'

        # Listas
        elif paragrafo._p.pPr is not None and paragrafo._p.pPr.numPr is not None:
            numId = paragrafo._p.pPr.numPr.numId.val
            if numId == 1:
                html += f'<ul><li>{texto_formatado}</li></ul>'
            elif numId == 2:
                html += f'<ol><li>{texto_formatado}</li></ol>'
            else:
                html += f'<p>{texto_formatado}</p>'
        else:
            subparagrafos = split(r'(?<=[.?!])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕ])', texto_formatado)
            for sub in subparagrafos:
                html += f'<p>{sub.strip()}</p>'

    # Tabelas
    for tabela in document.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(p.text for p in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    # Agrupamento de listas
    soup = BeautifulSoup(html, 'html.parser')
    novas_tags = []
    lista_atual = []
    tipo_lista = None

    for el in soup.contents:
        if el.name in ['ul', 'ol']:
            if tipo_lista == el.name:
                lista_atual.extend(el.find_all('li'))
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                    lista_atual = []
                tipo_lista = el.name
                lista_atual = el.find_all('li')
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                lista_atual = []
                tipo_lista = None
            novas_tags.append(el)

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)

    soup.clear()
    for tag in novas_tags:
        soup.append(tag)

    if not titulo_extraido:
        titulo_extraido = "Artigo sem Título"

    return str(soup), titulo_extraido
