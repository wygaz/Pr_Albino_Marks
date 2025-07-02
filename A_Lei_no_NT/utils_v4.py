
import os
import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document
from io import BytesIO
from datetime import datetime
from django.utils.html import strip_tags

def docx_para_html(arquivo):
    from docx import Document
    from bs4 import BeautifulSoup
    from re import split
    document = Document(arquivo)
    html = ''
    titulo_extraido = ''
    buffer_paragrafo = None  # Para unir "1." com linha seguinte

    for paragrafo in document.paragraphs:
        texto_raw = paragrafo.text.strip()
        if not texto_raw:
            continue

        # Une títulos numerados (1., 2., 3.) ao texto seguinte
        if texto_raw.strip().rstrip('.') in ['1', '2', '3', '4']:
            buffer_paragrafo = texto_raw
            continue
        if buffer_paragrafo:
            texto_raw = f'{buffer_paragrafo} {texto_raw}'
            buffer_paragrafo = None

        estilo = paragrafo.style.name.lower()

        # NOVO: heurística para detectar título por fonte e negrito
        if not titulo_extraido:
            texto_para_titulo = "".join(run.text for run in paragrafo.runs).strip()
            maior_fonte = 0
            contem_negrito = False

            for run in paragrafo.runs:
                if run.bold:
                    contem_negrito = True
                if run.font.size:
                    try:
                        pontos = run.font.size.pt
                        maior_fonte = max(maior_fonte, pontos)
                    except:
                        continue

            if maior_fonte > 12 and contem_negrito:
                titulo_extraido = texto_para_titulo

        # 🎨 Formatação de runs (negrito, itálico, sublinhado, quebras, página)
        texto_formatado = ''
        for i, run in enumerate(paragrafo.runs):
            texto = run.text.replace('\n', '<br>')
            if not texto.strip() and '<w:br' not in run._element.xml:
                continue

            if '<w:br w:type="page"/>' in run._element.xml:
                texto_formatado += '<hr class="page-break">'

            if run.bold:
                texto = f'<strong>{texto}</strong>'
            if run.italic:
                texto = f'<em>{texto}</em>'
            if run.underline:
                texto = f'<u>{texto}</u>'

            if texto_formatado and texto and not texto_formatado[-1].isspace() and not texto[0].isspace():
                texto_formatado += ' '

            texto_formatado += texto

        # 🎯 Títulos
        if estilo.startswith('heading'):
            nivel = ''.join(filter(str.isdigit, estilo)) or '1'
            html += f'<h{nivel}>{texto_formatado}</h{nivel}>'

        # 📑 Listas
        elif paragrafo._p.pPr is not None and paragrafo._p.pPr.numPr is not None:
            numId = paragrafo._p.pPr.numPr.numId.val
            if numId == 1:
                html += f'<ul><li>{texto_formatado}</li></ul>'
            elif numId == 2:
                html += f'<ol><li>{texto_formatado}</li></ol>'
            else:
                html += f'<p>{texto_formatado}</p>'
        else:
            # Quebrar em múltiplos <p> com base em pontuação
            subparagrafos = split(r'(?<=[.?!])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕ])', texto_formatado)
            for sub in subparagrafos:
                html += f'<p>{sub.strip()}</p>'

    # 📊 Tabelas (após parágrafos, por enquanto)
    for tabela in document.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(p.text for p in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    # 🔧 Agrupamento correto das listas
    soup = BeautifulSoup(html, 'html.parser')
    novas_tags = []
    lista_atual = []
    tipo_lista = None

    for el in soup.contents:
        if el.name in ['ul', 'ol']:
            if tipo_lista == el.name:
                lista_atual.extend(el.find_all('li'))
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                    lista_atual = []
                tipo_lista = el.name
                lista_atual = el.find_all('li')
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                lista_atual = []
                tipo_lista = None
            novas_tags.append(el)

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)

    soup.clear()
    for tag in novas_tags:
        soup.append(tag)

    return str(soup), titulo_extraido

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')

    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(titulo__startswith=titulo_base).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1

    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)

    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()

    return f"{titulo_base} - {total} de {total}"

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)
