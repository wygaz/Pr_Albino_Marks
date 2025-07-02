
import os
import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document


def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r' - \d+ de \d+$')
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub('', artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"


def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug


def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join('capas', novo_nome)


def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join('uploads', novo_nome)


def aplicar_formatacao(paragraph):
    partes = []
    for run in paragraph.runs:
        texto = run.text.replace('\n', '<br>')
        if not texto.strip():
            continue
        if run.bold:
            texto = f"<strong>{texto}</strong>"
        if run.italic:
            texto = f"<em>{texto}</em>"
        if run.underline:
            texto = f"<u>{texto}</u>"
        partes.append(texto)
    return ''.join(partes)


def detectar_titulo_possivel(paragraphs):
    for p in paragraphs[:10]:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith('Heading'):
            return text
        if p.alignment == 1 and len(text.split()) <= 10:
            return text
        size = 0
        for run in p.runs:
            if run.font.size and run.font.size.pt > size:
                size = run.font.size.pt
        if size > 14 and (any([r.bold for r in p.runs]) or any([r.underline for r in p.runs])):
            return text
    return None


def docx_para_html(docx_file):
    doc = Document(docx_file)
    html = ''
    titulo = detectar_titulo_possivel(doc.paragraphs)
    if not titulo:
        titulo = "Sem Título"
    for p in doc.paragraphs:
        texto = aplicar_formatacao(p).strip()
        if not texto:
            continue
        estilo = p.style.name.lower()
        if estilo.startswith('heading'):
            nivel = ''.join(filter(str.isdigit, estilo)) or '1'
            html += f'<h{nivel}>{texto}</h{nivel}>'
        elif estilo == 'normal':
            html += f'<p>{texto}</p>'
        elif estilo.startswith('list bullet'):
            html += f'<ul><li>{texto}</li></ul>'
        elif estilo.startswith('list number'):
            html += f'<ol><li>{texto}</li></ol>'
        else:
            if re.match(r'^[-*•‣◦]', texto.strip()):
                texto = re.sub(r'^[-*•‣◦]\s*', '', texto.strip())
                html += f'<ul><li>{texto}</li></ul>'
            elif re.match(r'^\(?[0-9A-Za-z]+\)?[\.\-]?', texto.strip()):
                texto = re.sub(r'^\(?[0-9A-Za-z]+\)?[\.\-]?\s*', '', texto.strip())
                html += f'<ol><li>{texto}</li></ol>'
            else:
                html += f'<p>{texto}</p>'
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup.find_all(['ul', 'ol']):
        next_tag = tag.find_next_sibling()
        while next_tag and next_tag.name == tag.name:
            tag.append(next_tag.extract())
            next_tag = tag.find_next_sibling()
    return str(soup), titulo
