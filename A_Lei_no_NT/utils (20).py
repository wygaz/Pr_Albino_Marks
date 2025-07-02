import re
import uuid
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def gerar_slug(texto):
    return slugify(texto)[:50]

def gerar_titulo_numerado():
    return f'artigo-{uuid.uuid4().hex[:8]}'

def detectar_titulo_possivel(paragrafos):
    for p in paragrafos[:5]:
        texto = ''.join(run.text for run in p.runs).strip()
        if texto and len(texto.split()) <= 12 and p.style.name.startswith('Heading'):
            return texto
    return None

def detectar_autor(paragrafos):
    for i, p in enumerate(paragrafos[1:4], start=1):
        texto = ''.join(run.text for run in p.runs).strip()
        if not texto or len(texto) > 100:
            continue
        if i in (1, 2) and len(texto.split()) <= 5 and len(p.runs) <= 3:
            return texto
        if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return texto
    return None

def remover_autor_do_conteudo(html, autor):
    if not autor:
        return html
    soup = BeautifulSoup(html, 'html.parser')
    for p in soup.find_all('p'):
        if autor in p.text:
            p.decompose()
    return str(soup)

def formatar_paragrafo(p):
    texto_formatado = ''
    for run in p.runs:
        t = run.text.replace('
', '<br>').replace(' ', '&nbsp;')
        if '<w:br w:type="page"/>' in run._element.xml:
            texto_formatado += '<hr class="page-break">'
        if run.bold:
            t = f'<strong>{t}</strong>'
        if run.italic:
            t = f'<em>{t}</em>'
        if run.underline:
            t = f'<u>{t}</u>'
        texto_formatado += t
    return texto_formatado

def converter_para_html(paragrafos):
    html = ''
    padrao_lista = re.compile(r'^(\d+\.|[ivxlc]+\.|[IVXLC]+\.|[a-zA-Z]\.|[a-zA-Z]\))\s+')
    for p in paragrafos:
        texto_bruto = ''.join(run.text for run in p.runs).strip()
        if not texto_bruto:
            continue
        texto_formatado = formatar_paragrafo(p)
        if p.style.name.lower().startswith('heading'):
            html += f'<h2>{texto_formatado}</h2>\n'
        elif padrao_lista.match(texto_bruto):
            html += f'<ol><li>{texto_formatado}</li></ol>\n'
        elif texto_bruto.startswith(('-', '*', '•')):
            html += f'<ul><li>{texto_formatado[1:].strip()}</li></ul>\n'
        else:
            html += f'<p>{texto_formatado}</p>\n'
    return html

def docx_para_html(arquivo):
    doc = Document(arquivo)
    paragrafos = doc.paragraphs
    titulo = detectar_titulo_possivel(paragrafos)
    autor = detectar_autor(paragrafos)
    html = converter_para_html(paragrafos)

    for tabela in doc.tables:
        html += '<table border="1" style="border-collapse: collapse; margin-top: 1em;">'
        for linha in tabela.rows:
            html += '<tr>'
            for celula in linha.cells:
                conteudo = '<br>'.join(p.text for p in celula.paragraphs)
                html += f'<td style="padding: 4px;">{conteudo}</td>'
            html += '</tr>'
        html += '</table>'

    html = remover_autor_do_conteudo(html, autor)

    soup = BeautifulSoup(html, 'html.parser')
    novas_tags, lista_atual, tipo_lista = [], [], None

    for el in soup.contents:
        if el.name in ['ul', 'ol']:
            if tipo_lista == el.name:
                lista_atual.extend(el.find_all('li'))
            else:
                if lista_atual:
                    nova_lista = soup.new_tag(tipo_lista)
                    for item in lista_atual:
                        nova_lista.append(item)
                    novas_tags.append(nova_lista)
                tipo_lista = el.name
                lista_atual = el.find_all('li')
        else:
            if lista_atual:
                nova_lista = soup.new_tag(tipo_lista)
                for item in lista_atual:
                    nova_lista.append(item)
                novas_tags.append(nova_lista)
                lista_atual, tipo_lista = [], None
            novas_tags.append(el)

    if lista_atual:
        nova_lista = soup.new_tag(tipo_lista)
        for item in lista_atual:
            nova_lista.append(item)
        novas_tags.append(nova_lista)

    soup.clear()
    for tag in novas_tags:
        soup.append(tag)

    return str(soup), titulo, autor