
import os
import uuid
import re
from bs4 import BeautifulSoup
from django.utils.text import slugify
from docx import Document
from django.utils.html import escape

def detectar_titulo_possivel(documento):
    for paragrafo in documento.paragraphs[:10]:
        texto = paragrafo.text.strip()
        if not texto:
            continue
        estilo = paragrafo.style.name.lower()
        if "heading" in estilo or "título" in estilo or "title" in estilo:
            return texto

        props = paragrafo.runs[0].font if paragrafo.runs else None
        if props:
            if props.size and props.size.pt >= 14:
                return texto
            if props.bold or props.underline:
                return texto
        if paragrafo.alignment:
            return texto
    return None

def docx_para_html(arquivo):
    doc = Document(arquivo)
    html = ""
    titulo = detectar_titulo_possivel(doc) or "TÍTULO NÃO DETECTADO"
    pilha_niveis = []

    for paragrafo in doc.paragraphs:
        texto = escape(paragrafo.text.strip())
        if not texto:
            continue

        estilo = paragrafo.style.name.lower()

        if estilo.startswith("heading"):
            nivel = int(estilo.replace("heading", "").strip())
            html += f"<h{nivel}>{texto}</h{nivel}>\n"
            continue

        if texto.startswith(("-", "*", "•")):
            while pilha_niveis and pilha_niveis[-1] != "ul":
                html += f"</{pilha_niveis.pop()}>\n"
            if not pilha_niveis or pilha_niveis[-1] != "ul":
                html += "<ul>\n"
                pilha_niveis.append("ul")
            html += f"<li>{texto[1:].strip()}</li>\n"
        elif re.match(r"^\d+[\).]", texto):
            while pilha_niveis and pilha_niveis[-1] != "ol":
                html += f"</{pilha_niveis.pop()}>\n"
            if not pilha_niveis or pilha_niveis[-1] != "ol":
                html += "<ol>\n"
                pilha_niveis.append("ol")
            html += f"<li>{texto.split(maxsplit=1)[1] if len(texto.split(maxsplit=1)) > 1 else ''}</li>\n"
        else:
            while pilha_niveis:
                html += f"</{pilha_niveis.pop()}>\n"
            html += f"<p>{texto}</p>\n"

    while pilha_niveis:
        html += f"</{pilha_niveis.pop()}>\n"

    soup = BeautifulSoup(html, "html.parser")
    html_final = soup.prettify()
    return html_final, titulo

def gerar_slug(titulo):
    from .models import Artigo
    base_slug = slugify(titulo)
    slug = base_slug
    contador = 1
    while Artigo.objects.filter(slug=slug).exists():
        slug = f"{base_slug}-{contador}"
        contador += 1
    return slug

def gerar_titulo_numerado(titulo_base, ordem_por='id'):
    from django.apps import apps
    Artigo = apps.get_model('A_Lei_no_NT', 'Artigo')
    padrao_numerado = re.compile(r" - \d+ de \d+$")
    artigos_com_titulo_base = Artigo.objects.filter(
        titulo__startswith=titulo_base
    ).order_by(ordem_por)
    total = artigos_com_titulo_base.count() + 1
    artigos_limpos = []
    for artigo in artigos_com_titulo_base:
        titulo_sem_num = padrao_numerado.sub("", artigo.titulo).strip()
        artigo.titulo = titulo_sem_num
        artigos_limpos.append(artigo)
    for i, artigo in enumerate(artigos_limpos, start=1):
        artigo.titulo = f"{titulo_base} - {i} de {total}"
        artigo.save()
    return f"{titulo_base} - {total} de {total}"

def renomear_imagem_capa(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:8])
    novo_nome = f"{nome_base}-capa{ext.lower()}"
    return os.path.join("capas", novo_nome)

def renomear_arquivo_word(instance, filename):
    base, ext = os.path.splitext(filename)
    nome_base = slugify(instance.slug or instance.titulo or uuid.uuid4().hex[:10])
    novo_nome = f"{nome_base}{ext.lower()}"
    return os.path.join("uploads", novo_nome)
