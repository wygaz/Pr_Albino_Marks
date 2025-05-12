from docx import Document
import os
from bs4 import BeautifulSoup

def convert_docx_to_html(docx_path, html_path):
    if not os.path.exists(docx_path):
        print(f"O arquivo {docx_path} não existe.")
        return
    
    document = Document(docx_path)
    html_content = '<html><body>'

    # Itera sobre cada parágrafo no documento
    for paragraph in document.paragraphs:
        p_html = '<p>'  # Inicia a tag de parágrafo
        for run in paragraph.runs:  # Itera sobre os 'runs' dentro do parágrafo
            if run.bold and run.italic:  # Verifica se o texto é negrito e itálico
                p_html += f'<b><i>{run.text}</i></b>'
            elif run.bold:  # Verifica se o texto é negrito
                p_html += f'<b>{run.text}</b>'
            elif run.italic:  # Verifica se o texto é itálico
                p_html += f'<i>{run.text}</i>'
            else:  # Texto normal
                p_html += run.text
        p_html += '</p>'  # Fecha a tag de parágrafo
        html_content += p_html

    # Itera sobre cada tabela no documento
    for table in document.tables:
        html_content += '<table border="1">'  # Inicia a tag de tabela
        for row in table.rows:  # Itera sobre cada linha da tabela
            html_content += '<tr>'  # Inicia a tag de linha da tabela
            for cell in row.cells:  # Itera sobre cada célula da linha
                cell_html = '<td>'  # Inicia a tag de célula
                for paragraph in cell.paragraphs:  # Itera sobre os parágrafos dentro da célula
                    for run in paragraph.runs:  # Itera sobre os 'runs' dentro do parágrafo
                        if run.bold and run.italic:  # Verifica se o texto é negrito e itálico
                            cell_html += f'<b><i>{run.text}</i></b>'
                        elif run.bold:  # Verifica se o texto é negrito
                            cell_html += f'<b>{run.text}</b>'
                        elif run.italic:  # Verifica se o texto é itálico
                            cell_html += f'<i>{run.text}</i>'
                        else:  # Texto normal
                            cell_html += run.text
                cell_html += '</td>'  # Fecha a tag de célula
                html_content += cell_html
            html_content += '</tr>'  # Fecha a tag de linha
        html_content += '</table>'  # Fecha a tag de tabela

    html_content += '</body></html>'

    # Usa BeautifulSoup para embelezar o HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    pretty_html = soup.prettify()

    with open(html_path, 'w', encoding='utf-8') as file:
        file.write(pretty_html)

    print(f"Conversão concluída: {html_path}")

# Caminhos para os arquivos
input_directory = 'C:\\Users\\Wanderley\\Documents\\Textos_Pr_Albino'
output_directory = 'C:\\Users\\Wanderley\\Apps\\Albino_Marks\\A_Lei_no_NT\\Templates\\A_Lei_no_NT\\Textos'


# Arquivos para converter
files_to_convert = [
    "1_OS_ESCRITORES_DO_NT_E_A_LEI.docx",
    "2_O_APOSTOLO_PAULO_E_A_LEI.docx",
    "3_O_NT_JESUS_E_A_LEI.docx",
    "4_JESUS_E_A_LEI-NOMOS.docx",
    "5_JESUS_NAO_REVOGOU_MAS_MAGNIFICOU_A_LEI.docx",
    "6_A_JUSTIÇA_DE_DEUS_E_A_DOS_FARISEUS.docx"
]

for file_name in files_to_convert:
    docx_path = os.path.join(input_directory, file_name)
    html_file_name = file_name.replace('.docx', '.html')
    html_path = os.path.join(output_directory, html_file_name)
    convert_docx_to_html(docx_path, html_path)
