from __future__ import annotations

import argparse
import csv
import re
import shutil
import unicodedata
from dataclasses import dataclass
from difflib import SequenceMatcher, get_close_matches
from pathlib import Path

from docx import Document


@dataclass
class MatchResult:
    serie_id: int
    serie_nome: str
    ordem_na_serie: int
    titulo_canonico: str
    metodo: str
    score: float


def strip_accents(s: str) -> str:
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", s or "")
        if not unicodedata.combining(ch)
    )


def normalize(s: str) -> str:
    s = strip_accents(s or "").lower()
    s = s.replace("“", '"').replace("”", '"').replace("’", "'")
    s = s.replace("_", " ")
    s = re.sub(r"\.docx?$", "", s, flags=re.I)
    s = re.sub(r"^\d{4}-\d{2}-\d{2}\s*[-_ ]*", "", s)
    s = re.sub(r"^\d+\s*[-_. ]+\s*", "", s)
    s = re.sub(r"_[A-Za-z0-9]{7}$", "", s)
    s = re.sub(r"\bsm\s*\d+\b", "", s, flags=re.I)
    s = re.sub(r"[^\w\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def clean_filename_title(name: str) -> str:
    name = Path(name).stem
    name = re.sub(r"_[A-Za-z0-9]{7}$", "", name)
    name = re.sub(r"^\d{4}-\d{2}-\d{2}_", "", name)
    return name.replace("_", " ").strip()


def docx_internal_title(path: Path) -> str | None:
    try:
        doc = Document(str(path))
    except Exception:
        return None

    try:
        title = (doc.core_properties.title or "").strip()
        if title:
            return title
    except Exception:
        pass

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            return txt
    return None


def extract_series_map(esboco_path: Path):
    doc = Document(str(esboco_path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    def find_idx_by_contains(*terms, start=0):
        termos = [normalize(t) for t in terms]
        for i in range(start, len(paras)):
            n = normalize(paras[i])
            if all(t in n for t in termos):
                return i
        raise ValueError(f"Nao encontrei cabecalho com termos: {terms}")

    idx1 = find_idx_by_contains("grande conflito", "cristo", "satanas")
    idx2 = find_idx_by_contains("apocalipse 17", "principio do simbolismo")
    idx3 = find_idx_by_contains("torah", "nomos", "mandamentos", "ordenancas", "graca")

    idx4_all = [
        i for i, p in enumerate(paras)
        if "biblia" in normalize(p) and "historia da humanidade" in normalize(p)
    ]
    if len(idx4_all) < 2:
        raise ValueError("Nao encontrei as duas ocorrencias de 'A Biblia e a Historia da Humanidade'.")
    idx4_first, idx4_second = idx4_all[0], idx4_all[1]

    series = {
        1: {"nome": "O GRANDE CONFLITO ENTRE CRISTO E SATANÁS", "titulos": []},
        2: {"nome": "APOCALIPSE 17 E O PRINCÍPIO DO SIMBOLISMO", "titulos": []},
        3: {"nome": "A LEI, TORAH, NÓMOS, MANDAMENTOS, ORDENANÇAS E A GRAÇA", "titulos": []},
        4: {"nome": "A BÍBLIA E A HISTÓRIA DA HUMANIDADE", "titulos": []},
    }

    for line in paras[idx1 + 1:idx2]:
        if normalize(line) == "esboco":
            continue
        series[1]["titulos"].append(line)

    for line in paras[idx2 + 1:idx3]:
        if normalize(line) == "esboco":
            continue
        series[2]["titulos"].append(line)

    for line in paras[idx3:idx4_first]:
        if normalize(line) == "esboco":
            continue
        series[3]["titulos"].append(line)

    for line in paras[idx4_second:]:
        if normalize(line) == "esboco":
            continue
        series[4]["titulos"].append(line)

    normalized_index = {}
    for serie_id, data in series.items():
        for ordem, titulo in enumerate(data["titulos"], start=1):
            normalized_index[normalize(titulo)] = (serie_id, data["nome"], ordem, titulo)

    return series, normalized_index


def match_title(candidate: str, normalized_index: dict[str, tuple], cutoff: float = 0.84) -> MatchResult | None:
    cand_norm = normalize(candidate)
    if not cand_norm:
        return None

    if cand_norm in normalized_index:
        serie_id, serie_nome, ordem, titulo = normalized_index[cand_norm]
        return MatchResult(serie_id, serie_nome, ordem, titulo, "exato", 1.0)

    keys = list(normalized_index.keys())

    if all(t in cand_norm for t in ["torah", "nomos", "mandamentos", "ordenancas", "graca"]):
        cand_norm = "a lei torah nomos mandamentos ordenancas e a graca"

    for key in keys:
        if cand_norm == key or cand_norm in key or key in cand_norm:
            serie_id, serie_nome, ordem, titulo = normalized_index[key]
            score = SequenceMatcher(None, cand_norm, key).ratio()
            if score >= 0.72:
                return MatchResult(serie_id, serie_nome, ordem, titulo, "substring", score)

    close = get_close_matches(cand_norm, keys, n=1, cutoff=cutoff)
    if close:
        key = close[0]
        serie_id, serie_nome, ordem, titulo = normalized_index[key]
        score = SequenceMatcher(None, cand_norm, key).ratio()
        return MatchResult(serie_id, serie_nome, ordem, titulo, "aproximado", score)

    best_key = None
    best_score = 0.0
    for key in keys:
        score = SequenceMatcher(None, cand_norm, key).ratio()
        if score > best_score:
            best_score = score
            best_key = key
    if best_key and best_score >= 0.72:
        serie_id, serie_nome, ordem, titulo = normalized_index[best_key]
        return MatchResult(serie_id, serie_nome, ordem, titulo, "similaridade", best_score)

    return None


def main():
    ap = argparse.ArgumentParser(description="Distribui DOCX em 4 diretórios de série com base no ESBOCO_Geral_Series_1_a_4.docx.")
    ap.add_argument("--fonte", required=True, help="Pasta com os DOCX a classificar.")
    ap.add_argument("--esboco", required=True, help="Caminho do ESBOCO_Geral_Series_1_a_4.docx.")
    ap.add_argument("--saida", required=True, help="Pasta-raiz de saída para as séries.")
    ap.add_argument("--mover", action="store_true", help="Move os arquivos em vez de copiar.")
    args = ap.parse_args()

    fonte = Path(args.fonte)
    esboco = Path(args.esboco)
    saida = Path(args.saida)

    if not fonte.exists():
        raise FileNotFoundError(f"Pasta fonte não encontrada: {fonte}")
    if not esboco.exists():
        raise FileNotFoundError(f"Arquivo de esboço não encontrado: {esboco}")

    series, normalized_index = extract_series_map(esboco)

    saida.mkdir(parents=True, exist_ok=True)
    serie_dirs = {}
    for serie_id, data in series.items():
        dir_name = f"Serie_{serie_id:02d} - {data['nome']}"
        dir_name = re.sub(r'[<>:"/\\|?*]', "", dir_name)
        serie_dir = saida / dir_name
        serie_dir.mkdir(parents=True, exist_ok=True)
        serie_dirs[serie_id] = serie_dir

    csv_path = saida / "classificacao_series.csv"
    nao_classificados = []
    total = 0
    classificados = 0

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow([
            "arquivo_docx", "titulo_do_arquivo", "titulo_interno", "serie_id",
            "serie_nome", "ordem_na_serie", "titulo_canonico", "metodo", "score", "destino"
        ])

        for path in sorted(fonte.glob("*.docx")):
            total += 1
            titulo_arquivo = clean_filename_title(path.name)
            titulo_interno = docx_internal_title(path) or ""

            match = None
            for candidate in [titulo_interno, titulo_arquivo]:
                if candidate:
                    match = match_title(candidate, normalized_index)
                    if match:
                        break

            if match:
                classificados += 1
                prefixo = f"{match.ordem_na_serie:02d} - "
                destino = serie_dirs[match.serie_id] / f"{prefixo}{path.name}"

                if args.mover:
                    shutil.move(str(path), str(destino))
                else:
                    shutil.copy2(str(path), str(destino))

                w.writerow([
                    path.name, titulo_arquivo, titulo_interno, match.serie_id,
                    match.serie_nome, match.ordem_na_serie, match.titulo_canonico,
                    match.metodo, f"{match.score:.3f}", str(destino)
                ])
            else:
                nao_classificados.append(path.name)
                w.writerow([
                    path.name, titulo_arquivo, titulo_interno, "", "", "", "",
                    "nao_classificado", "", ""
                ])

    if nao_classificados:
        (saida / "nao_classificados.txt").write_text(
            "\n".join(nao_classificados) + "\n", encoding="utf-8"
        )

    print(f"Total de DOCX na fonte : {total}")
    print(f"Classificados         : {classificados}")
    print(f"Nao classificados     : {len(nao_classificados)}")
    print(f"CSV                   : {csv_path}")
    if nao_classificados:
        print(f"Lista pendentes       : {saida / 'nao_classificados.txt'}")


if __name__ == "__main__":
    main()