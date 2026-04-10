import argparse
import html
import re
from pathlib import Path
from typing import List, Tuple, Dict

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SECTION_RE = re.compile(r"^(\d+)\)\s+(.*)$")


def set_cell_shading(cell, fill: str = "F3F4F6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def parse_sermao_md(text: str) -> Dict:
    lines = text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
    sections: List[Tuple[str, List[str]]] = []
    current_title = None
    current_lines: List[str] = []

    for raw in lines:
        line = raw.rstrip()
        m = SECTION_RE.match(line.strip())
        if m:
            if current_title is not None:
                sections.append((current_title, current_lines))
            current_title = m.group(2).strip()
            current_lines = []
        else:
            if current_title is None:
                continue
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, current_lines))

    data = {"sections": sections}

    def cleaned(lines: List[str]) -> List[str]:
        out = []
        for l in lines:
            s = l.strip()
            if s:
                out.append(s)
        return out

    sec_map = {title.lower(): cleaned(lines) for title, lines in sections}
    data["title"] = sec_map.get("título do sermão", [""])[0] if sec_map.get("título do sermão") else "Sermão"
    data["verso_chave"] = sec_map.get("verso-chave", [""])[0] if sec_map.get("verso-chave") else ""
    data["oracao"] = sec_map.get("sugestão breve de oração", [""])[0] if sec_map.get("sugestão breve de oração") else ""
    data["leitura"] = sec_map.get("leitura do texto bíblico central", [""])[0] if sec_map.get("leitura do texto bíblico central") else ""
    return data


def md_lines_to_html(lines: List[str]) -> str:
    chunks = []
    in_list = False
    for line in lines:
        s = line.strip()
        if not s:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            continue
        if s.startswith("- "):
            if not in_list:
                chunks.append('<ul class="bullet-list">')
                in_list = True
            chunks.append(f"<li>{html.escape(s[2:].strip())}</li>")
        else:
            if in_list:
                chunks.append("</ul>")
                in_list = False
            chunks.append(f"<p>{html.escape(s)}</p>")
    if in_list:
        chunks.append("</ul>")
    return "\n".join(chunks)


def build_html(data: Dict, page_mode: str) -> str:
    title = html.escape(data["title"])
    verse = html.escape(data["verso_chave"])
    prayer = html.escape(data["oracao"])
    leitura = html.escape(data["leitura"])

    if page_mode == "tablet":
        page_css = """
        body { background:#f5f7fb; font-family: Georgia, 'Times New Roman', serif; color:#1f2937; }
        .page { max-width: 820px; margin: 0 auto; background:#fff; min-height:100vh; padding: 28px 22px 44px; }
        h1 { font-size: 2rem; line-height:1.2; margin:0 0 14px; }
        h2 { font-size:1.15rem; color:#1d4ed8; margin:28px 0 10px; }
        p, li { font-size: 1.18rem; line-height:1.72; }
        .meta { background:#f8fafc; border-left:4px solid #1d4ed8; padding:14px 16px; border-radius:8px; margin:16px 0; }
        .section { margin-top: 18px; }
        """
        extra = ""
    elif page_mode == "a5":
        page_css = """
        @page { size: A5; margin: 1.3cm; }
        body { font-family: Georgia, 'Times New Roman', serif; color:#111827; }
        .page { max-width: 100%; margin: 0 auto; }
        h1 { font-size: 18pt; line-height:1.15; margin:0 0 10pt; }
        h2 { font-size: 11.5pt; color:#1f2937; margin:16pt 0 6pt; }
        p, li { font-size: 10.8pt; line-height:1.45; }
        .meta { border:1px solid #d1d5db; padding:9pt 10pt; border-radius:6pt; margin:10pt 0; }
        .section { break-inside: avoid-page; }
        """
        extra = ""
    else:
        page_css = """
        @page { size: A4; margin: 2cm 1.8cm 2cm 1.8cm; }
        body { font-family: Georgia, 'Times New Roman', serif; color:#111827; }
        .page { max-width: 100%; margin: 0 auto; }
        h1 { font-size: 22pt; line-height:1.15; margin:0 0 12pt; }
        h2 { font-size: 13pt; color:#1f2937; margin:18pt 0 7pt; }
        p, li { font-size: 11.3pt; line-height:1.5; }
        .meta { border:1px solid #d1d5db; padding:11pt 12pt; border-radius:6pt; margin:12pt 0; }
        .section { break-inside: avoid-page; }
        """
        extra = ""

    sections_html = []
    for title_sec, lines in data["sections"]:
        lower = title_sec.lower()
        if lower in {"título do sermão", "verso-chave", "sugestão breve de oração", "leitura do texto bíblico central"}:
            continue
        sections_html.append(f'<section class="section"><h2>{html.escape(title_sec)}</h2>{md_lines_to_html(lines)}</section>')

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{title}</title>
<style>
{page_css}
body {{ margin:0; }}
.bullet-list {{ padding-left: 1.1rem; margin: 0.35rem 0 0.9rem; }}
.lead {{ font-style: italic; color:#374151; }}
.label {{ font-weight:700; }}
{extra}
</style>
</head>
<body>
  <main class="page">
    <h1>{title}</h1>
    <div class="meta"><span class="label">Verso-chave:</span> {verse}</div>
    <div class="meta"><span class="label">Sugestão breve de oração:</span> {prayer}</div>
    <div class="meta"><span class="label">Texto bíblico central:</span> {leitura}</div>
    {''.join(sections_html)}
  </main>
</body>
</html>
"""


def add_docx_paragraphs_from_lines(doc: Document, lines: List[str], body_style: str = "Body Text"):
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s[2:].strip())
            set_paragraph_spacing(p, after=2, line=1.2)
        else:
            p = doc.add_paragraph(style=body_style)
            p.add_run(s)
            set_paragraph_spacing(p, after=4, line=1.22)


def build_docx_a4(data: Dict, outpath: Path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Georgia'
    normal.font.size = Pt(11)

    if 'Body Text' not in [s.name for s in styles]:
        styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    body = styles['Body Text']
    body.base_style = styles['Normal']
    body.font.name = 'Georgia'
    body.font.size = Pt(11)

    if 'MetaBox' not in [s.name for s in styles]:
        styles.add_style('MetaBox', WD_STYLE_TYPE.PARAGRAPH)
    meta = styles['MetaBox']
    meta.base_style = styles['Normal']
    meta.font.name = 'Georgia'
    meta.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data['title'])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    set_paragraph_spacing(title, after=10, line=1.1)

    for label, value in [
        ('Verso-chave', data['verso_chave']),
        ('Sugestão breve de oração', data['oracao']),
        ('Texto bíblico central', data['leitura']),
    ]:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        set_cell_shading(cell, 'F6F8FB')
        p = cell.paragraphs[0]
        p.style = meta
        r1 = p.add_run(f'{label}: ')
        r1.bold = True
        r2 = p.add_run(value)
        set_paragraph_spacing(p, after=0, line=1.15)
        doc.add_paragraph()

    for title_sec, lines in data['sections']:
        lower = title_sec.lower()
        if lower in {"título do sermão", "verso-chave", "sugestão breve de oração", "leitura do texto bíblico central"}:
            continue
        h = doc.add_paragraph()
        hr = h.add_run(title_sec)
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        set_paragraph_spacing(h, before=8, after=4, line=1.1)
        add_docx_paragraphs_from_lines(doc, lines)

    doc.save(outpath)


def export_formats(md_path: Path, outdir: Path):
    text = md_path.read_text(encoding='utf-8')
    data = parse_sermao_md(text)
    base = md_path.stem
    outdir.mkdir(parents=True, exist_ok=True)

    html_tablet = outdir / f"{base}__tablet.html"
    html_a4 = outdir / f"{base}__A4.html"
    html_a5 = outdir / f"{base}__A5.html"
    docx_a4 = outdir / f"{base}__A4.docx"

    html_tablet.write_text(build_html(data, 'tablet'), encoding='utf-8')
    html_a4.write_text(build_html(data, 'a4'), encoding='utf-8')
    html_a5.write_text(build_html(data, 'a5'), encoding='utf-8')
    build_docx_a4(data, docx_a4)

    return html_tablet, html_a4, html_a5, docx_a4


def main():
    ap = argparse.ArgumentParser(description='Exporta um sermão em .md para HTML tablet, HTML A4/A5 e DOCX A4.')
    ap.add_argument('--md', required=True, help='Arquivo .md do sermão')
    ap.add_argument('--outdir', required=True, help='Pasta de saída')
    args = ap.parse_args()

    md_path = Path(args.md)
    outdir = Path(args.outdir)

    if not md_path.exists():
        raise FileNotFoundError(f'Arquivo não encontrado: {md_path}')

    outputs = export_formats(md_path, outdir)
    for p in outputs:
        print(f'[OK] {p}')


if __name__ == '__main__':
    main()
