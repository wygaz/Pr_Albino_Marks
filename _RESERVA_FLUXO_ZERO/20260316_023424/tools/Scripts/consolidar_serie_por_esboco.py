import os
import re
import csv
import shutil
import unicodedata
import argparse
from pathlib import Path
from difflib import SequenceMatcher

try:
    from docx import Document
except ImportError:
    Document = None


def strip_accents(s: str) -> str:
    s = unicodedata.normalize("NFKD", s or "")
    return "".join(ch for ch in s if not unicodedata.combining(ch))


def is_esboco_heading(line: str) -> bool:
    raw = strip_accents(line).upper().strip()
    key = re.sub(r"[^A-Z]", "", raw)
    return key == "ESBOCO"


DATE_DIR_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
WS_RE = re.compile(r"\s+")
INVALID_WIN = re.compile(r'[<>:"/\\|?*]')
SM_RE = re.compile(r"^\s*sm\s*\d+\s*[-–—:]\s*|\s*^\s*sm\s*\d+\s*", re.IGNORECASE)
DATE_RE1 = re.compile(r"^\s*\d{4}-\d{2}-\d{2}\s*[-–—:]\s*")
DATE_RE2 = re.compile(r"^\s*\d{2}-\d{2}-\d{4}\s*[-–—:]\s*")

STOPWORDS = {
    "A", "O", "OS", "AS", "UM", "UMA", "UNS", "UMAS",
    "DE", "DO", "DA", "DOS", "DAS", "E", "EM", "NO", "NA", "NOS", "NAS",
    "PARA", "POR", "COM", "SEM", "AO", "AOS", "À", "ÀS", "ENTRE"
}


LOWER_WORDS = {
    "a", "o", "os", "as", "um", "uma", "uns", "umas",
    "de", "do", "da", "dos", "das", "e", "em", "no", "na", "nos", "nas",
    "para", "por", "com", "sem", "ao", "aos", "à", "às", "entre", "ou"
}


def clean_title(s: str) -> str:
    s = (s or "").strip()
    s = DATE_RE1.sub("", s)
    s = DATE_RE2.sub("", s)
    s = SM_RE.sub("", s)
    s = WS_RE.sub(" ", s).strip()
    return s



def safe_series_name(name: str, max_len: int = 80) -> str:
    name = clean_title(name).upper()
    name = INVALID_WIN.sub("", name).strip().rstrip(".")
    name = WS_RE.sub(" ", name).strip()
    if not name:
        name = "SERIE_SEM_NOME"
    return name[:max_len] if len(name) > max_len else name



def norm_key(s: str) -> str:
    s = clean_title(s)
    s = strip_accents(s).upper()
    s = re.sub(r"[^A-Z0-9\s]", " ", s)
    s = WS_RE.sub(" ", s).strip()
    parts = [p for p in s.split(" ") if p and p not in STOPWORDS]
    return " ".join(parts)



def human_title(s: str) -> str:
    s = clean_title(s)
    if not s:
        return s
    words = s.split()
    out = []
    for i, w in enumerate(words):
        wl = w.lower()
        if i > 0 and wl in LOWER_WORDS:
            out.append(wl)
        else:
            out.append(w[:1].upper() + w[1:].lower())
    return " ".join(out)



def pick_latest_date_folder(base: Path) -> str | None:
    dates = []
    for d in base.iterdir():
        if d.is_dir() and DATE_DIR_RE.match(d.name):
            dates.append(d.name)
    return max(dates) if dates else None



def _strip_accents(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    return "".join(ch for ch in s if not unicodedata.combining(ch))



def find_esboco_file(folder: Path) -> Path | None:
    candidates = []
    for p in folder.iterdir():
        if p.is_file() and p.suffix.lower() in (".txt", ".docx"):
            stem_norm = _strip_accents(p.stem).upper()
            if "ESBOC" in stem_norm:
                candidates.append(p)

    if candidates:
        candidates.sort(key=lambda x: (0 if x.suffix.lower() == ".txt" else 1, x.name))
        return candidates[0]
    return None



def find_latest_esboco_in_series(series_dir: Path) -> Path | None:
    cands = list(series_dir.glob("ESBOCO_*.*")) + list(series_dir.glob("*ESBOCO*.*"))
    cands = [c for c in cands if c.suffix.lower() in (".txt", ".docx")]
    if not cands:
        return None
    cands.sort(key=lambda p: (p.name, p.stat().st_mtime))
    return cands[-1]



def read_esboco_titles(esboco: Path) -> list[str]:
    titles = []

    if esboco.suffix.lower() == ".txt":
        for ln in esboco.read_text(encoding="utf-8", errors="ignore").splitlines():
            ln = ln.strip().strip("\ufeff")
            if not ln or ln.startswith("#"):
                continue
            if is_esboco_heading(ln):
                continue
            titles.append(ln)
        return titles

    if esboco.suffix.lower() == ".docx":
        if not Document:
            raise RuntimeError("python-docx não instalado (pip install python-docx) para ler ESBOCO.docx.")
        doc = Document(str(esboco))
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue
            if re.fullmatch(r"\d+", t):
                continue
            if is_esboco_heading(t):
                continue
            titles.append(t)
        return titles

    return titles



def best_match(target_title: str, candidates: list[Path]) -> tuple[Path | None, float]:
    tkey = norm_key(target_title)
    if not tkey:
        return None, 0.0

    best = None
    best_r = 0.0
    for f in candidates:
        fkey = norm_key(f.stem)
        if not fkey:
            continue
        r = SequenceMatcher(None, tkey, fkey).ratio()
        if r > best_r:
            best_r = r
            best = f
    return best, best_r



def load_last_series(scripts_dir: Path) -> str | None:
    p = scripts_dir / ".last_series.txt"
    if p.exists():
        return p.read_text(encoding="utf-8", errors="ignore").strip() or None
    return None



def save_last_series(scripts_dir: Path, name: str) -> None:
    (scripts_dir / ".last_series.txt").write_text(name, encoding="utf-8")



def count_docx_words(docx_path: Path) -> int:
    if not Document or not docx_path.exists():
        return 0
    try:
        doc = Document(str(docx_path))
        texts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                texts.append(t)
        full = "\n".join(texts)
        return len(re.findall(r"\S+", full))
    except Exception:
        return 0



def main():
    ap = argparse.ArgumentParser(description="Consolida um lote (YYYY-MM-DD) em uma pasta de série usando ESBOCO e fuzzy match.")
    ap.add_argument("--lote", help="Nome da pasta do lote (YYYY-MM-DD). Se omitido, usa a mais recente.")
    ap.add_argument("--series", help="Nome fixo da série (opcional). Se omitido, usa heurística do 1º item do ESBOCO.")
    ap.add_argument("--continue-series", action="store_true", help="Força continuar a última série (.last_series.txt) se existir.")
    ap.add_argument("--threshold", type=float, default=0.84, help="Limiar fuzzy (default 0.84).")
    ap.add_argument("--serie-codigo", default="SR01", help="Código curto da série (ex.: SR01).")
    ap.add_argument("--gerar-manifest-sermoes", action="store_true",
                    help="Gera também um manifest_sermoes.csv simples para controle do fluxo dos sermões.")
    args = ap.parse_args()

    scripts_dir = Path(__file__).resolve().parent
    base = scripts_dir.parent
    series_root = base / "SERIES"
    series_root.mkdir(exist_ok=True)

    latest = pick_latest_date_folder(base)
    lote = (args.lote or latest or "").strip()
    if not lote:
        raise SystemExit("Nenhuma pasta YYYY-MM-DD encontrada e você não informou --lote.")
    lote_dir = base / lote
    if not lote_dir.exists():
        raise SystemExit(f"Pasta do lote não encontrada: {lote_dir}")

    last_series = load_last_series(scripts_dir)

    series_name = None
    if args.continue_series and last_series:
        series_name = last_series
    elif args.series:
        series_name = args.series

    esboco = find_esboco_file(lote_dir)

    if not esboco and not series_name and last_series:
        resp = input(f"Não achei ESBOCO no lote. Usar ESBOCO da série anterior '{last_series}'? [S/n]: ").strip().lower()
        if resp in ("", "s", "sim"):
            series_name = last_series

    temp_series_dir = None
    if series_name:
        series_name = safe_series_name(series_name)
        temp_series_dir = series_root / series_name
        temp_series_dir.mkdir(parents=True, exist_ok=True)

    if not esboco and temp_series_dir:
        esboco = find_latest_esboco_in_series(temp_series_dir)

    if not esboco:
        raise SystemExit("Não encontrei ESBOCO no lote e não achei ESBOCO anterior na série.")

    titles = read_esboco_titles(esboco)
    if not titles:
        raise SystemExit("ESBOCO está vazio ou não consegui extrair títulos.")

    suggested_series = safe_series_name(titles[0])

    if not series_name:
        if last_series:
            resp = input(f"Continuar a série anterior '{last_series}'? [S/n]: ").strip().lower()
            if resp in ("", "s", "sim"):
                series_name = last_series
            else:
                series_name = input(f"Nome da série [default: {suggested_series}]: ").strip() or suggested_series
        else:
            series_name = input(f"Nome da série [default: {suggested_series}]: ").strip() or suggested_series

    series_name = safe_series_name(series_name)
    save_last_series(scripts_dir, series_name)

    series_dir = series_root / series_name
    docx_dir = series_dir / "DOCX"
    pdf_dir = series_dir / "PDF"
    img_dir = series_dir / "IMG"
    for d in (series_dir, docx_dir, pdf_dir, img_dir):
        d.mkdir(parents=True, exist_ok=True)

    stamp = lote
    esboco_copy = series_dir / f"ESBOCO_{stamp}{esboco.suffix.lower()}"
    if esboco.parent.resolve() == lote_dir.resolve():
        if not esboco_copy.exists():
            shutil.copy2(esboco, esboco_copy)

    lote_docx = sorted(lote_dir.glob("*.docx"))
    lote_imgs = []
    for ext in ("*.jpg", "*.jpeg", "*.png", "*.webp"):
        lote_imgs.extend(lote_dir.glob(ext))
    lote_imgs = sorted(lote_imgs)

    existing_docx = sorted(docx_dir.glob("*.docx"))
    all_docx_candidates = lote_docx + existing_docx

    manifest = []
    manifest_sermoes = []
    THRESH = float(args.threshold)
    serie_codigo = (args.serie_codigo or "SR01").strip().upper()

    for idx, title in enumerate(titles, start=1):
        match, ratio = best_match(title, all_docx_candidates)

        status = "FALTANDO"
        src_docx = ""
        dst_docx = ""
        notes = ""
        palavra_count = 0

        if match and ratio >= THRESH:
            status = "OK"
            src_docx = str(match)

            target_name = match.name
            dst_path = docx_dir / target_name
            dst_docx = str(dst_path)

            if match.parent.resolve() != docx_dir.resolve():
                if not dst_path.exists():
                    shutil.copy2(match, dst_path)

                pdf_src = match.with_suffix(".pdf")
                if pdf_src.exists():
                    pdf_dst = pdf_dir / pdf_src.name
                    if not pdf_dst.exists():
                        shutil.copy2(pdf_src, pdf_dst)

                for ext in (".png", ".jpg", ".jpeg", ".webp"):
                    img_src = match.with_suffix(ext)
                    if img_src.exists():
                        img_dst = img_dir / img_src.name
                        if not img_dst.exists():
                            shutil.copy2(img_src, img_dst)
            else:
                notes = "Já existia na série."

            palavra_count = count_docx_words(dst_path if dst_path.exists() else match)

        elif match:
            status = "DUVIDOSO"
            src_docx = str(match)
            notes = f"Melhor candidato abaixo do limiar (ratio={ratio:.2f}). Ajuste manual ou renomeie."
            palavra_count = count_docx_words(match)
        else:
            notes = "Nenhum candidato encontrado."

        manifest.append({
            "ordem": idx,
            "titulo_esboco": title,
            "chave_normalizada": norm_key(title),
            "status": status,
            "match_ratio": f"{ratio:.3f}",
            "origem_docx": src_docx,
            "destino_docx": dst_docx,
            "serie": series_name,
            "lote": lote,
            "notes": notes,
        })

        if args.gerar_manifest_sermoes:
            artigo_codigo = f"A{idx:03d}"
            id_item = f"{serie_codigo}-{artigo_codigo}"
            titulo_normalizado = human_title(Path(dst_docx).stem if dst_docx else title)
            arquivo_fonte = Path(dst_docx).name if dst_docx else ""

            if status == "OK":
                status_sermao = "PENDENTE"
                normalizado = "S"
            elif status == "DUVIDOSO":
                status_sermao = "REVISAR_MATCH"
                normalizado = "N"
            else:
                status_sermao = "SEM_FONTE"
                normalizado = "N"

            manifest_sermoes.append({
                "ordem": idx,
                "serie_codigo": serie_codigo,
                "serie_nome": series_name,
                "artigo_codigo": artigo_codigo,
                "id_item": id_item,
                "titulo": title,
                "titulo_normalizado": titulo_normalizado,
                "arquivo_fonte": arquivo_fonte,
                "palavras_artigo": palavra_count,
                "status_match": status,
                "normalizado": normalizado,
                "segmentado": "N",
                "prompt_ok": "N",
                "sermao_docx": "N",
                "pdf_sermao": "N",
                "aprovado": "N",
                "status_sermao": status_sermao,
                "observacoes": notes,
            })

    manifest_path = series_dir / "manifest.csv"
    with open(manifest_path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=list(manifest[0].keys()), delimiter=";")
        w.writeheader()
        w.writerows(manifest)

    manifest_sermoes_path = None
    if args.gerar_manifest_sermoes:
        manifest_sermoes_path = series_dir / "manifest_sermoes.csv"
        with open(manifest_sermoes_path, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=list(manifest_sermoes[0].keys()), delimiter=";")
            w.writeheader()
            w.writerows(manifest_sermoes)

    ok = sum(1 for r in manifest if r["status"] == "OK")
    falt = sum(1 for r in manifest if r["status"] == "FALTANDO")
    duv = sum(1 for r in manifest if r["status"] == "DUVIDOSO")

    print("\n✅ Série:", series_name)
    print("📦 Lote:", lote)
    print("📄 Itens ESBOÇO:", len(manifest), "| OK:", ok, "| FALTANDO:", falt, "| DUVIDOSO:", duv)
    print("🧾 manifest.csv:", manifest_path)
    if manifest_sermoes_path:
        print("🪵 manifest_sermoes.csv:", manifest_sermoes_path)


if __name__ == "__main__":
    main()
