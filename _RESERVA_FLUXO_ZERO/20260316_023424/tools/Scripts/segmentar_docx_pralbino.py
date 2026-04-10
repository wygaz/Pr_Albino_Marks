# segmentar_docx_pralbino.py
# Requisitos: pip install python-docx pandas
from __future__ import annotations

import argparse
import csv
import json
import os
import re
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document

# ---------- Heurísticas ----------
HEADING_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)

# Alguns Word em PT-BR podem usar nomes como "Título 1", "Título 2" etc.
PT_HEADING_RE = re.compile(r"^(T[ií]tulo)\s*(\d+)$", re.IGNORECASE)

def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s, flags=re.UNICODE)
    s = re.sub(r"[\s_-]+", "-", s)
    return s.strip("-")[:80] or "secao"

def clean_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def guess_heading_level(paragraph) -> Optional[int]:
    """
    Detecta nível de heading por estilo do parágrafo.
    Retorna int (1..9) ou None.
    """
    style_name = (paragraph.style.name or "").strip()

    m = HEADING_RE.match(style_name)
    if m:
        return int(m.group(1))

    m = PT_HEADING_RE.match(style_name)
    if m:
        return int(m.group(2))

    # Fallback heurístico: linha curta em CAIXA ALTA, isolada, sem ponto final.
    txt = (paragraph.text or "").strip()
    if not txt:
        return None
    if len(txt) <= 90 and txt.isupper() and not txt.endswith("."):
        return 2  # chute razoável: subtítulo
    return None

@dataclass
class Secao:
    artigo_arquivo: str
    artigo_slug: str
    titulo_documento: str
    area: str
    secao_id: str
    caminho: str                 # Ex.: "CAPÍTULO 1 > A LEI > A GRAÇA"
    nivel: int
    titulo_secao: str
    texto: str
    n_chars: int
    n_palavras: int

def contar_palavras(txt: str) -> int:
    return len(re.findall(r"\S+", txt))

def dividir_por_tamanho(paragrafos: List[str], alvo_palavras: int) -> List[List[str]]:
    """
    Divide lista de parágrafos em blocos aproximados de alvo_palavras.
    Mantém parágrafos inteiros (não corta no meio).
    """
    blocos = []
    atual = []
    w = 0
    for p in paragrafos:
        pw = contar_palavras(p)
        if atual and (w + pw) > alvo_palavras:
            blocos.append(atual)
            atual = []
            w = 0
        atual.append(p)
        w += pw
    if atual:
        blocos.append(atual)
    return blocos

def extrair_secoes_docx(
    docx_path: Path,
    area: str,
    artigo_slug: Optional[str],
    split_long: bool,
    max_words_section: int,
    split_target_words: int,
) -> List[Secao]:
    doc = Document(str(docx_path))

    # título do documento: primeira linha não vazia
    titulo_documento = ""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            titulo_documento = t
            break

    artigo_arquivo = docx_path.name
    artigo_slug = artigo_slug or slugify(docx_path.stem)

    # Varre parágrafos e cria blocos por headings
    stack: List[Tuple[int, str]] = []   # (level, title)
    current_level = 0
    current_title = "INÍCIO"
    buffer: List[str] = []
    secoes_raw: List[Tuple[int, str, List[str], List[Tuple[int, str]]]] = []

    def flush():
        nonlocal buffer, current_level, current_title, stack
        txt = clean_text("\n\n".join([b for b in buffer if b.strip()]))
        if txt:
            secoes_raw.append((current_level, current_title, buffer.copy(), stack.copy()))
        buffer = []

    for p in doc.paragraphs:
        txt = (p.text or "").strip()

        lvl = guess_heading_level(p)
        if lvl is not None and txt:
            # Encontrou um heading: fecha seção anterior
            flush()

            # Atualiza stack hierárquico
            while stack and stack[-1][0] >= lvl:
                stack.pop()
            stack.append((lvl, txt))

            current_level = lvl
            current_title = txt
            continue

        # Conteúdo normal
        if txt:
            buffer.append(txt)

    flush()

    # Converte em Secao(s), com possível split por tamanho
    secoes: List[Secao] = []
    for idx, (lvl, title, paras, st) in enumerate(secoes_raw, start=1):
        caminho = " > ".join([t for _, t in st]) if st else title
        texto = clean_text("\n\n".join(paras))
        n_words = contar_palavras(texto)

        # Split opcional para seções grandes
        if split_long and n_words > max_words_section:
            blocos = dividir_por_tamanho(paras, split_target_words)
            for bidx, bloco in enumerate(blocos, start=1):
                btxt = clean_text("\n\n".join(bloco))
                sec_id = f"{artigo_slug}__{slugify(title)}__p{idx:03d}_{bidx:02d}"
                secoes.append(
                    Secao(
                        artigo_arquivo=artigo_arquivo,
                        artigo_slug=artigo_slug,
                        titulo_documento=titulo_documento,
                        area=area,
                        secao_id=sec_id,
                        caminho=caminho + f" (parte {bidx})",
                        nivel=lvl or 9,
                        titulo_secao=title,
                        texto=btxt,
                        n_chars=len(btxt),
                        n_palavras=contar_palavras(btxt),
                    )
                )
        else:
            sec_id = f"{artigo_slug}__{slugify(title)}__p{idx:03d}"
            secoes.append(
                Secao(
                    artigo_arquivo=artigo_arquivo,
                    artigo_slug=artigo_slug,
                    titulo_documento=titulo_documento,
                    area=area,
                    secao_id=sec_id,
                    caminho=caminho,
                    nivel=lvl or 9,
                    titulo_secao=title,
                    texto=texto,
                    n_chars=len(texto),
                    n_palavras=n_words,
                )
            )

    return secoes

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-dir", required=True, help="Pasta com .docx (já normalizados/consolidados)")
    ap.add_argument("--saida", required=True, help="Pasta de saída")
    ap.add_argument("--area", default="", help="Nome da área (se vazio, tenta inferir pelo nome da pasta)")
    ap.add_argument("--split-long", action="store_true", help="Divide seções muito longas em partes menores")
    ap.add_argument("--max-words-section", type=int, default=3500, help="Seção acima disso será dividida (se --split-long)")
    ap.add_argument("--split-target-words", type=int, default=1800, help="Tamanho alvo (palavras) por parte ao dividir")
    ap.add_argument("--zip", action="store_true", help="Compacta a saída em ZIP ao final")
    args = ap.parse_args()

    input_dir = Path(args.input_dir)
    out_root = Path(args.saida)

    area = args.area.strip() or input_dir.name

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = out_root / f"pralbino_segmentado_{stamp}"
    textos_dir = out_dir / "textos_por_secao"
    textos_dir.mkdir(parents=True, exist_ok=True)

    secoes_all: List[Secao] = []
    docxs = sorted(input_dir.glob("*.docx"))
    if not docxs:
        raise SystemExit(f"[ERRO] Nenhum .docx encontrado em: {input_dir}")

    print(f"[INFO] DOCX encontrados: {len(docxs)} | Área: {area}")
    for i, docx_path in enumerate(docxs, start=1):
        print(f"[INFO] ({i}/{len(docxs)}) Segmentando: {docx_path.name}")
        secoes = extrair_secoes_docx(
            docx_path=docx_path,
            area=area,
            artigo_slug=None,
            split_long=args.split_long,
            max_words_section=args.max_words_section,
            split_target_words=args.split_target_words,
        )
        secoes_all.extend(secoes)

        # Escreve txt por seção
        for s in secoes:
            (textos_dir / f"{s.secao_id}.txt").write_text(s.texto, encoding="utf-8")

    # JSON
    json_path = out_dir / "secoes.json"
    json_path.write_text(json.dumps([asdict(s) for s in secoes_all], ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] JSON: {json_path}")

    # CSV
    csv_path = out_dir / "secoes.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(asdict(secoes_all[0]).keys()))
        w.writeheader()
        for s in secoes_all:
            d = asdict(s)
            d["texto"] = d["texto"][:200].replace("\n", " ")  # preview só
            w.writerow(d)
    print(f"[OK] CSV: {csv_path}")

    # ZIP opcional
    if args.zip:
        zip_path = out_dir.with_suffix(".zip")
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for p in out_dir.rglob("*"):
                z.write(p, p.relative_to(out_dir.parent))
        print(f"[OK] ZIP: {zip_path}")

    print(f"[OK] Seções geradas: {len(secoes_all)}")
    print(f"[OK] Pasta: {out_dir}")

if __name__ == "__main__":
    main()