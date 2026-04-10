import os
import re
import sys
import argparse
import datetime
import unicodedata
from io import BytesIO
from pathlib import Path
from imapclient import IMAPClient
import email
from email.header import decode_header
import subprocess

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

try:
    from docx import Document
except ImportError:
    Document = None


# =========================
# UTIL: localizar .env.local subindo pastas
# =========================


# =========================
# UTIL: localizar raiz do repo (manage.py/.git)
# =========================
def find_repo_root(start: Path, max_levels: int = 10) -> Path | None:
    cur = start
    for _ in range(max_levels + 1):
        if (cur / 'manage.py').exists() or (cur / '.git').exists():
            return cur
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


def resolve_data_root(arg_value: str | None, scripts_dir: Path) -> Path:
    """
    Resolve a pasta 'anexos_filtrados' (DataRoot).
    - Se o usuário passou --data-root: usa ele
    - Senão: tenta <RepoRoot>/Apenas_Local/anexos_filtrados
    - Fallback: comportamento legado (scripts_dir.parent)
    """
    if arg_value and str(arg_value).strip():
        return Path(arg_value).expanduser().resolve()

    repo = find_repo_root(scripts_dir)
    if repo:
        cand = (repo / 'Apenas_Local' / 'anexos_filtrados')
        if cand.exists():
            return cand.resolve()

    return scripts_dir.parent.resolve()
def find_upwards(start: Path, filename: str, max_levels: int = 6) -> Path | None:
    cur = start
    for _ in range(max_levels + 1):
        cand = cur / filename
        if cand.exists():
            return cand
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


# =========================
# FUNÇÕES
# =========================
def _parse_data(s: str) -> datetime.date:
    s = s.strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    raise ValueError("Formato inválido. Use AAAA-MM-DD (ex.: 2026-02-01) ou DD/MM/AAAA.")


def perguntar_periodo(data_ini_arg: str | None = None, data_fim_arg: str | None = None):
    if data_ini_arg:
        data_ini = _parse_data(data_ini_arg)
    else:
        while True:
            ini_txt = input("Data INICIAL (AAAA-MM-DD) [vazio = sem limite]: ").strip()
            if not ini_txt:
                data_ini = None
                break
            try:
                data_ini = _parse_data(ini_txt)
                break
            except ValueError as e:
                print("  ❌", e)

    if data_fim_arg:
        data_fim = _parse_data(data_fim_arg)
    else:
        while True:
            fim_txt = input("Data FINAL   (AAAA-MM-DD) [vazio = hoje]: ").strip()
            if not fim_txt:
                data_fim = datetime.date.today()
                break
            try:
                data_fim = _parse_data(fim_txt)
                break
            except ValueError as e:
                print("  ❌", e)

    if data_ini and data_ini > data_fim:
        raise ValueError("Data inicial maior que a final.")

    return data_ini, data_fim


def nome_unico(path: Path) -> Path:
    if not path.exists():
        return path
    base = path.with_suffix("")
    ext = path.suffix
    i = 2
    while True:
        cand = Path(f"{base}_{i}{ext}")
        if not cand.exists():
            return cand
        i += 1


def strip_sm_code(s: str) -> str:
    return re.sub(r"^\s*sm\s*\d+\s*[-–—:]\s*|\s*^\s*sm\s*\d+\s*", "", s or "", flags=re.IGNORECASE).strip()


def strip_prefixos_sujos(s: str) -> str:
    s = (s or "").strip()
    # remove prefixos de data no começo
    s = re.sub(r"^\s*\d{4}-\d{2}-\d{2}\s*[-–—:]\s*", "", s)
    s = re.sub(r"^\s*\d{2}-\d{2}-\d{4}\s*[-–—:]\s*", "", s)
    s = strip_sm_code(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def safe_filename(name: str, max_len: int = 140) -> str:
    name = strip_prefixos_sujos(name)
    name = re.sub(r'[<>:"/\\|?*]', "", name).strip().rstrip(".")
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "SEM_TITULO"
    return name[:max_len] if len(name) > max_len else name


def docx_best_title_from_bytes(data: bytes) -> str | None:
    if not Document:
        return None
    try:
        doc = Document(BytesIO(data))
        t = (doc.core_properties.title or "").strip()
        t = strip_prefixos_sujos(t)
        if t:
            return t
        for p in doc.paragraphs:
            txt = strip_prefixos_sujos(p.text or "")
            if txt:
                return txt
    except Exception:
        return None
    return None


def set_docx_internal_title(path: Path, title: str) -> None:
    if not Document:
        return
    try:
        doc = Document(str(path))
        doc.core_properties.title = title
        doc.save(str(path))
    except Exception:
        pass


def decode_filename(raw_filename: str) -> str:
    decoded_parts = decode_header(raw_filename)
    out = []
    for chunk, enc in decoded_parts:
        if isinstance(chunk, (bytes, bytearray)):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(chunk)
    return "".join(out)


def normalize_esboco_filename(name: str) -> str:
    # Se for esboço, padroniza
    low = (name or "").lower()
    if "esbo" in low:
        return "ESBOCO"
    return name


# =========================
# MAIN
# =========================
def main():
    ap = argparse.ArgumentParser(description="Baixa anexos do Pr. Albino (DOCX + ESBOCO) por período e salva em YYYY-MM-DD.")
    ap.add_argument("--ini", help="Data inicial AAAA-MM-DD (opcional).")
    ap.add_argument("--fim", help="Data final AAAA-MM-DD (opcional).")
    ap.add_argument("--data-root", default="", help="Pasta 'anexos_filtrados' (DataRoot). Se vazio, tenta repo/Apenas_Local/anexos_filtrados.")
    ap.add_argument("--remetente", default="pralbino@gmail.com")
    ap.add_argument("--so", nargs="*", default=[".docx", ".txt"], help="Extensões permitidas (default: .docx .txt).")
    ap.add_argument("--nao-consolidar", action="store_true", help="Não executa consolidação automática após baixar.")
    ap.add_argument("--nao-prompts", action="store_true", help="Não gera prompts automaticamente após consolidar.")
    args = ap.parse_args()

    # Caminhos (script dentro de ...\\anexos_filtrados\\Scripts)
    SCRIPTS_DIR = Path(__file__).resolve().parent
    BASE_DIR = resolve_data_root(args.data_root, SCRIPTS_DIR)
    OUTPUT_BASE = BASE_DIR

    # Carrega .env.local (subindo até raiz do projeto)
    if load_dotenv:
        env_path = find_upwards(SCRIPTS_DIR, ".env.local", max_levels=6)
        if env_path:
            load_dotenv(env_path, override=True)

    EMAIL_USER = os.getenv("EMAIL_USER") or "wygazeta@gmail.com"
    EMAIL_PASS = os.getenv("EMAIL_PASS")

    if not EMAIL_PASS:
        raise RuntimeError("EMAIL_PASS não encontrado. Coloque no .env.local (EMAIL_PASS=... senha de app ...)")

    EXTENSOES_PERMITIDAS = [e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.so]

    print("\n📅 Período para buscar anexos (deixe em branco para usar padrão):")
    data_ini, data_fim = perguntar_periodo(args.ini, args.fim)

    OUT_DIR = OUTPUT_BASE / data_fim.strftime("%Y-%m-%d")
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    criteria = ["FROM", args.remetente]
    if data_ini:
        criteria += ["SINCE", data_ini.strftime("%d-%b-%Y")]
    # BEFORE é exclusivo → soma 1 dia para incluir a data final
    criteria += ["BEFORE", (data_fim + datetime.timedelta(days=1)).strftime("%d-%b-%Y")]

    print("\n🔎 Filtro IMAP:", criteria)
    print("📁 Saída:", OUT_DIR)

    salvos = 0

    with IMAPClient("imap.gmail.com") as server:
        server.login(EMAIL_USER, EMAIL_PASS)
        server.select_folder("INBOX", readonly=True)

        messages = server.search(criteria)
        print(f"📨 E-mails encontrados: {len(messages)}")

        for uid, message_data in server.fetch(messages, ["RFC822"]).items():
            email_message = email.message_from_bytes(message_data[b"RFC822"])

            for part in email_message.walk():
                if part.get_content_maintype() == "multipart":
                    continue
                if part.get("Content-Disposition") is None:
                    continue

                raw = part.get_filename()
                if not raw:
                    continue

                filename = decode_filename(raw)
                ext = os.path.splitext(filename)[1].lower()
                nome_limpo = os.path.splitext(filename)[0]

                if ext not in EXTENSOES_PERMITIDAS:
                    continue

                payload = part.get_payload(decode=True) or b""

                if ext == ".docx":
                    titulo_doc = docx_best_title_from_bytes(payload)
                    base_name = safe_filename(titulo_doc or nome_limpo).upper()
                    out_path = nome_unico(OUT_DIR / f"{base_name}{ext}")
                    with open(out_path, "wb") as f:
                        f.write(payload)
                    set_docx_internal_title(out_path, base_name)
                    salvos += 1
                    print(f"📎 Salvo: {out_path.name}")
                    continue

                # TXT (principalmente ESBOCO)
                if ext == ".txt":
                    base = normalize_esboco_filename(safe_filename(nome_limpo))
                    out_path = nome_unico(OUT_DIR / f"{base}.txt")
                    with open(out_path, "wb") as f:
                        f.write(payload)
                    salvos += 1
                    print(f"📎 Salvo: {out_path.name}")
                    continue

    print(f"\n✅ Finalizado. Anexos salvos: {salvos}")

    # Pós-etapas (consolidar + prompts)
    if not args.nao_consolidar:
        resp = input("\n➡️  Consolidar este lote em uma SÉRIE agora? [S/n]: ").strip().lower()
        if resp in ("", "s", "sim"):
            cons_script = SCRIPTS_DIR / "consolidar_serie_por_esboco.py"
            subprocess.run([sys.executable, str(cons_script), "--lote", OUT_DIR.name, "--data-root", str(BASE_DIR)], check=False)

            if not args.nao_prompts:
                last_series = (SCRIPTS_DIR / ".last_series.txt")
                series_name = last_series.read_text(encoding="utf-8").strip() if last_series.exists() else ""
                if series_name:
                    gen_script = SCRIPTS_DIR / "gerar_prompts_imagens.py"
                    subprocess.run([sys.executable, str(gen_script), "--series", series_name, "--data-root", str(BASE_DIR)], check=False)
                else:
                    print("⚠️  Não encontrei .last_series.txt; rode gerar_prompts_imagens.py manualmente.")

if __name__ == "__main__":
    main()
