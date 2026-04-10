import os
import re
import sys
import argparse
import datetime
import unicodedata
from io import BytesIO
from pathlib import Path
from imapclient import IMAPClient
import email
from email.header import decode_header

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

try:
    from docx import Document
except ImportError:
    Document = None


# =========================
# UTIL: localizar .env.local subindo pastas
# =========================
def find_upwards(start: Path, filename: str, max_levels: int = 6) -> Path | None:
    cur = start
    for _ in range(max_levels + 1):
        cand = cur / filename
        if cand.exists():
            return cand
        if cur.parent == cur:
            break
        cur = cur.parent
    return None


# =========================
# NORMALIZACAO / TEXTO
# =========================
def strip_accents(s: str) -> str:
    s = s or ""
    return "".join(
        ch for ch in unicodedata.normalize("NFKD", s)
        if not unicodedata.combining(ch)
    )


def normalize_for_match(s: str) -> str:
    s = strip_accents(s).lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s


def contains_any_term(text: str, termos: list[str]) -> bool:
    norm = normalize_for_match(text)
    return any(normalize_for_match(t) in norm for t in termos)


def _parse_data(s: str) -> datetime.date:
    s = s.strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    raise ValueError("Formato invalido. Use AAAA-MM-DD (ex.: 2026-02-01) ou DD/MM/AAAA.")


def perguntar_periodo(data_ini_arg: str | None = None, data_fim_arg: str | None = None):
    if data_ini_arg:
        data_ini = _parse_data(data_ini_arg)
    else:
        while True:
            ini_txt = input("Data INICIAL (AAAA-MM-DD) [vazio = 2025-07-01]: ").strip()
            if not ini_txt:
                data_ini = datetime.date(2025, 7, 1)
                break
            try:
                data_ini = _parse_data(ini_txt)
                break
            except ValueError as e:
                print("  ERRO:", e)

    if data_fim_arg:
        data_fim = _parse_data(data_fim_arg)
    else:
        while True:
            fim_txt = input("Data FINAL   (AAAA-MM-DD) [vazio = 2026-02-25]: ").strip()
            if not fim_txt:
                data_fim = datetime.date(2026, 2, 25)
                break
            try:
                data_fim = _parse_data(fim_txt)
                break
            except ValueError as e:
                print("  ERRO:", e)

    if data_ini > data_fim:
        raise ValueError("Data inicial maior que a final.")

    return data_ini, data_fim


def nome_unico(path: Path) -> Path:
    if not path.exists():
        return path
    base = path.with_suffix("")
    ext = path.suffix
    i = 2
    while True:
        cand = Path(f"{base}_{i}{ext}")
        if not cand.exists():
            return cand
        i += 1


def strip_sm_code(s: str) -> str:
    return re.sub(r"^\s*sm\s*\d+\s*[-–—:]\s*|\s*^\s*sm\s*\d+\s*", "", s or "", flags=re.IGNORECASE).strip()


def strip_prefixos_sujos(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^\s*\d{4}-\d{2}-\d{2}\s*[-–—:]\s*", "", s)
    s = re.sub(r"^\s*\d{2}-\d{2}-\d{4}\s*[-–—:]\s*", "", s)
    s = strip_sm_code(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def safe_filename(name: str, max_len: int = 160) -> str:
    name = strip_prefixos_sujos(name)
    name = re.sub(r'[<>:"/\\|?*]', "", name).strip().rstrip(".")
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "SEM_TITULO"
    return name[:max_len] if len(name) > max_len else name


def decode_filename(raw_filename: str) -> str:
    decoded_parts = decode_header(raw_filename)
    out = []
    for chunk, enc in decoded_parts:
        if isinstance(chunk, (bytes, bytearray)):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(chunk)
    return "".join(out)


def decode_mime_header(value: str | None) -> str:
    if not value:
        return ""
    decoded_parts = decode_header(value)
    out = []
    for chunk, enc in decoded_parts:
        if isinstance(chunk, (bytes, bytearray)):
            out.append(chunk.decode(enc or "utf-8", errors="replace"))
        else:
            out.append(chunk)
    return "".join(out)


# =========================
# DOCX
# =========================
def docx_best_title_from_bytes(data: bytes) -> str | None:
    if not Document:
        return None
    try:
        doc = Document(BytesIO(data))

        t = (doc.core_properties.title or "").strip()
        t = strip_prefixos_sujos(t)
        if t:
            return t

        for p in doc.paragraphs:
            txt = strip_prefixos_sujos(p.text or "")
            if txt:
                return txt
    except Exception:
        return None
    return None


def set_docx_internal_title(path: Path, title: str) -> None:
    if not Document:
        return
    try:
        doc = Document(str(path))
        doc.core_properties.title = title
        doc.save(str(path))
    except Exception:
        pass


# =========================
# CRITERIO DE ESBOCO
# =========================
def anexo_eh_esboco(
    filename: str,
    ext: str,
    payload: bytes,
    termos: list[str],
    incluir_assunto: bool,
    assunto_email: str,
) -> tuple[bool, str, str | None]:
    """
    Retorna:
      (eh_esboco, motivo, titulo_docx)
    motivo: "nome_anexo", "titulo_interno", "assunto_email"
    """
    nome_sem_ext = os.path.splitext(filename)[0]

    if contains_any_term(nome_sem_ext, termos):
        return True, "nome_anexo", None

    if incluir_assunto and contains_any_term(assunto_email, termos):
        return True, "assunto_email", None

    if ext == ".docx":
        titulo_docx = docx_best_title_from_bytes(payload)
        if titulo_docx and contains_any_term(titulo_docx, termos):
            return True, "titulo_interno", titulo_docx
        return False, "", titulo_docx

    return False, "", None


# =========================
# MAIN
# =========================
def main():
    ap = argparse.ArgumentParser(
        description="Baixa APENAS anexos de ESBOCO do Pr. Albino por periodo, filtrando por nome do arquivo e/ou titulo interno do DOCX."
    )
    ap.add_argument("--ini", default="2025-07-01", help="Data inicial AAAA-MM-DD (default: 2025-07-01).")
    ap.add_argument("--fim", default="2026-02-25", help="Data final AAAA-MM-DD (default: 2026-02-25).")
    ap.add_argument("--remetente", default="pralbino@gmail.com")
    ap.add_argument("--so", nargs="*", default=[".docx", ".txt"], help="Extensoes permitidas (default: .docx .txt).")
    ap.add_argument("--termos", nargs="*", default=["esboço", "esboco", "esbo"], help="Termos que caracterizam esboco.")
    ap.add_argument("--incluir-assunto", action="store_true", help="Tambem considera o assunto do e-mail no filtro.")
    args = ap.parse_args()

    # Caminhos (script dentro de ...\\anexos_filtrados\\Scripts)
    SCRIPTS_DIR = Path(__file__).resolve().parent
    BASE_DIR = SCRIPTS_DIR.parent
    OUTPUT_BASE = BASE_DIR / "ESBOCOS_FILTRADOS"

    if load_dotenv:
        env_path = find_upwards(SCRIPTS_DIR, ".env.local", max_levels=6)
        if env_path:
            load_dotenv(env_path, override=True)

    EMAIL_USER = os.getenv("EMAIL_USER") or "wygazeta@gmail.com"
    EMAIL_PASS = os.getenv("EMAIL_PASS")

    if not EMAIL_PASS:
        raise RuntimeError("EMAIL_PASS nao encontrado. Coloque no .env.local (EMAIL_PASS=... senha de app ...)")

    EXTENSOES_PERMITIDAS = [e.lower() if e.startswith(".") else f".{e.lower()}" for e in args.so]

    print("\nPeriodo para buscar anexos de ESBOCO:")
    data_ini, data_fim = perguntar_periodo(args.ini, args.fim)

    lote_nome = f"{data_ini.strftime('%Y-%m-%d')}_a_{data_fim.strftime('%Y-%m-%d')}"
    OUT_DIR = OUTPUT_BASE / lote_nome
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    criteria = ["FROM", args.remetente]
    criteria += ["SINCE", data_ini.strftime("%d-%b-%Y")]
    criteria += ["BEFORE", (data_fim + datetime.timedelta(days=1)).strftime("%d-%b-%Y")]

    print("\nFiltro IMAP:", criteria)
    print("Saida:", OUT_DIR)
    print("Termos de busca:", args.termos)
    print("Considerar assunto do email:", "SIM" if args.incluir_assunto else "NAO")

    salvos = 0
    emails_encontrados = 0
    anexos_avaliados = 0
    anexos_esboco = 0

    with IMAPClient("imap.gmail.com") as server:
        server.login(EMAIL_USER, EMAIL_PASS)
        server.select_folder("INBOX", readonly=True)

        messages = server.search(criteria)
        emails_encontrados = len(messages)
        print(f"\nEmails encontrados: {emails_encontrados}")

        for uid, message_data in server.fetch(messages, ["RFC822"]).items():
            email_message = email.message_from_bytes(message_data[b"RFC822"])
            assunto_email = decode_mime_header(email_message.get("Subject", ""))

            for part in email_message.walk():
                if part.get_content_maintype() == "multipart":
                    continue
                if part.get("Content-Disposition") is None:
                    continue

                raw = part.get_filename()
                if not raw:
                    continue

                filename = decode_filename(raw)
                ext = os.path.splitext(filename)[1].lower()

                if ext not in EXTENSOES_PERMITIDAS:
                    continue

                anexos_avaliados += 1
                payload = part.get_payload(decode=True) or b""

                eh_esboco, motivo, titulo_docx = anexo_eh_esboco(
                    filename=filename,
                    ext=ext,
                    payload=payload,
                    termos=args.termos,
                    incluir_assunto=args.incluir_assunto,
                    assunto_email=assunto_email,
                )

                if not eh_esboco:
                    continue

                anexos_esboco += 1

                nome_sem_ext = os.path.splitext(filename)[0]

                if ext == ".docx":
                    base_name = safe_filename(titulo_docx or nome_sem_ext).upper()
                    if "ESBO" not in normalize_for_match(base_name).upper():
                        base_name = f"ESBOCO - {base_name}"
                    out_path = nome_unico(OUT_DIR / f"{base_name}{ext}")
                    with open(out_path, "wb") as f:
                        f.write(payload)
                    set_docx_internal_title(out_path, base_name)
                    salvos += 1
                    print(f"Salvo DOCX: {out_path.name}   [motivo={motivo}]")
                    continue

                if ext == ".txt":
                    base_name = safe_filename(nome_sem_ext).upper()
                    if "ESBO" not in normalize_for_match(base_name).upper():
                        base_name = f"ESBOCO - {base_name}"
                    out_path = nome_unico(OUT_DIR / f"{base_name}.txt")
                    with open(out_path, "wb") as f:
                        f.write(payload)
                    salvos += 1
                    print(f"Salvo TXT : {out_path.name}   [motivo={motivo}]")
                    continue

    print("\nFinalizado.")
    print(f"Emails encontrados : {emails_encontrados}")
    print(f"Anexos avaliados   : {anexos_avaliados}")
    print(f"Esbocos detectados : {anexos_esboco}")
    print(f"Arquivos salvos    : {salvos}")
    print(f"Pasta de saida     : {OUT_DIR}")


if __name__ == "__main__":
    main()