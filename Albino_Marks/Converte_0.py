import os
from docx import Document
import html

def convert_docx_to_html(input_directory, output_directory):
    for filename in os.listdir(input_directory):
        if filename.endswith(".docx"):
            docx_path = os.path.join(input_directory, filename)
            html_path = os.path.join(output_directory, filename.replace(".docx", ".html"))

            document = Document(docx_path)
            html_content = "<html><body>"

            # Adiciona parágrafos
            for para in document.paragraphs:
                html_content += "<p>"
                for run in para.runs:
                    if run.bold:
                        html_content += "<b>"
                    if run.italic:
                        html_content += "<i>"
                    html_content += html.escape(run.text)
                    if run.italic:
                        html_content += "</i>"
                    if run.bold:
                        html_content += "</b>"
                html_content += "</p>"

            # Adiciona tabelas
            for table in document.tables:
                html_content += "<table border='1'>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += "<td>"
                        for para in cell.paragraphs:
                            html_content += "<p>"
                            for run in para.runs:
                                if run.bold:
                                    html_content += "<b>"
                                if run.italic:
                                    html_content += "<i>"
                                html_content += html.escape(run.text)
                                if run.italic:
                                    html_content += "</i>"
                                if run.bold:
                                    html_content += "</b>"
                            html_content += "</p>"
                        html_content += "</td>"
                    html_content += "</tr>"
                html_content += "</table>"

            # Adiciona imagens
            image_index = 1
            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    img_path = os.path.join(output_directory, f"image_{image_index}.png")
                    with open(img_path, "wb") as img_file:
                        img_file.write(rel.target_part.blob)
                    html_content += f"<img src='image_{image_index}.png' />"
                    image_index += 1

            html_content += "</body></html>"

            with open(html_path, "w", encoding="utf-8") as html_file:
                html_file.write(html_content)

            print(f"Converted {docx_path} to {html_path}")

input_directory = r"C:\Users\Wanderley\Apps\Albino_Marks\A_Lei_no_NT\Templates\A_Lei_no_NT\Textos\DOCX"
output_directory = r"C:\Users\Wanderley\Apps\Albino_Marks\A_Lei_no_NT\Templates\A_Lei_no_NT\Textos"
convert_docx_to_html(input_directory, output_directory)
